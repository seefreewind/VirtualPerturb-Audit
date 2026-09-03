#!/usr/bin/env python3
"""Build the Cell Reports Methods submission-preparation package.

This script uses only frozen Phase 1/2A/2B/2C reports and tables. It does not
launch new model training, alter splits, or recompute primary metrics.
"""

from __future__ import annotations

import json
import math
import shutil
from datetime import datetime, timezone
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
MANUSCRIPT = ROOT / "manuscript"
SUBMISSION = ROOT / "submission"
FIG_MAIN = ROOT / "figures" / "main"
FIG_SUPP = ROOT / "figures" / "supplementary"
TABLES = ROOT / "results" / "tables"
PACKAGE = SUBMISSION / "cell_reports_methods"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")


PALETTE = {
    "ink": "#1F2933",
    "muted": "#667085",
    "blue": "#2E6F9E",
    "teal": "#3A8F7B",
    "red": "#B65A4A",
    "gold": "#C28A2C",
    "gray": "#D0D5DD",
}


def ensure_dirs() -> None:
    for path in [REPORTS, MANUSCRIPT, SUBMISSION, FIG_MAIN, FIG_SUPP, PACKAGE]:
        path.mkdir(parents=True, exist_ok=True)


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def fmt(x: float | int | str | None, digits: int = 4) -> str:
    if x is None:
        return "NA"
    if isinstance(x, str):
        return x
    if isinstance(x, (float, int)):
        if isinstance(x, float) and (math.isnan(x) or math.isinf(x)):
            return "NA"
        return f"{x:.{digits}f}"
    return str(x)


def md_table(df: pd.DataFrame, cols: list[str] | None = None) -> str:
    if cols:
        df = df[cols]
    out = ["| " + " | ".join(df.columns) + " |", "| " + " | ".join(["---"] * len(df.columns)) + " |"]
    for _, row in df.iterrows():
        vals = []
        for v in row.tolist():
            vals.append(fmt(v) if isinstance(v, (float, int)) else str(v))
        out.append("| " + " | ".join(vals) + " |")
    return "\n".join(out)


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def save_fig(fig, stem: str) -> list[Path]:
    paths = []
    for ext in ("pdf", "svg", "png"):
        p = FIG_MAIN / f"{stem}.{ext}"
        fig.savefig(p, bbox_inches="tight", dpi=300)
        paths.append(p)
    plt.close(fig)
    return paths


def build_figures(data: dict[str, pd.DataFrame]) -> list[str]:
    made: list[str] = []
    plt.rcParams.update({
        "font.family": "DejaVu Sans",
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.labelcolor": PALETTE["ink"],
        "xtick.color": PALETTE["ink"],
        "ytick.color": PALETTE["ink"],
        "text.color": PALETTE["ink"],
    })

    # Figure 1: conceptual audit framework.
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    ax.axis("off")
    stages = [
        ("Frozen inputs", "datasets, splits,\nmodel outputs"),
        ("Metric families", "global fit,\nretrieval,\nUER, sign flips"),
        ("Stress tests", "matched targets,\ncross context,\nprobe controls"),
        ("Bounded claims", "support level,\nlimitations,\nrelease gate"),
    ]
    xs = [0.08, 0.32, 0.56, 0.80]
    colors = [PALETTE["blue"], PALETTE["teal"], PALETTE["gold"], PALETTE["red"]]
    for i, ((title, body), x, c) in enumerate(zip(stages, xs, colors)):
        ax.add_patch(plt.Rectangle((x - 0.095, 0.46), 0.19, 0.30, facecolor="white", edgecolor=c, linewidth=2))
        ax.text(x, 0.68, title, ha="center", va="center", fontsize=11, fontweight="bold")
        ax.text(x, 0.56, body, ha="center", va="center", fontsize=9, color=PALETTE["muted"], linespacing=1.25)
        if i < len(xs) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.11, 0.61), xytext=(x + 0.11, 0.61),
                        arrowprops=dict(arrowstyle="->", lw=1.8, color=PALETTE["muted"]))
    ax.text(0.5, 0.28, "VirtualPerturb-Audit asks whether a perturbation model remains specific and stable when the evaluation target changes.",
            ha="center", fontsize=10)
    ax.text(0.5, 0.18, "Outputs are interpreted as falsification evidence, not as a direct leaderboard across incompatible metric spaces.",
            ha="center", fontsize=9, color=PALETTE["muted"])
    save_fig(fig, "crm_figure1_audit_framework")
    made.append("crm_figure1_audit_framework")

    # Figure 2: Norman/Replogle metric divergence.
    div = data["metric_divergence_profile"].copy()
    div["short"] = div["setting"].str.replace(" GEARS", "", regex=False).str.replace("Replogle ", "Replogle\n", regex=False)
    fig, ax1 = plt.subplots(figsize=(8.2, 4.8))
    x = range(len(div))
    ax1.bar([i - 0.18 for i in x], div["pearson"], width=0.36, color=PALETTE["blue"], label="Pearson")
    ax2 = ax1.twinx()
    ax2.bar([i + 0.18 for i in x], div["mrr"], width=0.36, color=PALETTE["gold"], label="MRR")
    ax1.set_ylabel("Global fit (Pearson)")
    ax2.set_ylabel("Perturbation retrieval (MRR)")
    ax1.set_xticks(list(x), div["short"], rotation=25, ha="right")
    ax1.set_ylim(0.94, 1.0)
    ax2.set_ylim(0, max(0.36, div["mrr"].max() * 1.25))
    ax1.set_title("High global agreement can mask weak perturbation specificity")
    lines = [plt.Rectangle((0, 0), 1, 1, color=PALETTE["blue"]), plt.Rectangle((0, 0), 1, 1, color=PALETTE["gold"])]
    ax1.legend(lines, ["Pearson", "MRR"], frameon=False, loc="upper right")
    save_fig(fig, "crm_figure2_norman_metric_divergence")
    made.append("crm_figure2_norman_metric_divergence")

    # Figure 3: Replogle within-context GEARS and probes.
    probes = data["replogle_gears_vs_probes"].copy()
    probes = probes[probes["model"].isin(["B1_global_perturbed_mean", "FP3_label_shuffled_mean_effect", "GEARS_cell_gears_0.1.2"])]
    probes["label"] = probes["context"] + "\n" + probes["model"].map({
        "B1_global_perturbed_mean": "mean",
        "FP3_label_shuffled_mean_effect": "shuffled",
        "GEARS_cell_gears_0.1.2": "GEARS",
    })
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.2), sharex=False)
    colors = [PALETTE["teal"] if "GEARS" in m else PALETTE["gray"] if "B1" in m else PALETTE["red"] for m in probes["model"]]
    axes[0].bar(range(len(probes)), probes["pearson_delta"], color=colors)
    axes[0].set_ylabel("Audit-delta Pearson")
    axes[0].set_title("Global perturbation effect")
    axes[1].bar(range(len(probes)), probes["retrieval_mrr"], color=colors)
    axes[1].set_ylabel("MRR")
    axes[1].set_title("Perturbation-specific retrieval")
    for ax in axes:
        ax.set_xticks(range(len(probes)), probes["label"], rotation=35, ha="right", fontsize=8)
    fig.suptitle("Replogle within-context audit separates mean-effect fit from target retrieval", y=1.02)
    save_fig(fig, "crm_figure3_replogle_within_context")
    made.append("crm_figure3_replogle_within_context")

    # Figure 4: GEARS matched transfer.
    sens = data["replogle_matched_rl1_rl4_sensitivity"]
    pearson = sens[(sens["comparison_role"] == "primary_source_context_comparison") & (sens["metric"] == "pearson_delta")].copy()
    directions = pearson["direction"].str.replace("_within_vs_", "\nvs\n", regex=False)
    fig, ax = plt.subplots(figsize=(8.2, 5.0))
    idx = range(len(pearson))
    ax.bar([i - 0.18 for i in idx], pearson["within_estimate"], width=0.36, color=PALETTE["teal"], label="Within")
    ax.bar([i + 0.18 for i in idx], pearson["cross_estimate"], width=0.36, color=PALETTE["red"], label="Cross")
    for i, row in enumerate(pearson.itertuples()):
        y_text = min(max(row.within_estimate, row.cross_estimate) + 0.04, 0.61)
        ax.text(i, y_text,
                f"drop {row.paired_difference:.3f}\n95% CI {row.ci_low:.3f}-{row.ci_high:.3f}",
                ha="center", fontsize=8)
    ax.axhline(0, lw=1, color=PALETTE["ink"])
    ax.set_xticks(list(idx), directions, fontsize=9)
    ax.set_ylabel("Matched-target Pearson")
    ax.set_ylim(-0.06, 0.68)
    ax.set_title("Matched targets do not rescue GEARS cross-context transfer", pad=16)
    ax.legend(frameon=False, loc="upper left")
    save_fig(fig, "crm_figure4_matched_gears_transfer")
    made.append("crm_figure4_matched_gears_transfer")

    # Figure 5: STATE partial confirmation.
    st = data["state_transfer_drop"].copy()
    st = st[st["metric"].isin(["pearson_delta", "spearman_delta", "cosine_delta", "uer50", "sign_flip_rate"])]
    fig, ax = plt.subplots(figsize=(8.0, 4.4))
    vals = st["mean_drop_source_minus_cross"].to_numpy()
    lows = vals - st["ci95_low"].to_numpy()
    highs = st["ci95_high"].to_numpy() - vals
    colors = [PALETTE["teal"] if v > 0 else PALETTE["red"] for v in vals]
    ax.bar(range(len(st)), vals, color=colors)
    ax.errorbar(range(len(st)), vals, yerr=[lows, highs], fmt="none", color=PALETTE["ink"], capsize=3, lw=1)
    ax.axhline(0, color=PALETTE["ink"], lw=1)
    ax.set_xticks(range(len(st)), st["metric"], rotation=25, ha="right")
    ax.set_ylabel("Within minus cross-context")
    ax.set_title("STATE supports the transfer-drop direction, with endpoint caveats")
    ax.text(0.02, 0.92, "n=15 matched targets", transform=ax.transAxes, fontsize=9, color=PALETTE["muted"])
    save_fig(fig, "crm_figure5_state_partial_confirmation")
    made.append("crm_figure5_state_partial_confirmation")
    return made


def build_krt() -> None:
    rows = [
        ["Software and algorithms", "VirtualPerturb-Audit", "This paper", "Current repository", "Audit framework scripts and frozen result tables"],
        ["Software and algorithms", "GEARS / cell-gears", "Roohani et al., 2024; official implementation", "https://github.com/snap-stanford/GEARS", "Primary perturbation-response model under audit"],
        ["Software and algorithms", "STATE", "Arc Institute / local installation", "See reports/STATE_*", "Second deep-architecture confirmatory model"],
        ["Software and algorithms", "Python", "Python Software Foundation", "3.12 environment used for post-processing", "Analysis and plotting runtime"],
        ["Software and algorithms", "pandas", "The pandas development team", "Local environment", "Table processing"],
        ["Software and algorithms", "matplotlib", "Matplotlib developers", "Local environment", "Figure generation"],
        ["Deposited data", "Norman et al. Perturb-seq", "Norman et al., 2019", "See reports/NORMAN_ACQUISITION_REPORT.md", "Frozen GEARS-compatible mirror used in audit"],
        ["Deposited data", "Replogle et al. Perturb-seq", "Replogle et al., 2022", "Filtered GEARS-compatible essential-screen subset", "Not the complete Figshare+ processed objects"],
        ["Other", "Frozen Phase 2A/2B/2C result tables", "This paper", "results/tables/", "Primary source for manuscript values"],
    ]
    wb = Workbook()
    ws = wb.active
    ws.title = "Key Resources Table"
    headers = ["REAGENT or RESOURCE", "SOURCE", "IDENTIFIER", "AVAILABILITY", "NOTES"]
    ws.append(headers)
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2E6F9E")
    widths = [28, 38, 38, 42, 50]
    for i, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width
    wb.save(MANUSCRIPT / "KEY_RESOURCES_TABLE.xlsx")
    krt_md = pd.DataFrame(rows, columns=headers)
    write(MANUSCRIPT / "KEY_RESOURCES_TABLE.md", "# Key Resources Table\n\n" + md_table(krt_md))


def build_cover_letter() -> None:
    text = f"""# Cover Letter to Cell Reports Methods

Generated: {GENERATED}

Dear Editors,

We are pleased to submit the Article manuscript, "VirtualPerturb-Audit: a falsification framework for perturbation-response model evaluation," for consideration in Cell Reports Methods.

Perturbation-response models are often summarized by global transcriptional agreement, but this single view can miss failures in perturbation specificity and cross-context transfer. VirtualPerturb-Audit provides a reusable evaluation framework that separates global fit, perturbation retrieval, unsupported-effect behavior, sign-flip burden, leakage risk, and matched-target context transfer. The manuscript is framed as a methods paper: the contribution is the audit design and claim discipline, with GEARS and STATE used as worked examples rather than as a direct leaderboard.

The study reports frozen Norman and GEARS-compatible filtered Replogle analyses. Matched-target GEARS tests show strong cross-context degradation in both K562-to-RPE1 and RPE1-to-K562 directions. A full GPU STATE audit provides partial cross-architecture support: matched Replogle targets show a Pearson drop, while retrieval and unsupported-effect endpoints remain mixed in full-summary comparisons. We therefore present a bounded conclusion that the framework can reveal transfer-specific failure modes that aggregate metrics obscure.

All claims in the manuscript preserve the limitations of the current evidence. Replogle analyses use GEARS-compatible filtered essential-screen data rather than the complete Figshare+ processed objects; BNS is unverified; UER is sensitivity-only; and GEARS/STATE absolute values are not treated as a direct leaderboard where target universes and metric spaces differ.

This manuscript has not been submitted elsewhere. Author, conflict-of-interest, funding, and data/code availability details should be finalized by the submitting author before journal submission.

Sincerely,

[Corresponding author name]
"""
    write(SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS.md", text)
    doc = Document()
    for block in text.splitlines():
        if block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        elif block.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(block)
    doc.save(SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS.docx")


def build_manuscript(data: dict[str, pd.DataFrame]) -> None:
    matched = data["replogle_matched_rl1_rl4_sensitivity"]
    k2r_p = matched[(matched["direction"] == "K562_within_vs_K562_to_RPE1") & (matched["metric"] == "pearson_delta")].iloc[0]
    k2r_u = matched[(matched["direction"] == "K562_within_vs_K562_to_RPE1") & (matched["metric"] == "uer50")].iloc[0]
    k2r_s = matched[(matched["direction"] == "K562_within_vs_K562_to_RPE1") & (matched["metric"] == "sign_flip_rate")].iloc[0]
    r2k_p = matched[(matched["direction"] == "RPE1_within_vs_RPE1_to_K562") & (matched["metric"] == "pearson_delta")].iloc[0]
    st = data["state_phase2c_primary_metrics"]
    st_audit = st[st["metric_space"].isin(["audit_delta", "target_control_audit_delta"])]
    state_rows = {r.setting: r for r in st_audit.itertuples()}
    state_drop = data["state_transfer_drop"]
    st_p = state_drop[state_drop["metric"] == "pearson_delta"].iloc[0]
    manuscript = f"""# VirtualPerturb-Audit: a falsification framework for perturbation-response model evaluation

Draft version: CRM_MANUSCRIPT_v1.0

Generated: {GENERATED}

## Author Information

Authors: [To be completed]

Affiliations: [To be completed]

Correspondence: [To be completed]

## Summary

Perturbation-response models are commonly evaluated by aggregate transcriptomic similarity, but this can obscure whether a model identifies the correct perturbation, avoids unsupported effects, and transfers across cellular contexts. We developed VirtualPerturb-Audit as a falsification framework that separates global fit, perturbation-level retrieval, unsupported-effect behavior, sign-flip burden, leakage risk, and matched-target transfer. In frozen GEARS analyses, matched Replogle K562-to-RPE1 transfer fell from Pearson {fmt(k2r_p.within_estimate)} within context to {fmt(k2r_p.cross_estimate)} cross context, with a paired drop of {fmt(k2r_p.paired_difference)}. A second-architecture STATE audit partially reproduced this transfer phenotype on matched targets, with Pearson decreasing by {fmt(st_p.mean_drop_source_minus_cross)}. These results support VirtualPerturb-Audit as a reusable method for stress-testing perturbation-response claims rather than a direct model leaderboard.

## Introduction

Single-cell perturbation screens create a direct setting for testing whether computational models can predict cellular responses to genetic or chemical interventions. A model that performs well in this setting should do more than reproduce the average shape of a transcriptional response. It should preserve perturbation identity, avoid high-confidence unsupported effects, remain robust under split changes, and expose when its predictions stop transferring across cellular contexts.

Many current evaluations still compress model behavior into global similarity metrics. These metrics are useful but incomplete. A prediction can be globally close to the observed expression profile while ranking the wrong perturbation, borrowing strength from mean-effect structure, or degrading sharply when the same perturbation must be inferred in another cellular context. Virtual perturbation models therefore need evaluation frameworks that make failure modes visible before claims move from benchmark performance to biological use.

VirtualPerturb-Audit was designed for this purpose. The framework freezes split definitions and result tables, evaluates perturbation-level outputs across complementary endpoint families, introduces falsification probes, and uses matched-target sensitivity analyses to separate target-composition effects from context-transfer degradation. The method treats evaluation as an audit of claim robustness. It asks which statements survive stricter endpoint definitions and which must be narrowed.

We apply the framework to Norman and GEARS-compatible filtered Replogle perturbation data. GEARS provides the primary worked example across within-context and cross-context settings. STATE provides a second deep-architecture check after full GPU execution became available. The manuscript uses these analyses to demonstrate the framework and its reporting discipline. It does not claim that GEARS and STATE absolute metrics define a universal ranking, because their metric spaces, target universes, and adapter requirements differ.

## Results

### VirtualPerturb-Audit separates model fit into falsifiable endpoint families

VirtualPerturb-Audit organizes perturbation-response evaluation into four linked stages (Figure 1). Frozen inputs define datasets, splits, model outputs, and permitted post-processing. Metric families then separate global transcriptional agreement from perturbation-specific retrieval, unsupported-effect rate, and sign-flip burden. Stress tests ask whether the same conclusion survives matched targets, cross-context inference, and simple probe controls. The final output is a bounded claim with explicit evidence and limitation status.

This structure is intended to prevent a single strong endpoint from carrying claims that it does not support. For example, a high Pearson correlation can support global response similarity, but it does not by itself establish perturbation identity recovery. A low unsupported-effect rate can support one aspect of stability, but it cannot validate biological realism when the null is sensitivity-only. The audit therefore reports endpoint families together and labels unresolved assumptions.

### Norman and Replogle expose divergence between global agreement and target retrieval

In the frozen GEARS comparison, Norman and Replogle showed different behavior when global agreement and perturbation specificity were viewed together (Figure 2). Norman L1 GEARS had raw-space Pearson {fmt(data["norman_replogle_rl1_comparison"].iloc[0].pearson_delta)} and MRR {fmt(data["norman_replogle_rl1_comparison"].iloc[0].retrieval_mrr)}. Replogle K562 R-L1 retained high raw-space Pearson {fmt(data["norman_replogle_rl1_comparison"].iloc[3].pearson_delta)} but had MRR {fmt(data["norman_replogle_rl1_comparison"].iloc[3].retrieval_mrr)}. Replogle RPE1 R-L1 had raw-space Pearson {fmt(data["norman_replogle_rl1_comparison"].iloc[4].pearson_delta)} and MRR {fmt(data["norman_replogle_rl1_comparison"].iloc[4].retrieval_mrr)}.

This divergence shows why the framework reports target retrieval alongside aggregate expression agreement. The Replogle analyses used GEARS-compatible filtered essential-screen data, not the complete Figshare+ processed objects, so the result is framed as filtered-data evidence. Within that scope, the audit shows that global similarity can remain high while perturbation-specific retrieval becomes weak.

### Within-context Replogle tests show why probe controls are needed

The within-context Replogle audit compared GEARS against simple probes and baselines (Figure 3). The point of this analysis is not to demote a deep model because a simple mean-effect estimate can be strong on a global endpoint. It is to identify which part of the signal is perturbation-specific. In K562 and RPE1, mean-effect probes achieved strong audit-delta Pearson while retrieval remained low. GEARS improved some retrieval endpoints, but the absolute retrieval values remained modest.

These findings support a methods claim: perturbation-response evaluation should include falsification probes that reveal when a model is capturing shared response structure rather than target-specific signal. Probe controls also make manuscript wording more precise, because they separate a statement about global expression fit from a statement about perturbation identification.

### Matched-target GEARS analysis supports cross-context transfer collapse

The strongest GEARS transfer result comes from matched-target sensitivity analysis (Figure 4). In K562-to-RPE1, matched-source Pearson decreased from {fmt(k2r_p.within_estimate)} within context to {fmt(k2r_p.cross_estimate)} cross context. The paired difference was {fmt(k2r_p.paired_difference)}, with a 95% interval of [{fmt(k2r_p.ci_low)}, {fmt(k2r_p.ci_high)}]. UER50 increased from {fmt(k2r_u.within_estimate)} to {fmt(k2r_u.cross_estimate)}, and sign-flip rate increased from {fmt(k2r_s.within_estimate)} to {fmt(k2r_s.cross_estimate)}.

The reverse RPE1-to-K562 direction gave the same conclusion. Matched-source Pearson decreased from {fmt(r2k_p.within_estimate)} to {fmt(r2k_p.cross_estimate)}, with a paired drop of {fmt(r2k_p.paired_difference)} and a 95% interval of [{fmt(r2k_p.ci_low)}, {fmt(r2k_p.ci_high)}]. This matched-target design reduces the possibility that transfer degradation is explained only by different test-target composition. The supported conclusion is `MATCHED_SUPPORTS_TRANSFER_COLLAPSE`.

### STATE provides partial cross-architecture support, with endpoint-level caveats

Phase 2C evaluated STATE as an independent deep architecture on four locked tasks. STATE achieved audit-delta Pearson {fmt(state_rows["Norman L1 STATE"].pearson_delta)} for Norman L1, {fmt(state_rows["Norman L2 STATE"].pearson_delta)} for Norman L2, {fmt(state_rows["Replogle K562 R-L1 STATE"].pearson_delta)} for Replogle K562 R-L1, and {fmt(state_rows["Replogle K562 -> RPE1 R-L4 STATE"].pearson_delta)} for Replogle K562-to-RPE1 R-L4 (Figure 5).

Matched Replogle STATE targets gave the clearest cross-architecture signal. Across {int(st_p.n_matched_targets)} shared targets, Pearson decreased from {fmt(st_p.source_mean)} within context to {fmt(st_p.cross_context_mean)} cross context. The mean drop was {fmt(st_p.mean_drop_source_minus_cross)}, with a 95% interval of [{fmt(st_p.ci95_low)}, {fmt(st_p.ci95_high)}]. Spearman and cosine moved in the same direction. UER50 and sign-flip rate also moved toward worse cross-context behavior on matched targets, although the UER50 interval crossed zero.

The STATE result is partial rather than uniform. In full-summary comparisons, STATE R-L4 had higher retrieval MRR and slightly lower UER50 than STATE R-L1, partly reflecting a smaller normalized R-L4 target universe. We therefore state the Phase 2C conclusion as partial cross-architecture support for matched-target transfer degradation, not as a universal claim that all endpoints or all model architectures fail.

## Discussion

VirtualPerturb-Audit reframes perturbation-response evaluation as a claim-stress problem. The framework asks whether a model conclusion remains true when the endpoint changes from global expression agreement to perturbation-specific retrieval, unsupported-effect burden, sign direction, leakage risk, or cross-context transfer. This approach gives editors and reviewers a transparent way to see which claims are supported and which claims depend on a narrow metric choice.

The GEARS worked example illustrates the value of this framing. Raw global agreement remains high in several settings, but target retrieval and matched transfer tell a different story. The matched Replogle analysis is especially informative because it preserves the target set while changing the context-transfer condition. The resulting Pearson drops in both transfer directions are large and have intervals that do not approach zero.

The STATE analysis strengthens but also narrows the interpretation. It shows that the Replogle transfer-drop direction is not confined to one GEARS run, while its mixed endpoint profile prevents a broad architecture-independent failure claim. This is exactly the reporting behavior the audit is meant to enforce: a result can be supportive and still demand narrower language.

Several limitations are permanent in the current submission package. Replogle analyses use GEARS-compatible filtered essential-screen data rather than the complete Figshare+ processed objects. BNS remains unverified, so UER is sensitivity-only. The GEARS R-L4 workflow is a GEARS-compatible cross-context inference adapter using source-context training, target-context control basal inputs, and target-context evaluation; it is not a native cell-line-aware GEARS split. STATE support is partial and not uniform across endpoints. GEARS and STATE absolute metric values should not be treated as a direct leaderboard where metric spaces and target universes differ.

The resulting manuscript is therefore a methods submission, not a model competition. The contribution is a reusable audit structure, a transparent evidence matrix, and a disciplined way to report perturbation-model robustness. Future work should apply the framework to complete Replogle processed objects, verified replicate-derived nulls, additional architectures, and prospective perturbation settings.

## STAR Methods

### Resource availability

#### Lead contact

Lead contact information will be supplied by the corresponding author before submission.

#### Materials availability

This computational study did not generate new physical reagents.

#### Data and code availability

The audit used public Norman perturbation data and GEARS-compatible filtered Replogle essential-screen data. The complete Figshare+ processed Replogle objects were not used in the current frozen analyses. Code-release status, environment gaps, and deposit requirements are audited in `reports/CRM_CODE_RELEASE_AUDIT.md` and `reports/CODE_RELEASE_GATE.md`.

### Method details

#### Audit design

VirtualPerturb-Audit begins by freezing datasets, split definitions, model outputs, and post-processing rules. All subsequent analyses operate on saved result objects or tabular summaries. The framework reports complementary endpoint families rather than optimizing a single scalar score.

#### Datasets and task levels

Norman analyses used frozen GEARS-compatible Norman perturbation splits. Replogle analyses used GEARS-compatible filtered essential-screen data for K562 and RPE1. R-L1 denotes within-context training and evaluation. R-L4 denotes source-context training with target-context control basal inputs and target-context evaluation through a GEARS-compatible cross-context inference adapter.

#### Endpoint definitions

Delta-expression metrics were computed after subtracting the appropriate control mean. The audit reports Pearson, Spearman, RMSE, cosine similarity, retrieval Top1/Top5/MRR, UER@20/50/100, and sign-flip rate where available. UER remains a sensitivity-only endpoint because BNS is unverified.

#### Matched-target sensitivity

Matched-target sensitivity restricts paired comparisons to perturbation targets shared between within-context and cross-context outputs. This analysis separates transfer degradation from target-composition changes. Bootstrap intervals use perturbation-level resampling as reported in frozen result tables.

#### STATE Phase 2C audit

STATE was evaluated after GPU/Linux execution became available. Four locked tasks were run: Norman L1, Norman L2, Replogle K562 R-L1, and Replogle K562-to-RPE1 R-L4. STATE perturbation labels were normalized by collapsing explicit control partners before target-level evaluation.

### Quantification and statistical analysis

Perturbation-level bootstrap intervals are taken from frozen result tables. Matched GEARS transfer uses paired target-level differences. STATE matched-target transfer uses perturbation-level bootstrap intervals over shared targets. No additional model training or new benchmark reruns were performed during CRM manuscript preparation.

## References

Reference metadata and verification status are listed in `reports/REFERENCE_AUDIT.md`. The working reference set includes Norman et al. Science 2019, Replogle et al. Cell 2022, Roohani et al. Nature Biotechnology 2024, PerturBench, Systema, scArchon, VCBench, and related virtual-cell benchmarking work. Items without fully verified bibliographic metadata are flagged for manual reference-manager confirmation before journal upload.

## Figure Legends

**Figure 1. VirtualPerturb-Audit framework.** Frozen inputs are evaluated through complementary metric families, stress-tested under matched targets and cross-context conditions, and translated into bounded claims.

**Figure 2. Global agreement and perturbation-specific retrieval can diverge.** Frozen GEARS results show high raw-space Pearson across Norman and Replogle settings, while retrieval MRR drops in Replogle.

**Figure 3. Within-context Replogle probe controls.** Mean-effect and shuffled probes help distinguish global response structure from perturbation-specific retrieval.

**Figure 4. Matched-target GEARS cross-context transfer.** Shared-target restriction preserves large Pearson drops in both K562-to-RPE1 and RPE1-to-K562 transfer directions.

**Figure 5. STATE partial cross-architecture confirmation.** Matched Replogle STATE targets show lower cross-context agreement, while endpoint-level caveats remain visible.
"""
    write(MANUSCRIPT / "CRM_MANUSCRIPT_v1.0.md", manuscript)

    supplement = f"""# VirtualPerturb-Audit Supplementary Information

Draft version: CRM_SUPPLEMENT_v1.0

Generated: {GENERATED}

## Supplementary Note 1: Frozen Analysis State

CRM preparation used the frozen Phase 2A, Phase 2B, and Phase 2C outputs. No GEARS or STATE model was rerun during this step.

## Supplementary Note 2: Permanent Limitations

- Replogle analyses use GEARS-compatible filtered essential-screen data, not complete Figshare+ processed objects.
- BNS remains `UNVERIFIED`.
- UER remains `sensitivity_only`.
- GEARS R-L4 is a cross-context inference adapter, not a native cell-line-aware GEARS split.
- STATE provides partial support and is not a uniform endpoint-level confirmation.
- GEARS and STATE absolute metrics are not direct leaderboard values where metric spaces and target universes differ.

## Supplementary Table S1: STATE Primary Metrics

{md_table(st_audit[["setting", "split", "metric_space", "n_test_perturbations", "pearson_delta", "retrieval_mrr", "uer50", "sign_flip_rate"]])}

## Supplementary Table S2: STATE Matched Transfer Contrast

{md_table(data["state_transfer_drop"])}

## Supplementary Table S3: GEARS Matched Transfer Sensitivity

{md_table(matched[matched["comparison_role"] == "primary_source_context_comparison"][["direction", "metric", "n_targets", "within_estimate", "cross_estimate", "paired_difference", "ci_low", "ci_high"]])}

## Supplementary Figures

- `figures/supplementary/phase2c_endpoint_heatmap.*`
- `figures/supplementary/phase2c_retrieval_rank_distribution.*`
"""
    write(MANUSCRIPT / "CRM_SUPPLEMENT_v1.0.md", supplement)


def build_front_matter() -> None:
    highlights = """# Highlights

- Audit separates global fit from perturbation-specific model behavior
- Matched targets expose cross-context transfer collapse in GEARS
- STATE gives partial support for the same transfer-degradation signal
- Claim-evidence matrices keep virtual-cell benchmarks interpretable
"""
    write(MANUSCRIPT / "HIGHLIGHTS.md", highlights)
    graphical = """# Graphical Abstract Brief

Design a four-step horizontal graphical abstract: frozen perturbation datasets and model outputs enter VirtualPerturb-Audit; the audit separates global fit, target retrieval, unsupported-effect burden, and sign-flip endpoints; matched-target cross-context stress tests expose transfer degradation; the output is a bounded claim-evidence matrix rather than a direct model leaderboard.

Keep visual text short. Emphasize the audit workflow and the GEARS/STATE worked examples. Do not present GEARS and STATE as a head-to-head ranking.
"""
    write(MANUSCRIPT / "GRAPHICAL_ABSTRACT_BRIEF.md", graphical)
    in_brief = """# In Brief

VirtualPerturb-Audit is a falsification framework for perturbation-response model evaluation. Applied to frozen GEARS and STATE analyses, it shows that global transcriptomic agreement can mask weak perturbation specificity and cross-context transfer degradation. The method reports bounded claim-evidence relationships rather than a direct model leaderboard.
"""
    write(MANUSCRIPT / "IN_BRIEF.md", in_brief)
    author = """# Author Contributions

Use this placeholder until the final author list is fixed.

- Conceptualization: [Name(s)]
- Methodology: [Name(s)]
- Software: [Name(s)]
- Formal analysis: [Name(s)]
- Investigation: [Name(s)]
- Data curation: [Name(s)]
- Visualization: [Name(s)]
- Writing - original draft: [Name(s)]
- Writing - review and editing: [Name(s)]
- Supervision: [Name(s)]
- Project administration: [Name(s)]
- Funding acquisition: [Name(s), if applicable]
"""
    write(MANUSCRIPT / "AUTHOR_CONTRIBUTIONS.md", author)
    doi = """# Declaration of Interests

The authors declare no competing interests.

Final submission note: replace this placeholder if any author has a financial, advisory, employment, patent, data-access, software-licensing, or other relationship that Cell Press requires to be disclosed.
"""
    write(MANUSCRIPT / "DECLARATION_OF_INTERESTS.md", doi)


def build_reports(data: dict[str, pd.DataFrame], made_figs: list[str]) -> None:
    requirements = f"""# Cell Reports Methods Requirements Audit

Generated: {GENERATED}

Sources checked were restricted to official Cell Press, Cell Reports Methods, Elsevier, and STAR Methods pages. Automated page access to several Cell.com author pages was restricted by Cloudflare, so exact journal-specific word limits that could not be read directly are marked for manual confirmation.

## Official Sources Located

| Source | URL | Status |
|---|---|---|
| Cell Reports Methods information for authors | https://www.cell.com/cell-reports-methods/information-for-authors | Located; automated content access restricted |
| Submit manuscript: Cell Reports Methods | https://www.cell.com/cell-reports-methods/information-for-authors/submit-manuscript | Located; search snippet indicates PDF or Word initial submission allowed |
| Revise manuscript: Cell Reports Methods | https://www.cell.com/cell-reports-methods/information-for-authors/revise-manuscript | Located; search snippet states STAR Methods required for acceptance |
| Final submission: Cell Reports Methods | https://www.cell.com/cell-reports-methods/information-for-authors/final-submission | Located; search snippet describes graphical abstract role |
| Cell Press article templates | https://www.cell.com/information-for-authors/article-templates | Located; search snippet notes KRT handling for STAR Methods journals |
| Cell Press resource availability | https://www.cell.com/pb-assets/journals/assets/info-for-authors/resource-availability.html | Located; search snippet notes newly generated items in KRT and DOI expectations |
| Cell Press figure guidelines | https://www.cell.com/information-for-authors/figure-guidelines | Located; search snippet points to graphical abstract guidelines |
| Elsevier Highlights guide | https://www.elsevier.com/researcher/author/tools-and-resources/highlights | Readable; three to four highlights for Cell Press, each 85 characters or fewer including spaces |
| STAR Methods Key Resources Table | https://star-methods.com/ | Readable landing page; KRT is a STAR Methods component |

## Practical Requirements Applied to This Package

- Prepare an Article-format manuscript with Summary, Introduction, Results, Discussion, STAR Methods, data/code availability, and figure legends.
- Include a Key Resources Table as a separate editable file.
- Provide Highlights, In Brief/eTOC-style text, and a Graphical Abstract brief for final-file readiness.
- Include a cover letter in editable form.
- Keep figures available as PDF/SVG/PNG exports for editorial handling.
- Keep resource availability, code availability, and reproducibility limitations explicit.

## Manual Confirmation Before Upload

- Exact Cell Reports Methods Article word limits and display-item limits.
- Whether the journal upload system requires the Key Resources Table as Word, Excel, main manuscript text, or STAR Methods form output at the current submission stage.
- Current final-file graphical abstract dimensions and file-format requirements from Cell Press PDF guidance.
"""
    write(REPORTS / "CELL_REPORTS_METHODS_REQUIREMENTS.md", requirements)

    article_type = """# Article Type Audit

Decision: `ARTICLE_METHODS_FRAME_READY_WITH_MANUAL_FORMAT_CHECK`.

VirtualPerturb-Audit should be submitted as a methods Article. The core contribution is a reusable falsification and reporting framework for perturbation-response model evaluation, demonstrated on frozen GEARS and STATE worked examples. The manuscript should not be framed as a GEARS benchmark paper, STATE benchmark paper, or universal virtual-cell leaderboard.

## Fit Rationale

- The paper introduces a transferable evaluation workflow.
- The strongest novelty is endpoint separation and claim-evidence discipline.
- GEARS and STATE are examples used to demonstrate the workflow.
- The framework produces editor-facing audit artifacts: limitations, code-release gate, reviewer simulation, and claim-evidence matrix.

## Risks

- If framed as a benchmark, the incomplete Replogle data and mixed STATE endpoints become major weaknesses.
- If framed as a method audit, those same constraints become visible boundaries of the worked example.
"""
    write(REPORTS / "ARTICLE_TYPE_AUDIT.md", article_type)

    novelty = """# CRM Novelty Matrix

| Comparator | Overlap | VirtualPerturb-Audit distinction | Risk |
|---|---|---|---|
| Ahlmann-Eltze-style perturbation benchmarks | Benchmarking perturbation prediction | Adds explicit falsification probes, unsupported-effect burden, and claim-evidence gating | Need exact reference verification |
| PerturBench | Standardized datasets, metrics, and model comparison | Emphasizes audit deltas, matched-transfer collapse, and bounded reporting rather than platform benchmarking | High conceptual overlap; position carefully |
| Systema | Perturbation-specific effects and interpretable perturbation landscape evaluation | Adds cross-context matched-target stress tests and permanent limitation labels | Strong related-work comparator |
| scArchon | Scalable benchmarking and biological hallucination framing | Uses narrower frozen worked examples but stronger submission-facing claim audit | Need avoid claiming broader scope |
| VCBench / virtual-cell benchmarks | In-the-wild virtual-cell evaluation and context shifts | Focuses on falsification workflow and manuscript-ready evidence matrices | Field is moving quickly |
| Simple baseline benchmark papers | Linear/mean-effect controls | Integrates baselines into falsification logic and manuscript claim boundaries | Must cite directly |
| scContam or leakage-focused work | Data leakage / contamination risk | Leakage is one component, not the full framework | Include only if directly relevant after reference verification |

## Novelty Statement

VirtualPerturb-Audit contributes a methods-level audit grammar for perturbation-response claims. Its novelty lies in pairing endpoint-specific metrics with matched-target cross-context stress tests and a submission-ready claim-evidence matrix, making limitations part of the result rather than post hoc caveats.
"""
    write(REPORTS / "CRM_NOVELTY_MATRIX.md", novelty)

    editorial = """# CRM Editorial Audit

Decision: `PROMISING_METHODS_SUBMISSION_AFTER_MANUAL_METADATA_AND_CODE_RELEASE_CLEANUP`.

## Strengths

- Clear methods framing with reusable audit artifacts.
- Strong GEARS matched-target evidence for transfer degradation.
- Second-architecture STATE result supports the direction of the transfer phenotype while preserving caveats.
- Claim boundaries are explicit and defensible.

## Main Editorial Risks

- Replogle data are filtered GEARS-compatible essential-screen data, not complete Figshare+ processed objects.
- BNS is unverified, so UER cannot be presented as biological-replicate hallucination.
- STATE support is partial and endpoint-mixed.
- Repository release metadata are incomplete until license, environment, and data availability decisions are finalized.
- Author metadata, funding, declarations, and exact journal formatting limits still require user input.

## Recommended Positioning

Lead with VirtualPerturb-Audit as a reusable falsification framework. Present GEARS and STATE as worked examples showing why global fit, target retrieval, unsupported-effect behavior, and cross-context transfer should be reported separately.
"""
    write(REPORTS / "CRM_EDITORIAL_AUDIT.md", editorial)

    reviewer = """# CRM Reviewer Simulation

## Reviewer 1: Computational Methods

Major concerns likely to focus on whether the audit adds methodological value beyond existing benchmarks. Response strategy: emphasize endpoint separation, matched-target transfer sensitivity, falsification probes, and claim-evidence reporting rather than claiming a larger benchmark platform.

## Reviewer 2: Single-Cell Perturbation Biology

Major concerns likely to focus on Replogle data completeness and biological-null validity. Response strategy: state that the current Replogle evidence is GEARS-compatible filtered essential-screen evidence; keep BNS unverified; interpret UER as sensitivity-only; avoid biological validation language.

## Reviewer 3: Model Evaluation and Reproducibility

Major concerns likely to focus on code release, split freezing, and whether GEARS/STATE comparisons are fair. Response strategy: point to frozen split manifests, no-rerun CRM preparation, reproducibility file, code-release gate, and the explicit statement that absolute GEARS/STATE metrics are not direct leaderboard values.

## Cross-Reviewer Blocking Questions

- Are the exact data objects and split files recoverable by an external reader?
- Is the repository licensed and environment-pinned enough for review?
- Is the methods contribution clear without overclaiming the worked-example results?
"""
    write(REPORTS / "CRM_REVIEWER_SIMULATION.md", reviewer)

    claims = pd.DataFrame([
        ["VirtualPerturb-Audit separates endpoint families", "Framework design, Figure 1, manuscript STAR Methods", "SUPPORTED", "General method claim"],
        ["High global fit can mask weak perturbation retrieval", "Norman/Replogle GEARS comparison; Figure 2", "SUPPORTED_WITH_FILTERED_REPLOGLE_SCOPE", "Do not imply complete Figshare+ Replogle"],
        ["Probe controls identify mean-effect reliance", "Replogle GEARS vs probes; Figure 3", "SUPPORTED", "Interpret as audit signal, not model demotion"],
        ["Matched GEARS transfer collapses across contexts", "Phase 2B matched sensitivity; Figure 4", "SUPPORTED", "Target composition alone cannot explain drop"],
        ["STATE independently confirms all GEARS failures", "Phase 2C", "NOT_SUPPORTED", "Replace with partial cross-architecture support"],
        ["STATE supports matched-target transfer-drop direction", "state_transfer_drop.csv; Figure 5", "SUPPORTED_PARTIAL", "Endpoint caveats remain"],
        ["UER measures biological hallucination rate", "BNS unverified", "NOT_SUPPORTED", "Use sensitivity-only wording"],
        ["GEARS and STATE define a direct leaderboard", "Different metric spaces and target universes", "NOT_SUPPORTED", "Use worked-example wording"],
    ], columns=["Claim", "Evidence", "Status", "Required wording"])
    write(REPORTS / "CLAIM_EVIDENCE_MATRIX.md", "# Claim-Evidence Matrix\n\n" + md_table(claims))

    code_audit = """# CRM Code Release Audit

Decision: `NEEDS_CLEANUP_BEFORE_PUBLIC_RELEASE`.

## Present Assets

- Frozen reports and tables for Phase 1, Phase 2A, Phase 2B, and Phase 2C.
- Scripts for acquisition, GEARS audit runs, STATE post-processing, supplementary package generation, and CRM package generation.
- Figure exports in PDF/SVG/PNG.
- Raw Phase 2C h5ad outputs retained locally and intentionally not copied into the submission package.

## Blocking Release Items

- No top-level LICENSE file was found.
- No top-level `requirements.txt`, `environment.yml`, or `pyproject.toml` was found during audit.
- README is older than the final Phase 2A/2B/2C state and should be updated before repository release.
- Data availability language must state the filtered Replogle scope and identify external datasets precisely.
- Large local raw outputs should remain excluded or moved to a formal data repository if required.

## Safe Next Cleanup

Add a license selected by the user, export a minimal reproducible environment, update README with frozen run commands and limitations, and prepare a release manifest listing exactly which files are included.
"""
    write(REPORTS / "CRM_CODE_RELEASE_AUDIT.md", code_audit)

    gate = """# Code Release Gate

Gate decision: `NOT_READY_FOR_PUBLIC_RELEASE_WITHOUT_USER_DECISIONS`.

## Pass

- Frozen result artifacts are present.
- Submission-facing manuscript and audit files are generated.
- Main figures are exported.

## Hold

- User must choose repository license.
- User must confirm whether raw Phase 2C outputs should be deposited, excluded, or archived privately.
- Environment specification should be pinned before public release.
- Author/funding/declaration metadata remain placeholders.
"""
    write(REPORTS / "CODE_RELEASE_GATE.md", gate)

    refs = """# Reference Audit

Generated: {generated}

This audit records reference metadata gathered from official publisher, PubMed, arXiv, and journal pages during CRM preparation. Items marked `VERIFY_IN_REFERENCE_MANAGER` should be checked in Zotero/EndNote/Crossref before final upload.

## Verified Core References

| Topic | Reference | DOI/URL | Status |
|---|---|---|---|
| GEARS | Roohani Y, Huang K, Leskovec J. Predicting transcriptional outcomes of novel multigene perturbations with GEARS. Nature Biotechnology 42, 927-935 (2024). | https://www.nature.com/articles/s41587-023-01905-6 | VERIFIED_WEB |
| Norman dataset | Norman TM et al. Exploring genetic interaction manifolds constructed from rich single-cell phenotypes. Science 365, 786-793 (2019). | https://doi.org/10.1126/science.aax4438 | VERIFIED_WEB |
| Replogle dataset | Replogle JM et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq. Cell 185, 2559-2575.e28 (2022). | https://doi.org/10.1016/j.cell.2022.05.013 | VERIFIED_WEB |
| Replogle processed data manifest | Figshare+ processed Perturb-seq datasets for Replogle et al. 2022. | https://doi.org/10.25452/figshare.plus.20029387 | VERIFIED_WEB |
| Replogle SRA/GEO manifest | Figshare+ SRA and GEO file manifest for Replogle et al. 2022. | https://doi.org/10.25452/figshare.plus.20022944 | VERIFIED_WEB |
| PerturBench | Wu Y et al. PerturBench: Benchmarking Machine Learning Models for Cellular Perturbation Analysis. arXiv:2408.10609. | https://arxiv.org/abs/2408.10609 | VERIFIED_WEB |
| Systema | Viñas Torné R et al. Systema: a framework for evaluating genetic perturbation response prediction. Nature Biotechnology (2025). | https://www.nature.com/articles/s41587-025-02777-8 | VERIFIED_WEB |
| scArchon | Radig J et al. scArchon: a scalable benchmarking framework for assessing single-cell perturbation models. Genome Biology 27, 162 (2026). | https://doi.org/10.1186/s13059-026-04104-z | VERIFIED_WEB |
| VCBench | Mao X et al. Benchmarking virtual cell models for in-the-wild perturbation response. arXiv:2604.27646. | https://arxiv.org/abs/2604.27646 | VERIFIED_WEB |

## Manual Verification Needed

- Confirm the final accepted citation style required by Cell Reports Methods.
- Confirm whether STATE should be cited as software, manuscript, preprint, repository, or challenge resource based on the final source selected by the authors.
- Add any exact benchmark references requested by editors or reviewers after journal-specific scope review.
""".format(generated=GENERATED)
    write(REPORTS / "REFERENCE_AUDIT.md", refs)

    readiness = f"""# CRM Submission Readiness

Generated: {GENERATED}

Decision: `READY_FOR_SCIENTIFIC_REVIEW_NOT_READY_FOR_FINAL_JOURNAL_UPLOAD`.

## Completed

- Cell Reports Methods requirements audit created.
- CRM manuscript draft created.
- CRM supplement created.
- Highlights, In Brief, graphical abstract brief, author contributions placeholder, declaration placeholder, KRT, cover letter, novelty matrix, editorial audit, reviewer simulation, claim-evidence matrix, reproducibility file, and code-release gates created.
- Main CRM figures generated: {", ".join(made_figs)}.
- Submission package directory created at `submission/cell_reports_methods/`.

## Remaining User-Dependent Items

- Final author list, affiliations, corresponding author, funding, acknowledgements, and declaration details.
- Repository license choice.
- Public data/code repository decision and DOI/accession details.
- Manual Cell Reports Methods formatting confirmation for current upload system requirements.
- Optional professional graphical abstract artwork based on `manuscript/GRAPHICAL_ABSTRACT_BRIEF.md`.

## Scientific Boundary Check

The package keeps the frozen conclusion: GEARS matched-target analyses support cross-context transfer collapse; STATE provides partial cross-architecture support with mixed endpoint caveats; UER is sensitivity-only; BNS remains unverified; Replogle scope is filtered and GEARS-compatible.
"""
    write(REPORTS / "CRM_SUBMISSION_READINESS.md", readiness)

    repro = """# Reproducibility

This file summarizes how to inspect and reproduce the frozen VirtualPerturb-Audit evidence package.

## Frozen Status

CRM preparation did not run new GEARS/STATE benchmarks, alter split files, download complete Replogle Figshare+ objects, or improve model results. It generated manuscript, audit, figure, and submission-organization artifacts from existing Phase 1/2A/2B/2C outputs.

## Core Evidence Files

- `reports/PHASE2A_RL1_FULL_REPORT.md`
- `reports/PHASE2A_RL4_FULL_REPORT.md`
- `reports/PHASE2B_MATCHED_TARGET_SENSITIVITY.md`
- `reports/PHASE2C_DECISION.md`
- `reports/PHASE2C_RESULT_INTERPRETATION.md`
- `results/tables/replogle_matched_rl1_rl4_sensitivity.csv`
- `results/tables/replogle_rl1_rl4_gears_comparison.csv`
- `results/tables/state_phase2c_primary_metrics.csv`
- `results/tables/state_transfer_drop.csv`
- `results/tables/gears_state_primary_comparison.csv`

## CRM Build Command

```bash
environment/state-postprocess-venv/bin/python scripts/build_crm_submission_package.py
```

## Required Cautions

- Replogle data scope: GEARS-compatible filtered essential-screen data.
- BNS: unverified.
- UER: sensitivity-only.
- GEARS R-L4: cross-context inference adapter.
- STATE: partial endpoint-mixed support.
- GEARS/STATE absolute values: not a direct leaderboard.
"""
    write(ROOT / "REPRODUCIBILITY.md", repro)


def build_package() -> None:
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE, ignore_errors=True)
    subdirs = [
        "01_main_manuscript",
        "02_figures",
        "03_supplement",
        "04_cover_letter",
        "05_graphical_abstract",
        "06_highlights_in_brief",
        "07_key_resources",
        "08_code_data_statements",
        "09_editorial_audits",
    ]
    for sd in subdirs:
        (PACKAGE / sd).mkdir(parents=True, exist_ok=True)

    copies = [
        (MANUSCRIPT / "CRM_MANUSCRIPT_v1.0.md", PACKAGE / "01_main_manuscript"),
        (MANUSCRIPT / "CRM_SUPPLEMENT_v1.0.md", PACKAGE / "03_supplement"),
        (SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS.md", PACKAGE / "04_cover_letter"),
        (SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS.docx", PACKAGE / "04_cover_letter"),
        (MANUSCRIPT / "GRAPHICAL_ABSTRACT_BRIEF.md", PACKAGE / "05_graphical_abstract"),
        (MANUSCRIPT / "HIGHLIGHTS.md", PACKAGE / "06_highlights_in_brief"),
        (MANUSCRIPT / "IN_BRIEF.md", PACKAGE / "06_highlights_in_brief"),
        (MANUSCRIPT / "KEY_RESOURCES_TABLE.xlsx", PACKAGE / "07_key_resources"),
        (MANUSCRIPT / "KEY_RESOURCES_TABLE.md", PACKAGE / "07_key_resources"),
        (ROOT / "REPRODUCIBILITY.md", PACKAGE / "08_code_data_statements"),
        (REPORTS / "CRM_CODE_RELEASE_AUDIT.md", PACKAGE / "08_code_data_statements"),
        (REPORTS / "CODE_RELEASE_GATE.md", PACKAGE / "08_code_data_statements"),
    ]
    for src, dst in copies:
        dst.mkdir(parents=True, exist_ok=True)
        if src.exists():
            shutil.copy2(src, dst / src.name)

    for stem in [
        "crm_figure1_audit_framework",
        "crm_figure2_norman_metric_divergence",
        "crm_figure3_replogle_within_context",
        "crm_figure4_matched_gears_transfer",
        "crm_figure5_state_partial_confirmation",
    ]:
        for ext in ("pdf", "svg", "png"):
            src = FIG_MAIN / f"{stem}.{ext}"
            if src.exists():
                shutil.copy2(src, PACKAGE / "02_figures" / src.name)

    for pattern in ["phase2c_endpoint_heatmap.*", "phase2c_retrieval_rank_distribution.*"]:
        for src in FIG_SUPP.glob(pattern):
            if not src.name.startswith("._"):
                shutil.copy2(src, PACKAGE / "03_supplement" / src.name)

    for src in [
        REPORTS / "CELL_REPORTS_METHODS_REQUIREMENTS.md",
        REPORTS / "ARTICLE_TYPE_AUDIT.md",
        REPORTS / "CRM_NOVELTY_MATRIX.md",
        REPORTS / "CRM_EDITORIAL_AUDIT.md",
        REPORTS / "CRM_REVIEWER_SIMULATION.md",
        REPORTS / "CLAIM_EVIDENCE_MATRIX.md",
        REPORTS / "REFERENCE_AUDIT.md",
        REPORTS / "CRM_SUBMISSION_READINESS.md",
    ]:
        if src.exists():
            shutil.copy2(src, PACKAGE / "09_editorial_audits" / src.name)

    for path in PACKAGE.rglob("._*"):
        if path.is_file():
            path.unlink()

    write(PACKAGE / "README_SUBMISSION_PACKAGE.md", """# Cell Reports Methods Submission Package

This directory organizes the CRM manuscript-preparation artifacts. It is not a final journal upload ZIP because author metadata, license, repository DOI, and exact upload-system formatting still require user confirmation.

## Directory Map

- `01_main_manuscript/`: main Markdown manuscript draft.
- `02_figures/`: five CRM main figures in PDF/SVG/PNG.
- `03_supplement/`: supplement Markdown and selected supplementary figures.
- `04_cover_letter/`: cover letter Markdown and DOCX.
- `05_graphical_abstract/`: graphical abstract production brief.
- `06_highlights_in_brief/`: highlights and In Brief text.
- `07_key_resources/`: Key Resources Table.
- `08_code_data_statements/`: reproducibility and code-release gate files.
- `09_editorial_audits/`: requirements, novelty, claim, reviewer, and readiness audits.
""")

    for path in PACKAGE.rglob("._*"):
        if path.is_file():
            path.unlink()

    manifest = []
    for path in sorted(PACKAGE.rglob("*")):
        if path.is_file() and not path.name.startswith("._") and path.name != "PACKAGE_MANIFEST.json":
            manifest.append({"path": str(path.relative_to(PACKAGE)), "bytes": path.stat().st_size})
    write(PACKAGE / "PACKAGE_MANIFEST.json", json.dumps(manifest, indent=2, ensure_ascii=False))

    for path in PACKAGE.rglob("._*"):
        if path.is_file():
            path.unlink()


def main() -> None:
    ensure_dirs()
    data = {
        "replogle_matched_rl1_rl4_sensitivity": read_csv("replogle_matched_rl1_rl4_sensitivity.csv"),
        "state_phase2c_primary_metrics": read_csv("state_phase2c_primary_metrics.csv"),
        "state_transfer_drop": read_csv("state_transfer_drop.csv"),
        "gears_state_primary_comparison": read_csv("gears_state_primary_comparison.csv"),
        "replogle_rl1_rl4_gears_comparison": read_csv("replogle_rl1_rl4_gears_comparison.csv"),
        "norman_replogle_rl1_comparison": read_csv("norman_replogle_rl1_comparison.csv"),
        "metric_divergence_profile": read_csv("metric_divergence_profile.csv"),
        "replogle_gears_vs_probes": read_csv("replogle_gears_vs_probes.csv"),
    }
    made_figs = build_figures(data)
    build_krt()
    build_cover_letter()
    build_manuscript(data)
    build_front_matter()
    build_reports(data, made_figs)
    build_package()
    print("CRM package generated")
    print(PACKAGE)


if __name__ == "__main__":
    main()
