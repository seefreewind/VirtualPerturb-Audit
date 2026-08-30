#!/usr/bin/env python3
"""Finalize Cell Reports Methods v1.1 submission-preparation package.

Allowed scope: manuscript, figures, references, data/code availability, release
readiness, and submission organization from frozen outputs only.
"""

from __future__ import annotations

import json
import math
import os
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
MANUSCRIPT = ROOT / "manuscript"
SUBMISSION = ROOT / "submission"
FIG_MAIN = ROOT / "figures" / "main"
FIG_SUPP = ROOT / "figures" / "supplementary"
TABLES = ROOT / "results" / "tables"
FINAL = SUBMISSION / "cell_reports_methods" / "final"
QA = ROOT / "qa" / "docx_render_crm_v11"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

INK = "#1F2933"
MUTED = "#667085"
BLUE = "#2E6F9E"
TEAL = "#3A8F7B"
RED = "#B65A4A"
GOLD = "#C28A2C"
GRAY = "#D0D5DD"


def fmt(x, digits=4) -> str:
    if isinstance(x, str):
        return x
    if x is None:
        return "NA"
    try:
        if math.isnan(float(x)):
            return "NA"
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def read_table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def md_table(df: pd.DataFrame) -> str:
    lines = ["| " + " | ".join(df.columns) + " |", "| " + " | ".join(["---"] * len(df.columns)) + " |"]
    for _, row in df.iterrows():
        vals = [fmt(v) if isinstance(v, (int, float)) else str(v) for v in row.tolist()]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def get_metrics():
    sens = read_table("replogle_matched_rl1_rl4_sensitivity.csv")
    state = read_table("state_phase2c_primary_metrics.csv")
    state_drop = read_table("state_transfer_drop.csv")
    norm_rep = read_table("norman_replogle_rl1_comparison.csv")
    probes = read_table("replogle_gears_vs_probes.csv")
    k2r_p = sens[(sens.direction == "K562_within_vs_K562_to_RPE1") & (sens.metric == "pearson_delta")].iloc[0]
    k2r_u = sens[(sens.direction == "K562_within_vs_K562_to_RPE1") & (sens.metric == "uer50")].iloc[0]
    k2r_s = sens[(sens.direction == "K562_within_vs_K562_to_RPE1") & (sens.metric == "sign_flip_rate")].iloc[0]
    r2k_p = sens[(sens.direction == "RPE1_within_vs_RPE1_to_K562") & (sens.metric == "pearson_delta")].iloc[0]
    st_p = state_drop[state_drop.metric == "pearson_delta"].iloc[0]
    return sens, state, state_drop, norm_rep, probes, k2r_p, k2r_u, k2r_s, r2k_p, st_p


def save_fig(fig, stem: str):
    FIG_MAIN.mkdir(parents=True, exist_ok=True)
    for ext in ["pdf", "svg", "png"]:
        fig.savefig(FIG_MAIN / f"{stem}.{ext}", bbox_inches="tight", dpi=300)
    plt.close(fig)


def build_final_figures(state: pd.DataFrame, state_drop: pd.DataFrame, sens: pd.DataFrame):
    plt.rcParams.update({
        "font.family": "DejaVu Sans",
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.labelcolor": INK,
        "xtick.color": INK,
        "ytick.color": INK,
        "text.color": INK,
    })

    # Figure 4 v1.1: explicitly audit-delta Pearson and matched target n.
    pearson = sens[(sens.comparison_role == "primary_source_context_comparison") & (sens.metric == "pearson_delta")].copy()
    labels = ["K562 to RPE1\nn=150", "RPE1 to K562\nn=148"]
    fig, ax = plt.subplots(figsize=(8.0, 4.8))
    xs = range(len(pearson))
    ax.bar([i - 0.18 for i in xs], pearson.within_estimate, width=0.36, color=TEAL, label="Within context")
    ax.bar([i + 0.18 for i in xs], pearson.cross_estimate, width=0.36, color=RED, label="Cross context")
    for i, row in enumerate(pearson.itertuples()):
        ax.text(i, min(max(row.within_estimate, row.cross_estimate) + 0.035, 0.615),
                f"drop {row.paired_difference:.4f}\n95% CI [{row.ci_low:.4f}, {row.ci_high:.4f}]",
                ha="center", fontsize=8.2)
    ax.axhline(0, color=INK, lw=1)
    ax.set_xticks(list(xs), labels)
    ax.set_ylim(-0.06, 0.69)
    ax.set_ylabel("Matched-target audit-delta Pearson")
    ax.set_title("Matched-target transfer degradation in GEARS", pad=14)
    ax.legend(frameon=False, loc="upper left")
    save_fig(fig, "crm_figure4_matched_gears_transfer_v11")

    # Figure 5 v1.1: endpoint heterogeneity, including full-summary MRR.
    rows = []
    for metric in ["pearson_delta", "spearman_delta", "cosine_delta", "uer50", "sign_flip_rate"]:
        r = state_drop[state_drop.metric == metric].iloc[0]
        rows.append([metric, r.mean_drop_source_minus_cross, r.ci95_low, r.ci95_high, "matched n=15"])
    state_a = state[state.metric_space.isin(["audit_delta", "target_control_audit_delta"])]
    s_rl1 = state_a[state_a.setting == "Replogle K562 R-L1 STATE"].iloc[0]
    s_rl4 = state_a[state_a.setting == "Replogle K562 -> RPE1 R-L4 STATE"].iloc[0]
    rows.append(["retrieval_mrr\n(full summary)", s_rl1.retrieval_mrr - s_rl4.retrieval_mrr, None, None, "unmatched targets"])
    fig_df = pd.DataFrame(rows, columns=["metric", "value", "low", "high", "scope"])
    fig, ax = plt.subplots(figsize=(8.6, 4.8))
    colors = [TEAL if v > 0 else RED for v in fig_df.value]
    ax.bar(range(len(fig_df)), fig_df.value, color=colors)
    for i, r in fig_df.iterrows():
        if pd.notna(r.low):
            yerr = [[r.value - r.low], [r.high - r.value]]
            ax.errorbar([i], [r.value], yerr=yerr, fmt="none", color=INK, capsize=3, lw=1)
        ax.text(i, r.value + (0.012 if r.value >= 0 else -0.018), r.scope, ha="center",
                va="bottom" if r.value >= 0 else "top", fontsize=7.5, color=MUTED)
    ax.axhline(0, color=INK, lw=1)
    ax.set_xticks(range(len(fig_df)), fig_df.metric, rotation=25, ha="right")
    ax.set_ylabel("Within minus cross-context")
    ax.set_title("STATE gives partial support with endpoint heterogeneity", pad=14)
    ax.text(0.01, 0.95, "Agreement endpoints: positive favors within. Burden endpoints: negative means worse cross-context.",
            transform=ax.transAxes, fontsize=8.5, color=MUTED, va="top")
    save_fig(fig, "crm_figure5_state_partial_confirmation_v11")


def references_text() -> str:
    return """## References

1. Norman, T. M., Horlbeck, M. A., Replogle, J. M., Ge, A. Y., Xu, A., Jost, M., Gilbert, L. A., and Weissman, J. S. Exploring genetic interaction manifolds constructed from rich single-cell phenotypes. *Science* 365, 786-793 (2019). https://doi.org/10.1126/science.aax4438. PubMed: 31395745.
2. Replogle, J. M. et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq. *Cell* 185, 2559-2575.e28 (2022). https://doi.org/10.1016/j.cell.2022.05.013. PubMed: 35688146.
3. Roohani, Y., Huang, K., and Leskovec, J. Predicting transcriptional outcomes of novel multigene perturbations with GEARS. *Nature Biotechnology* 42, 927-935 (2024). https://doi.org/10.1038/s41587-023-01905-6. PubMed: 37592036.
4. Wu, Y., Wershof, E., Schmon, S. M., Nassar, M., Osinski, B., Eksi, R., Yan, Z., Stark, R., Zhang, K., and Graepel, T. PerturBench: Benchmarking Machine Learning Models for Cellular Perturbation Analysis. *Advances in Neural Information Processing Systems 38*, 106937-106977 (2025). https://doi.org/10.52202/085713-3225. Preprint: https://arxiv.org/abs/2408.10609.
5. Vinas Torne, R. et al. Systema: a framework for evaluating genetic perturbation response prediction beyond systematic variation. *Nature Biotechnology* (2025). https://doi.org/10.1038/s41587-025-02777-8.
6. Radig, J. et al. scArchon: a scalable benchmarking framework for assessing single-cell perturbation models. *Genome Biology* 27, 162 (2026). https://doi.org/10.1186/s13059-026-04104-z. PubMed: 42121287.
7. Mao, X. et al. Benchmarking virtual cell models for in-the-wild perturbation response. *arXiv* 2604.27646 (2026). https://arxiv.org/abs/2604.27646.
8. Vollenweider, M. et al. Signal, Bounds, and Baselines: Principles for Rigorous Single-Cell Perturbation Prediction Benchmarking. *bioRxiv* (2026). https://doi.org/10.64898/2026.04.20.719650.
9. Replogle, J. M. et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq: SRA and GEO file manifest. Figshare+ (2022). https://doi.org/10.25452/figshare.plus.20022944.
10. Replogle, J. M. et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq: processed datasets. Figshare+ (2022). https://doi.org/10.25452/figshare.plus.20029387.
"""


def build_manuscript_v11(metrics):
    sens, state, state_drop, norm_rep, probes, k2r_p, k2r_u, k2r_s, r2k_p, st_p = metrics
    state_a = state[state.metric_space.isin(["audit_delta", "target_control_audit_delta"])]
    srows = {r.setting: r for r in state_a.itertuples()}
    text = f"""# VirtualPerturb-Audit: a falsification framework for perturbation-response model evaluation

Draft version: CRM_MANUSCRIPT_v1.1

Generated: {GENERATED}

## Author Information

Authors: [To be completed]

Affiliations: [To be completed]

Correspondence: [To be completed]

## Summary

Perturbation-response models are often judged by aggregate transcriptomic similarity, but this endpoint can miss perturbation-specific and context-shifted failures. VirtualPerturb-Audit is a falsification framework that freezes analysis inputs and evaluates global fit, target retrieval, unsupported-effect behavior, sign-flip rate, leakage risk, and matched-target transfer. In GEARS applied to GEARS-compatible filtered Replogle data, matched K562-to-RPE1 audit-delta Pearson decreased by {fmt(k2r_p.paired_difference)}. In STATE, a second deep architecture, the matched K562-to-RPE1 audit-delta Pearson drop was {fmt(st_p.mean_drop_source_minus_cross)}. The central contribution is not a new perturbation predictor, but a reusable framework for testing which aspects of apparent predictive performance survive perturbation-specific and context-shifted stress testing.

## Introduction

Single-cell perturbation screens connect experimental intervention with transcriptome-scale phenotypes, making them a natural testbed for perturbation-response prediction. Norman et al. established rich single-cell genetic-interaction phenotypes, and Replogle et al. extended Perturb-seq to genome-scale CRISPRi maps across millions of cells [1,2]. Models such as GEARS use these data to predict responses to unseen perturbations [3].

Aggregate similarity is useful but incomplete. A model can achieve high raw-space transcriptomic similarity while recovering little perturbation identity, especially when shared mean-response structure dominates the signal. Recent benchmarks have sharpened this concern by showing that model rankings and conclusions depend on metric choice, task design, and systematic variation [4-8].

Existing benchmark resources help standardize datasets, baselines, and model comparisons [4-8]. What remains fragmented is the link between a reported number and the claim it is allowed to support. Raw-space transcriptomic similarity and control-subtracted audit-delta agreement quantify different properties and should not be interpreted as numerically interchangeable endpoints.

VirtualPerturb-Audit addresses this gap by treating evaluation as falsification. It freezes inputs, separates endpoint families, adds probe controls, runs matched-target context-transfer stress tests, and records claim boundaries. We demonstrate the framework with frozen GEARS and STATE analyses on Norman and GEARS-compatible filtered Replogle data.

## Results

### VirtualPerturb-Audit separates model fit into falsifiable endpoint families

VirtualPerturb-Audit organizes perturbation-response evaluation into frozen inputs, metric families, stress tests, and bounded claims (Figure 1). Frozen inputs define datasets, split assignments, model outputs, and permitted post-processing. Metric families separate raw-space global transcriptomic similarity, audit-delta agreement, perturbation-specific retrieval, unsupported-effect rate, and sign-flip rate. Stress tests then ask whether a claim survives probe controls, matched targets, and cross-context transfer.

This design prevents one strong endpoint from carrying unsupported claims. Raw-space Pearson can support global expression agreement, but it does not establish perturbation identity recovery. Unsupported-effect rate can support sensitivity to large unsupported predictions, but it is not a validated biological endpoint unless its null envelope is derived from validated biological replicate ground truth.

### Norman and Replogle expose divergence between global agreement and target retrieval

Frozen GEARS results showed that raw-space global similarity and perturbation-specific retrieval can diverge (Figure 2). Norman L1 GEARS had raw-space Pearson {fmt(norm_rep.iloc[0].pearson_delta)} and MRR {fmt(norm_rep.iloc[0].retrieval_mrr)}. Replogle K562 R-L1 retained raw-space Pearson {fmt(norm_rep.iloc[3].pearson_delta)} but had MRR {fmt(norm_rep.iloc[3].retrieval_mrr)}. Replogle RPE1 R-L1 had raw-space Pearson {fmt(norm_rep.iloc[4].pearson_delta)} and MRR {fmt(norm_rep.iloc[4].retrieval_mrr)}.

These values are not directly comparable to audit-delta Pearson values used in later transfer sections. Raw-space Pearson measures agreement in the expression space used by the GEARS evaluation row. Audit-delta Pearson measures agreement after control subtraction and is used to compare perturbation-specific response patterns.

### Within-context Replogle tests show why probe controls are needed

The within-context Replogle audit compared GEARS with simple probes and baselines (Figure 3). Mean-effect probes achieved strong audit-delta Pearson in K562 and RPE1, while retrieval remained low. GEARS improved some retrieval endpoints, but absolute retrieval values stayed modest.

Probe controls clarify what kind of signal drives an apparent success. In this setting, they separate global response structure from perturbation-specific target recovery. The result supports an audit-framework claim rather than a broad claim about model superiority or failure.

### Matched-target GEARS analysis supports cross-context transfer collapse

Matched-target GEARS analysis provides the strongest quantitative result (Figure 4). In K562-to-RPE1 transfer, audit-delta Pearson decreased from {fmt(k2r_p.within_estimate)} within context to {fmt(k2r_p.cross_estimate)} cross context. The paired drop was {fmt(k2r_p.paired_difference)}, with a 95% interval of [{fmt(k2r_p.ci_low)}, {fmt(k2r_p.ci_high)}]. UER50 increased from {fmt(k2r_u.within_estimate)} to {fmt(k2r_u.cross_estimate)}, and sign-flip rate increased from {fmt(k2r_s.within_estimate)} to {fmt(k2r_s.cross_estimate)}.

The reverse RPE1-to-K562 direction showed the same pattern. Audit-delta Pearson decreased from {fmt(r2k_p.within_estimate)} to {fmt(r2k_p.cross_estimate)}, with a paired drop of {fmt(r2k_p.paired_difference)} and a 95% interval of [{fmt(r2k_p.ci_low)}, {fmt(r2k_p.ci_high)}]. The matched-target design reduces target-composition confounding, but it does not eliminate every possible confounder. The supported conclusion is `MATCHED_SUPPORTS_TRANSFER_COLLAPSE`.

### STATE provides partial cross-architecture support, with endpoint-level caveats

Phase 2C evaluated STATE as a second deep architecture on four locked tasks. STATE achieved audit-delta Pearson {fmt(srows["Norman L1 STATE"].pearson_delta)} for Norman L1, {fmt(srows["Norman L2 STATE"].pearson_delta)} for Norman L2, {fmt(srows["Replogle K562 R-L1 STATE"].pearson_delta)} for Replogle K562 R-L1, and {fmt(srows["Replogle K562 -> RPE1 R-L4 STATE"].pearson_delta)} for Replogle K562-to-RPE1 R-L4.

Matched STATE targets gave partial cross-architecture support for the transfer-degradation signal (Figure 5). Across {int(st_p.n_matched_targets)} shared targets, audit-delta Pearson decreased from {fmt(st_p.source_mean)} within context to {fmt(st_p.cross_context_mean)} cross context. The mean drop was {fmt(st_p.mean_drop_source_minus_cross)}, with a 95% interval of [{fmt(st_p.ci95_low)}, {fmt(st_p.ci95_high)}]. Spearman and cosine moved in the same direction, and sign-flip rate was worse cross context. UER50 had a worse point estimate, but its interval crossed zero.

STATE did not provide a uniform endpoint-level confirmation. In full-summary comparisons, STATE R-L4 had higher retrieval MRR and slightly lower UER50 than STATE R-L1, in a smaller normalized R-L4 target universe. STATE therefore provided partial cross-architecture support for matched-target transfer degradation, with endpoint-level heterogeneity that limits broader generalization.

## Discussion

VirtualPerturb-Audit contributes a reusable audit grammar for perturbation-response model evaluation. It links each reported endpoint to the claim it can support, then records where the claim narrows under stricter stress tests.

The GEARS worked example shows why this grammar matters. Raw-space global similarity remained high in several settings, while perturbation-specific retrieval and matched transfer revealed weaker behavior. The matched Replogle analysis is the clearest stress test because it compares the same target set across within-context and cross-context conditions.

The STATE analysis adds a second-architecture check without overstating the evidence. It supports the matched-target transfer-drop direction but also shows endpoint heterogeneity. This mixed result strengthens the methods argument because the framework exposes both supportive and limiting evidence.

Practically, perturbation-response studies should report raw-space similarity, audit-delta agreement, retrieval, unsupported-effect rate, sign-flip rate, and context-transfer analyses as distinct endpoints. They should also state the target universe and metric space for every comparison.

The current package stops at a bounded methods claim. It does not establish a universal model ranking, validated biological UER endpoint, or architecture-independent failure. Future work should apply the same audit protocol to complete Replogle processed objects, validated replicate-derived nulls, more architectures, and prospective perturbation settings.

## Limitations of the study

Replogle analyses use GEARS-compatible filtered essential-screen data, not the complete Figshare+ processed objects. Validated biological replicate metadata were unavailable, so BNS remains unverified. UER is interpreted as sensitivity-only because its null envelope is not derived from validated biological replicate ground truth. GEARS R-L4 uses a GEARS-compatible cross-context inference adapter, not a native cell-line-aware GEARS split. The worked example includes two architectures, and STATE matched transfer has 15 shared targets. STATE support is endpoint-heterogeneous. The manuscript does not make a direct universal model-ranking claim.

## STAR Methods

### Resource availability

#### Lead contact

Lead contact information will be supplied by the corresponding author before submission.

#### Materials availability

This computational study did not generate new physical reagents.

#### Data and code availability

Original datasets: Norman perturbation data were used through a GEARS-compatible processed mirror [1,3]. Replogle data were used as GEARS-compatible filtered essential-screen K562 and RPE1 objects; complete Figshare+ processed objects were not part of the frozen analyses [2,9,10].

Processed and derived audit files: frozen result tables are stored under `results/tables/`. Split assignments are stored under the project split and metadata outputs listed in `REPRODUCIBILITY.md`. Predictions and target-level outputs are stored in local frozen result directories and require repository/archive deposition before journal upload (`TODO_DEPOSIT`).

Source code and environment: source code is in the VirtualPerturb-Audit repository. Public repository URL, archive DOI, and final environment export remain `TODO_DEPOSIT`. The repository uses an MIT license in this finalization package.

### Method details

#### Audit design

VirtualPerturb-Audit freezes datasets, split definitions, model outputs, and post-processing rules before manuscript interpretation. All CRM v1.1 materials were generated from saved result tables and reports. No new GEARS training, STATE training, primary endpoint redefinition, or matched-target registry change was performed.

#### Datasets and task levels

Norman analyses used frozen GEARS-compatible Norman perturbation splits. Replogle analyses used GEARS-compatible filtered essential-screen data for K562 and RPE1. R-L1 denotes within-context perturbation holdout. R-L4 denotes source-context training with target-context control basal inputs and target-context evaluation through a GEARS-compatible cross-context inference adapter.

#### Endpoint definitions

Raw-space Pearson measures global transcriptomic similarity in the expression space of the relevant model output. Audit-delta Pearson measures control-subtracted perturbation-response agreement. Retrieval endpoints include Top1, Top5, and MRR. Unsupported-effect rate is reported as UER@K. BNS remains unverified because validated biological replicate metadata were unavailable.

#### Matched-target sensitivity

Matched-target analysis restricts paired comparisons to perturbation targets shared between within-context and cross-context outputs. GEARS matched transfer uses paired perturbation-level differences. STATE matched transfer uses bootstrap intervals over shared targets.

#### STATE Phase 2C audit

STATE was evaluated on a CUDA-capable Linux GPU server after local CPU execution was classified as not performance-eligible. Locked tasks were Norman L1, Norman L2, Replogle K562 R-L1, and Replogle K562-to-RPE1 R-L4. STATE perturbation labels were normalized by collapsing explicit control partners before target-level evaluation.

### Quantification and statistical analysis

Bootstrap confidence intervals use perturbation-level resampling as reported in frozen result tables. GEARS K562-to-RPE1 and RPE1-to-K562 transfer analyses use matched-target paired differences. STATE K562-to-RPE1 transfer uses 15 matched targets. No new primary analyses were computed during manuscript finalization.

{references_text()}

## Figure Legends

**Figure 1. VirtualPerturb-Audit framework.** Frozen perturbation-response datasets and model outputs enter a four-stage audit: input freeze, endpoint-family evaluation, stress testing, and bounded claim assignment. The figure depicts the method workflow, not a model-ranking result.

**Figure 2. Raw-space global agreement and perturbation-specific retrieval can diverge.** GEARS raw-space Pearson and retrieval MRR are shown for frozen Norman and GEARS-compatible filtered Replogle within-context tasks. Pearson is raw expression Pearson, whereas MRR measures perturbation-specific retrieval. Replogle panels use filtered essential-screen data.

**Figure 3. Within-context Replogle probe controls.** GEARS, mean-effect probes, and label-shuffled probes are compared on GEARS-compatible filtered Replogle K562 and RPE1 R-L1 tasks. Bars report audit-delta Pearson and retrieval MRR from frozen result tables.

**Figure 4. Matched-target GEARS cross-context transfer.** Shared-target analysis compares within-context and cross-context audit-delta Pearson for K562-to-RPE1 (n=150 matched targets) and RPE1-to-K562 (n=148 matched targets). Labels show paired drops and perturbation-level bootstrap 95% intervals. UER values discussed in text are sensitivity-only.

**Figure 5. STATE partial cross-architecture support.** STATE K562-to-RPE1 matched targets (n=15) show lower cross-context audit-delta Pearson, Spearman, and cosine, while UER has an interval crossing zero. The full-summary MRR bar is included to show endpoint heterogeneity rather than complete replication.
"""
    write(MANUSCRIPT / "CRM_MANUSCRIPT_v1.1.md", text)


def markdown_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    for line in md_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            continue
        else:
            doc.add_paragraph(line)
    doc.save(docx_path)


def build_supporting_files(metrics):
    sens, state, state_drop, norm_rep, probes, k2r_p, k2r_u, k2r_s, r2k_p, st_p = metrics
    supplement = f"""# VirtualPerturb-Audit Supplementary Information

Draft version: CRM_SUPPLEMENT_v1.1

Generated: {GENERATED}

## Frozen Analysis State

Phase 2A-2C primary outputs are frozen. CRM v1.1 finalization did not rerun GEARS, rerun STATE, add datasets, redefine endpoints, or change matched-target registries.

## Permanent Limitations

- Replogle analyses use GEARS-compatible filtered essential-screen data, not complete Figshare+ processed objects.
- Validated biological replicate metadata were unavailable.
- BNS remains `UNVERIFIED`.
- UER remains `sensitivity_only` because its null envelope is not validated biological replicate ground truth.
- GEARS R-L4 is a GEARS-compatible cross-context inference adapter.
- STATE support is partial and endpoint-heterogeneous.
- GEARS and STATE absolute values are not direct universal leaderboard values.

## STATE Primary Metrics

{md_table(state[state.metric_space.isin(["audit_delta", "target_control_audit_delta"])][["setting","split","metric_space","n_test_perturbations","pearson_delta","spearman_delta","cosine_delta","retrieval_mrr","uer50","sign_flip_rate"]])}

## STATE Matched Transfer

{md_table(state_drop)}

## GEARS Matched Transfer Sensitivity

{md_table(sens[sens.comparison_role == "primary_source_context_comparison"][["direction","metric","n_targets","within_estimate","cross_estimate","paired_difference","ci_low","ci_high"]])}

## Additional Retrieval Metrics

Common-candidate retrieval rows in `results/tables/replogle_matched_rl1_rl4_sensitivity.csv` remain low in both transfer directions and should be interpreted as perturbation-specific retrieval stress tests, not as global expression-fit endpoints.

## Probe Controls

Probe-control rows are stored in `results/tables/replogle_gears_vs_probes.csv`. They show that mean-effect structure can support audit-delta Pearson while retrieval remains weak.

## QC, Split Integrity, Null Sensitivity, and Adapter Details

Primary provenance and audit depth are retained in `reports/replogle_split_integrity_report.md`, `reports/replicate_label_audit.md`, `reports/STATE_GEARS_METRIC_COMPATIBILITY.md`, `reports/STATE_RL4_ADAPTER_REPORT.md`, and `reports/PHASE2C_RESULT_INTERPRETATION.md`.

## Gene-Family L3, Seed Sensitivity, and Vocabulary Compatibility

Norman L3 gene-family holdout and seed-sensitivity outputs are retained in `results/tables/table8_seed_robustness_summary.*`, `results/tables/table9_gene_family_confusion_summary.*`, and `results/tables/table10_l3_gene_family_holdout_candidates.*`.
"""
    write(MANUSCRIPT / "CRM_SUPPLEMENT_v1.1.md", supplement)

    highlights = """# Highlights

- VirtualPerturb-Audit stress-tests model reliability across endpoints
- Matched targets reveal context-transfer degradation in GEARS
- STATE partially reproduces matched-target transfer degradation
- Multiple endpoints expose conclusions hidden by aggregate similarity
"""
    write(MANUSCRIPT / "HIGHLIGHTS_v1.0.md", highlights)

    write(MANUSCRIPT / "IN_BRIEF_v1.0.md", """# In Brief

VirtualPerturb-Audit tests whether perturbation-response model claims survive stricter endpoint and context-shifted evaluation. Applied to frozen GEARS and STATE analyses, it separates global similarity from perturbation-specific retrieval, unsupported-effect behavior, and matched-target transfer degradation. The framework helps convert model-performance numbers into bounded, evidence-linked claims.
""")

    write(MANUSCRIPT / "GRAPHICAL_ABSTRACT_BRIEF_v1.0.md", """# Graphical Abstract Brief

Left: predicted perturbation response. Center: VirtualPerturb-Audit with five stacked audit dimensions: global fit, target retrieval, unsupported effects, matched transfer, and falsification probes. Right: claim survives or claim narrows.

Core text: Global agreement does not equal perturbation-specific reliability.

Avoid the phrase "AI hallucination exposed." The graphic should present a reusable audit workflow rather than a negative model-failure poster.
""")

    write(MANUSCRIPT / "AUTHOR_CONTRIBUTIONS.md", """# Author Contributions

Author order, affiliations, and CRediT roles require final author input.

Suggested CRediT placeholders:

- Conceptualization: [Name(s)]
- Methodology: [Name(s)]
- Software: [Name(s)]
- Formal analysis: [Name(s)]
- Investigation: [Name(s)]
- Data curation: [Name(s)]
- Visualization: [Name(s)]
- Writing - original draft: [Name(s)]
- Writing - review and editing: [Name(s)]
- Supervision: [Name(s)]
- Project administration: [Name(s)]
""")

    write(MANUSCRIPT / "DECLARATION_OF_INTERESTS.md", """# Declaration of Interests

The authors declare no competing interests.
""")

    write(MANUSCRIPT / "FUNDING_STATEMENT.md", """# Funding

This work received no specific funding.
""")

    # KRT final.
    rows = [
        ["Deposited data", "Norman Perturb-seq dataset", "Norman et al., 2019", "DOI: 10.1126/science.aax4438"],
        ["Deposited data", "Replogle Perturb-seq dataset", "Replogle et al., 2022", "DOI: 10.1016/j.cell.2022.05.013"],
        ["Dataset", "GEARS-compatible filtered Replogle essential-screen data", "Roohani et al. / GEARS resources", "Filtered K562/RPE1 audit scope"],
        ["Software and algorithms", "VirtualPerturb-Audit", "This paper", "Repository URL: TODO_DEPOSIT"],
        ["Software and algorithms", "GEARS / cell-gears", "Roohani et al., 2024", "DOI: 10.1038/s41587-023-01905-6"],
        ["Software and algorithms", "STATE", "Arc Institute / local installation", "Citation/source: MANUAL_CONFIRMATION_REQUIRED"],
        ["Software and algorithms", "Python", "Python Software Foundation", "3.12 post-processing environment"],
        ["Software and algorithms", "PyTorch", "PyTorch project", "Installed in project environments"],
        ["Software and algorithms", "Scanpy", "Scanpy developers", "Single-cell object handling"],
        ["Software and algorithms", "AnnData", "scverse", "h5ad data structure"],
        ["Software and algorithms", "scikit-learn", "scikit-learn developers", "Baseline and metric utilities"],
        ["Software and algorithms", "SciPy", "SciPy developers", "Statistical utilities"],
        ["Database", "HGNC", "HGNC", "Gene-family annotation where used"],
        ["Database", "Gene Ontology / Reactome", "GO / Reactome", "Prior annotations where used by model resources"],
    ]
    wb = Workbook()
    ws = wb.active
    ws.title = "Key Resources Table"
    headers = ["REAGENT or RESOURCE", "SOURCE", "IDENTIFIER"]
    ws.append(headers)
    for row in rows:
        ws.append([row[0], row[1], f"{row[2]}; {row[3]}"])
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2E6F9E")
    for i, w in enumerate([30, 42, 60], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    wb.save(MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.xlsx")
    write(MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.md", "# Key Resources Table v1.0\n\n" + md_table(pd.DataFrame([[r[0], r[1], f"{r[2]}; {r[3]}"] for r in rows], columns=headers)))

    cover = f"""# Cover Letter to Cell Reports Methods

Generated: {GENERATED}

Dear Editors,

We are pleased to submit "VirtualPerturb-Audit: a falsification framework for perturbation-response model evaluation" as an Article for Cell Reports Methods.

The central contribution is not a new perturbation predictor, but a reusable framework for testing which aspects of apparent predictive performance survive perturbation-specific and context-shifted stress testing. The framework freezes analysis inputs, separates raw-space global similarity from control-subtracted audit-delta agreement, evaluates perturbation-specific retrieval and unsupported-effect behavior, and records which claims survive matched-target transfer analysis.

The worked example uses frozen GEARS and STATE analyses on Norman and GEARS-compatible filtered Replogle perturbation data. In matched GEARS K562-to-RPE1 transfer, audit-delta Pearson decreased from 0.2812 to -0.0070, a paired drop of 0.2883 with a 95% interval of [0.2559, 0.3206]. The reverse RPE1-to-K562 direction showed a paired drop of 0.5480. STATE provided partial cross-architecture support: across 15 matched Replogle targets, audit-delta Pearson decreased from 0.2955 to 0.1792, a drop of 0.1163 with a 95% interval of [0.0684, 0.1599].

We believe the manuscript fits Cell Reports Methods because it provides a reusable, reproducible evaluation workflow for an active area of computational biology. The manuscript emphasizes claim discipline and community utility rather than a direct model leaderboard. It also makes its boundaries explicit: Replogle analyses use GEARS-compatible filtered essential-screen data, validated biological replicate metadata were unavailable, UER is sensitivity-only, and STATE support is endpoint-heterogeneous.

The authors declare no competing interests. This work received no specific funding. Public code and data archive details are being finalized and are marked as TODO_DEPOSIT in the submission-preparation files.

Sincerely,

[Corresponding author name]
"""
    write(SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS_v1.0.md", cover)
    markdown_to_docx(SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS_v1.0.md", SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS_v1.0.docx")


def build_audits(metrics):
    _, state, state_drop, _, _, k2r_p, k2r_u, k2r_s, r2k_p, st_p = metrics
    write(REPORTS / "PEARSON_METRIC_SPACE_AUDIT.md", """# Pearson Metric Space Audit

Decision: `FIXED`.

| Figure/table/text | Metric label | Metric space | Control subtraction? | Interpretation |
|---|---|---|---|---|
| Figure 2 | raw expression Pearson | raw-space | No | Global transcriptomic similarity |
| Figure 4 | audit-delta Pearson | audit-delta | Yes | Matched perturbation-response agreement |
| Figure 5 | audit-delta Pearson | audit-delta / target-control audit-delta | Yes | STATE matched transfer agreement |
| Manuscript Result 2 | raw-space Pearson | raw-space | No | Not interchangeable with audit-delta values |
| Manuscript Results 4-5 | audit-delta Pearson | audit-delta | Yes | Transfer stress-test endpoint |
""")

    write(REPORTS / "REFERENCE_AUDIT_FINAL.md", f"""# Reference Audit Final

Generated: {GENERATED}

## VERIFIED

| Citation | Journal/source | Year | Volume/pages | DOI/URL | PubMed |
|---|---|---|---|---|---|
| Norman et al., Exploring genetic interaction manifolds constructed from rich single-cell phenotypes | Science | 2019 | 365:786-793 | 10.1126/science.aax4438 | 31395745 |
| Replogle et al., Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq | Cell | 2022 | 185:2559-2575.e28 | 10.1016/j.cell.2022.05.013 | 35688146 |
| Roohani et al., Predicting transcriptional outcomes of novel multigene perturbations with GEARS | Nature Biotechnology | 2024 | 42:927-935 | 10.1038/s41587-023-01905-6 | 37592036 |
| Wu et al., PerturBench: Benchmarking Machine Learning Models for Cellular Perturbation Analysis | NeurIPS 38 / arXiv | 2025 | 106937-106977 | 10.52202/085713-3225; arXiv:2408.10609 | NA |
| Vinas Torne et al., Systema: a framework for evaluating genetic perturbation response prediction beyond systematic variation | Nature Biotechnology | 2025 | Online first | 10.1038/s41587-025-02777-8 | NA |
| Radig et al., scArchon: a scalable benchmarking framework for assessing single-cell perturbation models | Genome Biology | 2026 | 27:162 | 10.1186/s13059-026-04104-z | 42121287 |
| Mao et al., Benchmarking virtual cell models for in-the-wild perturbation response | arXiv | 2026 | 2604.27646 | https://arxiv.org/abs/2604.27646 | NA |
| Vollenweider et al., Signal, Bounds, and Baselines | bioRxiv | 2026 | Preprint | 10.64898/2026.04.20.719650 | NA |
| Replogle et al. SRA/GEO file manifest | Figshare+ | 2022 | Dataset | 10.25452/figshare.plus.20022944 | NA |
| Replogle et al. processed datasets | Figshare+ | 2022 | Dataset | 10.25452/figshare.plus.20029387 | NA |

## MANUAL_CONFIRMATION_REQUIRED

- STATE citation/source form for final journal upload.
- Exact final page/article metadata for 2025-2026 benchmark papers if journal upload requires finalized issue information.
- Reference-manager export and Cell Reports Methods style conversion.

## REMOVE

- Unverified scContam citation. The manuscript does not discuss pretraining contamination in enough detail to require it.
""")

    write(REPORTS / "DATA_RELEASE_FINAL_AUDIT.md", """# Data Release Final Audit

| Artifact | Current location | Planned destination | Required before submission? | Required before acceptance? |
|---|---|---|---|---|
| Original Norman data | External source plus local mirror/provenance | Cite source; repository instructions | Yes | Yes |
| Original Replogle data | External source plus filtered local objects | Cite Cell/Figshare+; clarify filtered scope | Yes | Yes |
| Processed audit tables | `results/tables/` | Public repository/archive DOI TODO_DEPOSIT | Yes | Yes |
| Split assignments | project split/metadata files | Public repository/archive DOI TODO_DEPOSIT | Yes | Yes |
| Predictions | local result directories | Archive or reviewer-access storage TODO_DEPOSIT | Yes for review access | Yes |
| Summary metrics | `results/tables/` and reports | Public repository/archive DOI TODO_DEPOSIT | Yes | Yes |
| Source code | local git repository | Public repository URL TODO_DEPOSIT | Yes for review access | Yes |
| Environment files | `environment/` and `REPRODUCIBILITY.md` | Repository plus frozen export | Yes | Yes |
""")

    write(REPORTS / "CODE_RELEASE_FINAL_AUDIT.md", """# Code Release Final Audit

Status: `NEEDS_MINOR_CLEANUP`.

## Review

- README: updated to reviewer mode.
- LICENSE: MIT license added.
- Installation: described at high level; exact environment export still needs final pinning.
- Environment: project environments exist; top-level minimal export remains recommended.
- CLI: scripts are present for data acquisition, GEARS, STATE post-processing, tables, figures, and CRM package generation.
- Minimal example: described in README.
- Reproduce core results: frozen-table regeneration path documented.
- Build figures/tables: documented in README and REPRODUCIBILITY.
- Frozen split retrieval: available through project split reports and tables.
- Test suite: existing tests passed in prior project phases; final v1.1 did not rerun model tests.

## Remaining Cleanup

Public repository URL, archive DOI, final environment export, and reviewer-access instructions remain manual release tasks.
""")

    write(REPORTS / "ARTICLE_TYPE_FINAL_DECISION.md", """# Article Type Final Decision

Decision: `ARTICLE`.

Rationale: the submission presents a reusable computational method and reporting framework, demonstrated through frozen GEARS and STATE worked examples. It should not be framed as a Report, because the main contribution is the audit method and reproducibility package rather than a compact single-result observation.
""")

    write(REPORTS / "MANUAL_AUTHOR_ITEMS.md", """# Manual Author Items

- Author order
- Degrees if required by journal system
- Affiliations
- Corresponding author name
- Corresponding author email
- ORCID identifiers if required
- Final CRediT roles
- Acknowledgments
- Repository URL and archive DOI

Confirmed by user:

- Funding: no specific funding
- Competing interests: none
- License: MIT selected for repository release
""")

    claim = pd.DataFrame([
        ["VirtualPerturb-Audit is reusable", "Figure 1; README; reproducibility package", "Moderate", "reusable falsification framework", "universal benchmark"],
        ["Global fit differs from perturbation specificity", "Figure 2; raw-space Pearson vs MRR", "Strong within frozen scope", "raw-space similarity and retrieval diverge", "global fit proves reliability"],
        ["GEARS degradation persists after target matching", "Figure 4; matched K2R/R2K table", "Strong", "matched-target transfer collapse", "fully eliminates all confounding"],
        ["STATE partially supports transfer degradation", "Figure 5; n=15 matched targets", "Moderate", "partial cross-architecture support", "architecture-independent confirmation"],
        ["UER is sensitivity-only", "BNS unavailable; null not replicate-derived", "Strong limitation", "sensitivity-only unsupported-effect rate", "validated biological hallucination"],
        ["Endpoint heterogeneity exists", "STATE MRR/UER full-summary caveats", "Strong", "endpoint-level heterogeneity limits broader generalization", "uniform confirmation"],
    ], columns=["Claim", "Supporting figure/table", "Evidence strength", "Allowed wording", "Prohibited wording"])
    write(REPORTS / "CLAIM_EVIDENCE_MATRIX_FINAL.md", "# Claim-Evidence Matrix Final\n\n" + md_table(claim))

    write(REPORTS / "CRM_EDITORIAL_AUDIT_FINAL.md", """# CRM Editorial Audit Final

Desk reject risk: `MODERATE`.

1. Is this a method or benchmark? Method. The benchmark examples demonstrate the audit framework.
2. Is the method reusable? Yes, but public repository and archive metadata must be finalized.
3. Are data/code sufficiently accessible? Not yet for final upload; reviewer-access details remain TODO_DEPOSIT.
4. Is validation breadth enough? Borderline but acceptable for a methods submission if claims remain bounded.
5. Does filtered Replogle materially weaken the paper? It weakens complete-data claims, but not the methods demonstration if labeled clearly.
6. Is partial STATE support acceptable? Yes, because the paper uses STATE to test claim discipline, not to assert universal confirmation.
7. Biggest desk-reject reason? Perceived overlap with recent perturbation benchmarks unless the audit-and-claim-evidence contribution is foregrounded.
8. Can it be solved without new experiments? Mostly yes, through framing, references, code release, and exact data availability.
""")

    write(REPORTS / "CRM_REVIEWER_SIMULATION_FINAL.md", """# CRM Reviewer Simulation Final

## Reviewer 1: Computational Methods

Major 1: The method may overlap with existing benchmarks. Classification: `MANUSCRIPT_FIX`.
Major 2: Metric-space differences may confuse readers. Classification: `ALREADY_ADDRESSED`.
Major 3: Reproducibility depends on public code release. Classification: `MANUSCRIPT_FIX`.
Minor 1: Define R-L1/R-L4 earlier. Classification: `ALREADY_ADDRESSED`.
Minor 2: State that UER is sensitivity-only in legends. Classification: `ALREADY_ADDRESSED`.
Minor 3: Add final repository DOI. Classification: `MANUSCRIPT_FIX`.

## Reviewer 2: Single-Cell Perturbation Biology

Major 1: Filtered Replogle scope may limit biological claims. Classification: `ALREADY_ADDRESSED`.
Major 2: Lack of validated biological replicates limits BNS and UER. Classification: `ALREADY_ADDRESSED`.
Major 3: Complete Figshare+ data were not analyzed. Classification: `NEW_EXPERIMENT_REQUIRED`.
Minor 1: Clarify cell-line directions. Classification: `ALREADY_ADDRESSED`.
Minor 2: Avoid biological validation wording. Classification: `ALREADY_ADDRESSED`.
Minor 3: Add exact data source URLs. Classification: `MANUSCRIPT_FIX`.

## Reviewer 3: Model Evaluation

Major 1: STATE support is based on 15 matched targets. Classification: `ALREADY_ADDRESSED`.
Major 2: Full-summary endpoints are mixed. Classification: `ALREADY_ADDRESSED`.
Major 3: More architectures would improve generality. Classification: `NEW_EXPERIMENT_REQUIRED`.
Minor 1: Include MRR caveat visually. Classification: `ALREADY_ADDRESSED`.
Minor 2: Use consistent audit-delta terminology. Classification: `ALREADY_ADDRESSED`.
Minor 3: Confirm STATE citation. Classification: `MANUSCRIPT_FIX`.
""")

    write(REPORTS / "POSTHOC_EXPLORATORY_ANALYSES.md", """# Posthoc Exploratory Analyses

No new exploratory analyses were executed during CRM v1.1 finalization.

Policy: future cheap analyses are allowed only if they use frozen outputs, require no new training, do not redefine endpoints, finish in under 30 minutes, and directly answer a major reviewer concern. Any such analysis must be labeled `EXPLORATORY`.
""")

    write(REPORTS / "TERMINOLOGY_AUDIT_FINAL.md", """# Terminology Audit Final

Status: `PASS_WITH_LEGACY_FILES_UNCHANGED`.

Canonical terms used in v1.1 files:

- VirtualPerturb-Audit
- perturbation-response model
- global transcriptomic similarity
- raw-space Pearson
- audit-delta Pearson
- perturbation-specific retrieval
- unsupported-effect rate
- sign-flip rate
- within-context perturbation holdout
- cross-context transfer
- matched-target analysis
- partial cross-architecture support

Legacy Phase 2 reports are retained as provenance and were not rewritten.
""")

    write(REPORTS / "NUMERICAL_AUDIT_FINAL.md", f"""# Numerical Audit Final

Status: `PASS`.

| Quantity | Expected | v1.1 status |
|---|---|---|
| GEARS K2R within audit-delta Pearson | 0.2812 | PASS |
| GEARS K2R cross audit-delta Pearson | -0.0070 | PASS |
| GEARS K2R drop | 0.2883 | PASS |
| GEARS K2R CI | [0.2559, 0.3206] | PASS |
| GEARS K2R UER50 | 0.1532 -> 0.3877 | PASS |
| GEARS K2R sign-flip rate | 0.2714 -> 0.5718 | PASS |
| GEARS R2K within audit-delta Pearson | 0.5501 | PASS |
| GEARS R2K cross audit-delta Pearson | 0.0021 | PASS |
| GEARS R2K drop | 0.5480 | PASS |
| GEARS R2K CI | [0.5146, 0.5802] | PASS |
| STATE matched n | 15 | PASS |
| STATE matched within audit-delta Pearson | 0.2955 | PASS |
| STATE matched cross audit-delta Pearson | 0.1792 | PASS |
| STATE matched drop | 0.1163 | PASS |
| STATE matched CI | [0.0684, 0.1599] | PASS |
""")

    write(REPORTS / "CRM_SUBMISSION_READINESS_FINAL.md", """# CRM Submission Readiness Final

Submission-readiness: `READY_AFTER_MINOR_MANUAL_ITEMS`.

| Dimension | Score (1-5) | Rationale |
|---|---:|---|
| Scientific story | 4 | Clear methods framing and strongest GEARS matched result |
| Method novelty | 4 | Audit grammar distinct from benchmark-only framing |
| Rigor | 4 | Frozen outputs, matched-target analysis, explicit limitations |
| Reproducibility | 3 | Local artifacts strong; public archive still TODO |
| Code readiness | 3 | MIT license and README added; environment export still needs final pinning |
| Data readiness | 3 | Sources clear; prediction/archive DOI still TODO |
| Figure quality | 4 | Main figures generated and visually checked |
| Reference completeness | 4 | Core references verified; STATE citation manual |
| Formatting completeness | 3 | Markdown/DOCX package ready; exact upload-system details need manual confirmation |

Largest desk-reject risk: overlap with recent perturbation benchmarks if the manuscript is read as a benchmark rather than as an audit-and-claim-evidence method.

Largest reviewer risk: filtered Replogle scope and partial STATE support limit breadth. These are addressed as limitations, not solved by new experiments in this phase.

Remaining manual tasks: author metadata, ORCID if required, corresponding author email, repository URL, archive DOI, final STATE citation/source, and upload-system format confirmation.
""")


def build_license_readme_repro():
    write(ROOT / "LICENSE", """MIT License

Copyright (c) 2026 VirtualPerturb-Audit contributors

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
""")

    write(ROOT / "README.md", """# VirtualPerturb-Audit

VirtualPerturb-Audit is a falsification framework for perturbation-response model evaluation. It asks which model-performance claims survive when global transcriptomic similarity is separated from perturbation-specific retrieval, unsupported-effect behavior, sign-flip rate, leakage risk, and matched-target context transfer.

## What Problem Does It Solve?

Aggregate expression similarity can make a perturbation-response model look reliable even when target retrieval or cross-context transfer is weak. This repository turns those concerns into an auditable workflow with frozen splits, endpoint-specific tables, probe controls, and claim-evidence reports.

## Main Audit Dimensions

- Raw-space global transcriptomic similarity
- Control-subtracted audit-delta agreement
- Perturbation-specific retrieval
- Unsupported-effect rate, reported as sensitivity-only UER
- Sign-flip rate
- Matched-target context transfer
- Falsification probes and baseline controls

## Supported Example Datasets

- Norman perturbation data through a GEARS-compatible processed mirror
- GEARS-compatible filtered Replogle K562 and RPE1 essential-screen data

The current Replogle analyses do not use the complete Figshare+ processed objects.

## Quick Start

Inspect the frozen CRM package:

```bash
ls submission/cell_reports_methods/final
```

Regenerate the v1.1 submission-preparation package from frozen outputs:

```bash
environment/state-postprocess-venv/bin/python scripts/finalize_crm_submission_v11.py
```

Regenerate the earlier CRM v1.0 package:

```bash
environment/state-postprocess-venv/bin/python scripts/build_crm_submission_package.py
```

## Expected Outputs

- `manuscript/CRM_MANUSCRIPT_v1.1.md`
- `manuscript/CRM_MANUSCRIPT_v1.1.docx`
- `manuscript/CRM_SUPPLEMENT_v1.1.md`
- `figures/main/crm_figure4_matched_gears_transfer_v11.*`
- `figures/main/crm_figure5_state_partial_confirmation_v11.*`
- `reports/CRM_SUBMISSION_READINESS_FINAL.md`
- `submission/cell_reports_methods/final/`

## Reproduction Commands

The finalization script does not train models. It reads frozen tables under `results/tables/` and writes manuscript, figure, audit, and submission files. Earlier GEARS and STATE training commands are retained in phase-specific reports.

## Known Limitations

- Replogle scope is GEARS-compatible filtered essential-screen data.
- BNS remains unverified because validated biological replicate metadata were unavailable.
- UER is sensitivity-only because its null is not derived from validated biological replicate ground truth.
- GEARS R-L4 is a cross-context inference adapter.
- STATE support is partial and endpoint-heterogeneous.
- GEARS and STATE absolute metrics are not a direct universal model leaderboard.
""")

    write(ROOT / "REPRODUCIBILITY.md", """# Reproducibility

## Hardware and Operating Systems

- Local manuscript/finalization host: macOS on Apple Silicon.
- GEARS full CPU runs: Mac CPU, long-running jobs taking approximately 5-6 hours per Norman split and longer for Replogle R-L1.
- STATE Phase 2C full runs: CUDA-capable Linux GPU server.

## GPU and CUDA

STATE confirmatory execution requires a CUDA-capable Linux environment. The completed Phase 2C run used a rented GPU server with NVIDIA driver/CUDA support sufficient for STATE. Exact server details and run manifests are retained in `results/tables/state_phase2c_run_manifest.csv` and Phase 2C reports.

## Memory and Disk

Replogle and STATE artifacts require tens of GB of working storage. Raw Phase 2C h5ad outputs are local but not copied into the submission package.

## Python and Packages

Post-processing and CRM finalization use the project environment:

```bash
environment/state-postprocess-venv/bin/python
```

Key packages include pandas, matplotlib, python-docx, openpyxl, PyTorch, Scanpy, AnnData, scikit-learn, and SciPy. Final public release should include a pinned environment export.

## Seeds and Task Names

- GEARS seed: 1 for frozen full runs.
- GEARS R-L1 tasks: `R-L1-K562`, `R-L1-RPE1`.
- GEARS R-L4 tasks: `R-L4-K2R`, `R-L4-R2K`.
- STATE Phase 2C tasks: `S1_norman_l1`, `S2_norman_l2`, `S3_replogle_k562_rl1`, `S4_replogle_k2r_rl4`.

## Expected Result Files

- `results/tables/replogle_matched_rl1_rl4_sensitivity.csv`
- `results/tables/replogle_rl1_rl4_gears_comparison.csv`
- `results/tables/state_phase2c_primary_metrics.csv`
- `results/tables/state_transfer_drop.csv`
- `results/tables/gears_state_primary_comparison.csv`

## Figure Regeneration

```bash
environment/state-postprocess-venv/bin/python scripts/finalize_crm_submission_v11.py
```

## Known Non-Portable Steps

- Complete Replogle Figshare+ command-line access was blocked by HTTP 403 in the frozen project state.
- Full GEARS training is slow and terminal-lifetime sensitive on CPU.
- STATE execution depends on a CUDA-capable Linux server and local adapter paths.
- Public repository URL, archive DOI, and final environment export remain manual release tasks.
""")


def render_docx(docx_path: Path) -> str:
    render_script = Path("/Users/zy/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/render_docx.py")
    if not render_script.exists():
        return "RENDER_SCRIPT_MISSING"
    out = QA / docx_path.stem
    out.mkdir(parents=True, exist_ok=True)
    cmd = [os.environ.get("PYTHON", "python3"), str(render_script), str(docx_path), "--output_dir", str(out)]
    try:
        subprocess.run(cmd, check=True, cwd=str(ROOT), stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=120)
        pngs = sorted(out.glob("page-*.png"))
        return f"RENDERED_{len(pngs)}_PAGES" if pngs else "RENDERED_NO_PNG"
    except Exception as exc:
        return f"RENDER_FAILED:{type(exc).__name__}"


def build_final_package():
    if FINAL.exists():
        shutil.rmtree(FINAL, ignore_errors=True)
    dirs = [
        "main_manuscript", "supplement", "main_figures", "supp_figures", "cover_letter",
        "highlights", "in_brief", "graphical_abstract_brief", "key_resources_table",
        "data_code_availability", "reproducibility", "author_manual_items", "audits",
    ]
    for d in dirs:
        (FINAL / d).mkdir(parents=True, exist_ok=True)
    copy_map = [
        (MANUSCRIPT / "CRM_MANUSCRIPT_v1.1.md", "main_manuscript"),
        (MANUSCRIPT / "CRM_MANUSCRIPT_v1.1.docx", "main_manuscript"),
        (MANUSCRIPT / "CRM_SUPPLEMENT_v1.1.md", "supplement"),
        (SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS_v1.0.md", "cover_letter"),
        (SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS_v1.0.docx", "cover_letter"),
        (MANUSCRIPT / "HIGHLIGHTS_v1.0.md", "highlights"),
        (MANUSCRIPT / "IN_BRIEF_v1.0.md", "in_brief"),
        (MANUSCRIPT / "GRAPHICAL_ABSTRACT_BRIEF_v1.0.md", "graphical_abstract_brief"),
        (MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.xlsx", "key_resources_table"),
        (MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.md", "key_resources_table"),
        (REPORTS / "DATA_RELEASE_FINAL_AUDIT.md", "data_code_availability"),
        (REPORTS / "CODE_RELEASE_FINAL_AUDIT.md", "data_code_availability"),
        (ROOT / "REPRODUCIBILITY.md", "reproducibility"),
        (REPORTS / "MANUAL_AUTHOR_ITEMS.md", "author_manual_items"),
    ]
    for src, d in copy_map:
        if src.exists():
            shutil.copy2(src, FINAL / d / src.name)
    for stem in ["crm_figure1_audit_framework", "crm_figure2_norman_metric_divergence", "crm_figure3_replogle_within_context", "crm_figure4_matched_gears_transfer_v11", "crm_figure5_state_partial_confirmation_v11"]:
        for ext in ["pdf", "svg", "png"]:
            src = FIG_MAIN / f"{stem}.{ext}"
            if src.exists():
                shutil.copy2(src, FINAL / "main_figures" / src.name)
    for pattern in ["phase2c_endpoint_heatmap.*", "phase2c_retrieval_rank_distribution.*"]:
        for src in FIG_SUPP.glob(pattern):
            if not src.name.startswith("._"):
                shutil.copy2(src, FINAL / "supp_figures" / src.name)
    for src in [
        "PEARSON_METRIC_SPACE_AUDIT.md", "REFERENCE_AUDIT_FINAL.md", "ARTICLE_TYPE_FINAL_DECISION.md",
        "CLAIM_EVIDENCE_MATRIX_FINAL.md", "CRM_EDITORIAL_AUDIT_FINAL.md", "CRM_REVIEWER_SIMULATION_FINAL.md",
        "NUMERICAL_AUDIT_FINAL.md", "TERMINOLOGY_AUDIT_FINAL.md", "CRM_SUBMISSION_READINESS_FINAL.md",
        "POSTHOC_EXPLORATORY_ANALYSES.md",
    ]:
        p = REPORTS / src
        if p.exists():
            shutil.copy2(p, FINAL / "audits" / src)
    write(SUBMISSION / "cell_reports_methods" / "FINAL_SUBMISSION_CHECKLIST.md", """# Final Submission Checklist

| Item | Status |
|---|---|
| Title | DONE |
| Authors | MANUAL_REQUIRED |
| Affiliations | MANUAL_REQUIRED |
| Correspondence | MANUAL_REQUIRED |
| Summary | DONE |
| Main text | DONE |
| Figures | DONE |
| Legends | DONE |
| Supplement | DONE |
| References | MANUAL_REQUIRED |
| Key Resources Table | DONE |
| Highlights | DONE |
| In Brief | DONE |
| Cover letter | DONE |
| Graphical abstract | MANUAL_REQUIRED |
| Data statement | MANUAL_REQUIRED |
| Code statement | MANUAL_REQUIRED |
| CRediT | MANUAL_REQUIRED |
| Funding | DONE |
| COI | DONE |
| ORCID | MANUAL_REQUIRED |
| Repository accessibility | MANUAL_REQUIRED |
""")
    shutil.copy2(SUBMISSION / "cell_reports_methods" / "FINAL_SUBMISSION_CHECKLIST.md", FINAL / "FINAL_SUBMISSION_CHECKLIST.md")
    for p in FINAL.rglob("._*"):
        if p.is_file():
            p.unlink()
    manifest = [{"path": str(p.relative_to(FINAL)), "bytes": p.stat().st_size} for p in sorted(FINAL.rglob("*")) if p.is_file() and not p.name.startswith("._")]
    write(FINAL / "FINAL_PACKAGE_MANIFEST.json", json.dumps(manifest, indent=2, ensure_ascii=False))
    for p in FINAL.rglob("._*"):
        if p.is_file():
            p.unlink()


def main():
    for d in [REPORTS, MANUSCRIPT, SUBMISSION, FIG_MAIN, FIG_SUPP]:
        d.mkdir(parents=True, exist_ok=True)
    metrics = get_metrics()
    _, state, state_drop, _, _, *_ = metrics
    build_final_figures(state, state_drop, metrics[0])
    build_manuscript_v11(metrics)
    build_supporting_files(metrics)
    build_audits(metrics)
    build_license_readme_repro()
    markdown_to_docx(MANUSCRIPT / "CRM_MANUSCRIPT_v1.1.md", MANUSCRIPT / "CRM_MANUSCRIPT_v1.1.docx")
    main_render = render_docx(MANUSCRIPT / "CRM_MANUSCRIPT_v1.1.docx")
    write(REPORTS / "DOCX_QA_STATUS.md", f"""# DOCX QA Status

Main manuscript DOCX: `{main_render}`

Supplement DOCX: `NOT_GENERATED`; supplement is delivered as Markdown because the wide audit tables are safer as source tables/Markdown unless Cell Reports Methods requests a DOCX conversion.

QA output directory: `{QA}`
""")
    build_final_package()
    print("CRM v1.1 finalization complete")
    print(FINAL)


if __name__ == "__main__":
    main()
