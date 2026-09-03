#!/usr/bin/env python3
"""Update Cell Reports Methods cover letter Markdown and DOCX files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]

MARKDOWN_TEXT = """Dear Editors,

We are pleased to submit our Article manuscript, **"VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models,"** for consideration in *Cell Reports Methods*.

Perturbation-response models are increasingly evaluated using aggregate transcriptomic agreement, yet strong global similarity does not necessarily establish that a model has recovered perturbation-specific effects or can transfer those effects across cellular contexts. VirtualPerturb-Audit addresses this gap by shifting perturbation-model evaluation from **model ranking to claim falsification**. The central contribution is not a new perturbation predictor or another benchmark leaderboard, but a reusable audit framework that asks which performance claims remain supported after target-information removal, perturbation-specific retrieval testing, matched-target context shift, and endpoint-specific stress testing. The framework additionally freezes analysis provenance and maps different endpoint families to explicit, bounded scientific claims.

We demonstrate the framework using frozen GEARS analyses of Norman and GEARS-compatible filtered Replogle K562/RPE1 data, together with an architecturally distinct STATE analysis. The GEARS results show that high global expression agreement can coexist with weak perturbation retrieval, while matched-target analyses reveal substantial cross-context degradation in both K562-to-RPE1 and RPE1-to-K562 directions. The STATE analysis provides more limited but directionally consistent cross-architecture support across agreement endpoints, while retrieval and error-burden measures remain heterogeneous. These results illustrate why discordant endpoints should constrain the strength of a model claim rather than be collapsed into a single performance score.

We believe this work is particularly suited to *Cell Reports Methods* because its primary contribution is a reproducible and model-agnostic evaluation methodology that can be applied across perturbation-prediction systems. Rather than proposing a universal ranking of GEARS, STATE, or other models, VirtualPerturb-Audit provides a practical framework for determining whether evidence supports global expression reconstruction, perturbation-identity recovery, context transfer, directional fidelity, or only a narrower response-structure interpretation. We deliberately retain the limitations of the frozen evidence, including the use of GEARS-compatible filtered Replogle essential-screen data and the sensitivity-only interpretation of unsupported-effect metrics.

This manuscript is original and is not under consideration elsewhere. Information regarding data and code availability, funding, conflicts of interest, and author contributions is provided in the submission materials.

Thank you for considering our manuscript. We believe VirtualPerturb-Audit will be useful to researchers developing, benchmarking, and interpreting perturbation-response models, and we would be grateful for the opportunity to have the work considered for publication in *Cell Reports Methods*.

Sincerely,

**Yu Zhang**  
Corresponding Author  
Department of Ophthalmology  
The Second Affiliated Hospital of Wenzhou Medical University  
Wenzhou, Zhejiang Province, China  
Email: zhangyu1@wzhealth.com
"""

DOCX_PARAGRAPHS = [
    ("Dear Editors,", False, False),
    (
        'We are pleased to submit our Article manuscript, "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models," for consideration in Cell Reports Methods.',
        False,
        False,
    ),
    (
        "Perturbation-response models are increasingly evaluated using aggregate transcriptomic agreement, yet strong global similarity does not necessarily establish that a model has recovered perturbation-specific effects or can transfer those effects across cellular contexts. VirtualPerturb-Audit addresses this gap by shifting perturbation-model evaluation from model ranking to claim falsification. The central contribution is not a new perturbation predictor or another benchmark leaderboard, but a reusable audit framework that asks which performance claims remain supported after target-information removal, perturbation-specific retrieval testing, matched-target context shift, and endpoint-specific stress testing. The framework additionally freezes analysis provenance and maps different endpoint families to explicit, bounded scientific claims.",
        False,
        False,
    ),
    (
        "We demonstrate the framework using frozen GEARS analyses of Norman and GEARS-compatible filtered Replogle K562/RPE1 data, together with an architecturally distinct STATE analysis. The GEARS results show that high global expression agreement can coexist with weak perturbation retrieval, while matched-target analyses reveal substantial cross-context degradation in both K562-to-RPE1 and RPE1-to-K562 directions. The STATE analysis provides more limited but directionally consistent cross-architecture support across agreement endpoints, while retrieval and error-burden measures remain heterogeneous. These results illustrate why discordant endpoints should constrain the strength of a model claim rather than be collapsed into a single performance score.",
        False,
        False,
    ),
    (
        "We believe this work is particularly suited to Cell Reports Methods because its primary contribution is a reproducible and model-agnostic evaluation methodology that can be applied across perturbation-prediction systems. Rather than proposing a universal ranking of GEARS, STATE, or other models, VirtualPerturb-Audit provides a practical framework for determining whether evidence supports global expression reconstruction, perturbation-identity recovery, context transfer, directional fidelity, or only a narrower response-structure interpretation. We deliberately retain the limitations of the frozen evidence, including the use of GEARS-compatible filtered Replogle essential-screen data and the sensitivity-only interpretation of unsupported-effect metrics.",
        False,
        False,
    ),
    (
        "This manuscript is original and is not under consideration elsewhere. Information regarding data and code availability, funding, conflicts of interest, and author contributions is provided in the submission materials.",
        False,
        False,
    ),
    (
        "Thank you for considering our manuscript. We believe VirtualPerturb-Audit will be useful to researchers developing, benchmarking, and interpreting perturbation-response models, and we would be grateful for the opportunity to have the work considered for publication in Cell Reports Methods.",
        False,
        False,
    ),
    ("Sincerely,", False, False),
    ("Yu Zhang", True, False),
    ("Corresponding Author", False, False),
    ("Department of Ophthalmology", False, False),
    ("The Second Affiliated Hospital of Wenzhou Medical University", False, False),
    ("Wenzhou, Zhejiang Province, China", False, False),
    ("Email: zhangyu1@wzhealth.com", False, False),
]

MD_PATHS = [
    ROOT / "submission/COVER_LETTER_CELL_REPORTS_METHODS.md",
    ROOT / "submission/COVER_LETTER_CELL_REPORTS_METHODS_FINAL.md",
    ROOT / "submission/cell_reports_methods/04_cover_letter/COVER_LETTER_CELL_REPORTS_METHODS.md",
    ROOT / "submission/cell_reports_methods/final/cover_letter/COVER_LETTER_CELL_REPORTS_METHODS_v1.0.md",
]

DOCX_PATHS = [
    ROOT / "submission/COVER_LETTER_CELL_REPORTS_METHODS.docx",
    ROOT / "submission/COVER_LETTER_CELL_REPORTS_METHODS_v1.0.docx",
    ROOT / "submission/cell_reports_methods/04_cover_letter/COVER_LETTER_CELL_REPORTS_METHODS.docx",
    ROOT / "submission/cell_reports_methods/final/cover_letter/COVER_LETTER_CELL_REPORTS_METHODS_v1.0.docx",
]


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    for text, bold, _italic in DOCX_PARAGRAPHS:
        add_paragraph(doc, text, bold=bold)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    for path in MD_PATHS:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(MARKDOWN_TEXT.rstrip() + "\n", encoding="utf-8")
    for path in DOCX_PATHS:
        build_docx(path)
    print("updated cover letter markdown/docx files")


if __name__ == "__main__":
    main()
