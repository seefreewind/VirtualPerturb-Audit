#!/usr/bin/env python3
"""Build CRM manuscript v1.6 as a combined full-text Word manuscript.

The v1.6 file is assembled from the current v1.5 manuscript text and the
approved main Figure 1-5 assets. Figures are inserted near their first Results
mention while the figure legends remain in the manuscript back matter.
"""

from __future__ import annotations

import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
REPORTS = ROOT / "reports"
FIGURES = ROOT / "figures" / "main"

SOURCE_MD = MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.md"
OUT_MD = MANUSCRIPT / "CRM_MANUSCRIPT_v1.6_FULL.md"
OUT_DOCX = MANUSCRIPT / "CRM_MANUSCRIPT_v1.6_FULL.docx"

FIGURE_FILES = {
    "Figure 1": FIGURES / "Figure1.png",
    "Figure 2": FIGURES / "Figure2.png",
    "Figure 3": FIGURES / "Figure3.png",
    "Figure 4": FIGURES / "Figure4.png",
    "Figure 5": FIGURES / "Figure5.png",
}


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def build_markdown() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")
    text = re.sub(r"Draft version: .*", "Draft version: CRM_MANUSCRIPT_v1.6_FULL", text, count=1)
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    text = re.sub(r"Generated: .*", f"Generated: {generated}", text, count=1)
    write(OUT_MD, text)


def apply_run_style(run, size: int = 10, bold: bool | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.bold = True


def add_plain_paragraph(doc: Document, text: str) -> None:
    clean = text.replace("**", "").replace("`", "")
    p = doc.add_paragraph(clean)
    p.paragraph_format.space_after = Pt(6)


def add_figure(doc: Document, label: str) -> None:
    image = FIGURE_FILES[label]
    if not image.exists():
        raise FileNotFoundError(f"Missing main figure image: {image}")
    if label == "Figure 4":
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Inches(6.2))
    cap = doc.add_paragraph(label)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for run in cap.runs:
        apply_run_style(run, size=8, bold=True)


def build_docx() -> None:
    doc = Document()
    style_doc(doc)
    inserted: set[str] = set()
    in_figure_legends = False

    lines = OUT_MD.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("## Figure Legends"):
            in_figure_legends = True

        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            apply_run_style(run, size=18, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for run in p.runs:
                        apply_run_style(run, size=8, bold=True)
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[: len(header)]):
                    cells[j].text = val
                    cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cells[j].paragraphs:
                        for run in p.runs:
                            apply_run_style(run, size=7)
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            add_plain_paragraph(doc, line)
            if not in_figure_legends:
                for label in FIGURE_FILES:
                    if label in line and label not in inserted:
                        add_figure(doc, label)
                        inserted.add(label)
                        break
        i += 1

    missing = sorted(set(FIGURE_FILES) - inserted)
    if missing:
        raise RuntimeError(f"Figures not inserted: {missing}")
    doc.save(OUT_DOCX)


def render_docx() -> str:
    renderer = Path(
        "/Users/zy/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/render_docx.py"
    )
    out = REPORTS / "docx_qc_v16_full_pages"
    proc = subprocess.run([sys.executable, str(renderer), str(OUT_DOCX), "--output_dir", str(out)], text=True, capture_output=True)
    return f"Render return code: {proc.returncode}; output dir: {out}; stdout: {proc.stdout.strip()}; stderr: {proc.stderr.strip()}"


def write_qc(render_status: str) -> None:
    write(
        REPORTS / "CRM_MANUSCRIPT_V16_FULL_QC.md",
        f"""# CRM Manuscript v1.6 Full QC

Status: PASS

## Outputs

- `manuscript/CRM_MANUSCRIPT_v1.6_FULL.md`
- `manuscript/CRM_MANUSCRIPT_v1.6_FULL.docx`

## Figure inclusion

| Figure | Source image | Inserted near first Results mention |
|---|---|---|
| Figure 1 | `figures/main/Figure1.png` | YES |
| Figure 2 | `figures/main/Figure2.png` | YES |
| Figure 3 | `figures/main/Figure3.png` | YES |
| Figure 4 | `figures/main/Figure4.png` | YES |
| Figure 5 | `figures/main/Figure5.png` | YES |

## Render

{render_status}
""",
    )


def main() -> None:
    build_markdown()
    build_docx()
    status = render_docx()
    if "Render return code: 0" not in status:
        raise RuntimeError(status)
    write_qc(status)
    print(status)


if __name__ == "__main__":
    main()
