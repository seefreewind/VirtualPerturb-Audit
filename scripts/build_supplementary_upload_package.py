#!/usr/bin/env python3
"""Build a clean supplementary-material upload folder for journal submission."""

from __future__ import annotations

import csv
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "supplementary_upload"
SUPP_MD = ROOT / "manuscript" / "CRM_SUPPLEMENT_v1.3.md"

SUPPLEMENTARY_FILES = [
    (
        ROOT / "figures/supplementary/phase2c_endpoint_heatmap.pdf",
        "Supplementary_Figure_S1_phase2c_endpoint_heatmap.pdf",
        "Supplementary Figure S1",
    ),
    (
        ROOT / "figures/supplementary/phase2c_endpoint_heatmap.png",
        "Supplementary_Figure_S1_phase2c_endpoint_heatmap.png",
        "Supplementary Figure S1 preview",
    ),
    (
        ROOT / "figures/supplementary/phase2c_retrieval_rank_distribution.pdf",
        "Supplementary_Figure_S2_phase2c_retrieval_rank_distribution.pdf",
        "Supplementary Figure S2",
    ),
    (
        ROOT / "figures/supplementary/phase2c_retrieval_rank_distribution.png",
        "Supplementary_Figure_S2_phase2c_retrieval_rank_distribution.png",
        "Supplementary Figure S2 preview",
    ),
    (
        ROOT / "figures/supplementary/state_matched_leave_one_out.pdf",
        "Supplementary_Figure_S3_state_matched_leave_one_out.pdf",
        "Supplementary Figure S3",
    ),
    (
        ROOT / "figures/supplementary/state_matched_leave_one_out.png",
        "Supplementary_Figure_S3_state_matched_leave_one_out.png",
        "Supplementary Figure S3 preview",
    ),
    (
        ROOT / "submission/KEY_RESOURCES_TABLE_FINAL.xlsx",
        "Key_Resources_Table.xlsx",
        "Key resources table",
    ),
    (
        ROOT / "submission/SOURCE_DATA_MANIFEST.tsv",
        "Source_Data_Manifest.tsv",
        "Source data manifest",
    ),
]


def reset_out() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for path in OUT.iterdir():
        if path.is_file():
            path.unlink()


def set_normal_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 10.5)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9 if style is None else 10)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [row for i, row in enumerate(rows) if i != 1 or not all(set(cell) <= {"-", ":", " "} for cell in row)]
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(7)
            run.bold = r == 0
    doc.add_paragraph()


def build_docx(md: str, out_path: Path) -> None:
    doc = Document()
    set_normal_style(doc)
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            add_paragraph(doc, line[2:].strip(), "Heading 1")
        elif line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), "Heading 2")
        elif line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), "Heading 3")
        elif line.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = "Arial"
            run.font.size = Pt(9)
        else:
            add_paragraph(doc, line)
    flush_table()
    doc.save(out_path)


def copy_files() -> list[dict[str, str]]:
    rows = []
    md_text = SUPP_MD.read_text(encoding="utf-8")
    md_out = OUT / "Supplementary_Information.md"
    docx_out = OUT / "Supplementary_Information.docx"
    md_out.write_text(md_text, encoding="utf-8")
    build_docx(md_text, docx_out)
    rows.append({"file": md_out.name, "role": "Supplementary Information source text"})
    rows.append({"file": docx_out.name, "role": "Supplementary Information upload document"})
    for src, name, role in SUPPLEMENTARY_FILES:
        dst = OUT / name
        shutil.copy2(src, dst)
        rows.append({"file": name, "role": role})
    return rows


def write_manifest(rows: list[dict[str, str]]) -> None:
    manifest = OUT / "SUPPLEMENTARY_UPLOAD_MANIFEST.tsv"
    with manifest.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=["file", "role"], delimiter="\t")
        writer.writeheader()
        writer.writerows(rows)


def write_zip() -> Path:
    zip_path = OUT.parent / "VirtualPerturb_Audit_supplementary_upload.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUT.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def main() -> None:
    reset_out()
    rows = copy_files()
    write_manifest(rows)
    zip_path = write_zip()
    print(f"folder={OUT}")
    print(f"zip={zip_path}")


if __name__ == "__main__":
    main()
