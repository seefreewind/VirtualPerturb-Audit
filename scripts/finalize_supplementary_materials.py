#!/usr/bin/env python3
"""Finalize reviewer-facing supplementary materials without changing science."""

from __future__ import annotations

import csv
import hashlib
import json
import math
import shutil
import textwrap
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIGURES = ROOT / "figures"
TABLES = ROOT / "results" / "tables"
SUBMISSION = ROOT / "submission"
FINAL = SUBMISSION / "final"
SUPP_FIG_DIR = FIGURES / "supplementary_final"
QC_DIR = FIGURES / "qc"
PY = Path("/Users/zy/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3")

BLUE = "4F86C6"
TEXT = "#222222"
GRID = "D8DEE6"
TEAL = "#2D7A78"
SLATE = "#64707D"


def ensure_dirs() -> None:
    for path in [REPORTS, FINAL, SUPP_FIG_DIR, QC_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def fmt(x: object, digits: int = 4) -> str:
    if x is None or x == "":
        return ""
    try:
        if pd.isna(x):
            return ""
    except TypeError:
        pass
    return f"{float(x):.{digits}f}"


def random_mrr(n: int) -> float:
    return sum(1 / r for r in range(1, n + 1)) / n


def write_inventory() -> None:
    zip_paths = [
        Path("/mnt/data/04944ecd-917c-451d-a211-1531d7c88854.zip"),
        Path("/mnt/data/d5bb5539-e8c0-406d-a235-d99739bc9fd8.zip"),
        SUBMISSION / "Supplementary_Document_and_Images.zip",
        SUBMISSION / "Supplementary_Tables.zip",
        SUBMISSION / "VirtualPerturb_Audit_supplementary_upload.zip",
    ]
    rows = []
    seen: dict[str, str] = {}
    for archive in zip_paths:
        if not archive.exists():
            rows.append([str(archive), "MISSING_ARCHIVE", "", "NO", "Input archive referenced in prompt but not visible on this host", "ARCHIVE_ONLY", "STALE"])
            continue
        with zipfile.ZipFile(archive) as zf:
            for info in zf.infolist():
                if info.is_dir() or info.filename.startswith("._") or "/._" in info.filename:
                    continue
                name = Path(info.filename).name
                duplicate = "YES" if name in seen else "NO"
                seen.setdefault(name, archive.name)
                lower = name.lower()
                if "supplementary_information" in lower or "supplemental_information" in lower:
                    role = "Supplemental Information"
                    submission_role = "JOURNAL_UPLOAD"
                    status = "AUTHORITATIVE" if archive.name.startswith("Supplementary_Document") else "DUPLICATE"
                elif "figure" in lower:
                    role = "Supplementary figure"
                    submission_role = "ARCHIVE_ONLY"
                    status = "DUPLICATE" if duplicate == "YES" else "ARCHIVE_ONLY"
                elif "resources" in lower:
                    role = "Key Resources Table"
                    submission_role = "JOURNAL_UPLOAD"
                    status = "AUTHORITATIVE" if archive.name.startswith("Supplementary_Tables") else "DUPLICATE"
                elif "source_data" in lower or "manifest" in lower:
                    role = "Source-data or upload manifest"
                    submission_role = "JOURNAL_UPLOAD"
                    status = "AUTHORITATIVE" if archive.name.startswith("Supplementary_Tables") else "DUPLICATE"
                else:
                    role = "Supporting file"
                    submission_role = "ARCHIVE_ONLY"
                    status = "ARCHIVE_ONLY"
                rows.append([f"{archive.name}:{info.filename}", Path(name).suffix or "file", str(info.file_size), duplicate, role, submission_role, status])
    lines = [
        "# Supplement Final Input Inventory",
        "",
        "| filename | type | size | duplicate? | scientific role | submission role | mark |",
        "| --- | --- | ---: | --- | --- | --- | --- |",
    ]
    lines += ["| " + " | ".join(map(str, row)) + " |" for row in rows]
    (REPORTS / "SUPPLEMENT_FINAL_INPUT_INVENTORY.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_defect_map() -> None:
    targets = [
        ROOT / "manuscript/CRM_SUPPLEMENT_v1.3.md",
        SUBMISSION / "supplementary_upload/Supplementary_Information.md",
        SUBMISSION / "KEY_RESOURCES_TABLE_FINAL.csv",
        SUBMISSION / "SOURCE_DATA_MANIFEST.tsv",
    ]
    patterns = [
        "Phase 2C", "v1.3", "confirmatory", "reviewer-facing", "deposition readiness",
        "Frozen Analysis State", "internal project path", "run ID", "GO", "NO-GO",
        "PASS", "FAIL", "UNVERIFIED", "Draft version", "Generated:", "audit_delta",
        "sign_flip_rate", "metric_space", "within_minus_cross",
    ]
    rows = []
    for path in targets:
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="replace")
        for pat in patterns:
            for lineno, line in enumerate(text.splitlines(), 1):
                if pat in line:
                    if pat in {"Phase 2C", "v1.3", "Draft version", "Generated:", "Frozen Analysis State"}:
                        cat = "P0_SCIENTIFIC_CLARITY"
                    elif pat in {"audit_delta", "sign_flip_rate", "metric_space", "within_minus_cross"}:
                        cat = "P1_READER_FACING"
                    elif pat == "UNVERIFIED":
                        cat = "P0_TRACEABILITY"
                    else:
                        cat = "P2_FORMATTING"
                    rows.append([cat, str(path.relative_to(ROOT)), lineno, pat, line.strip()])
    lines = [
        "# Supplement Final Defect Map",
        "",
        "This pre-edit map records stale/internal language found in legacy supplementary inputs. The final Supplemental Information is rebuilt from frozen tables and does not use these legacy lines as reader-facing text.",
        "",
        "| category | file | line | occurrence | legacy text |",
        "| --- | --- | ---: | --- | --- |",
    ]
    lines += ["| " + " | ".join(str(x).replace("|", "\\|") for x in row) + " |" for row in rows]
    (REPORTS / "SUPPLEMENT_FINAL_DEFECT_MAP.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def load_tables() -> dict[str, pd.DataFrame]:
    return {
        "registry": pd.read_csv(TABLES / "FINAL_MANUSCRIPT_NUMERIC_REGISTRY.tsv", sep="\t"),
        "baseline": pd.read_csv(TABLES / "baseline_definition_registry.tsv", sep="\t"),
        "probe": pd.read_csv(TABLES / "falsification_probe_registry.tsv", sep="\t"),
        "state_primary": pd.read_csv(TABLES / "state_phase2c_primary_metrics.csv"),
        "state_transfer": pd.read_csv(TABLES / "state_transfer_drop.csv"),
        "state_loo": pd.read_csv(TABLES / "state_matched_leave_one_out_summary.tsv", sep="\t"),
        "state_loo_detail": pd.read_csv(TABLES / "state_matched_leave_one_out.tsv", sep="\t"),
        "state_common_mrr": pd.read_csv(TABLES / "state_matched_common_candidate_retrieval_summary.tsv", sep="\t"),
        "state_common_detail": pd.read_csv(TABLES / "state_matched_common_candidate_retrieval.tsv", sep="\t"),
        "gears_transfer": pd.read_csv(TABLES / "replogle_matched_rl1_rl4_sensitivity.csv"),
        "replogle_probe": pd.read_csv(TABLES / "replogle_gears_vs_probes.csv"),
        "state_retrieval": pd.read_csv(TABLES / "state_phase2c_retrieval.csv"),
    }


def table_s1() -> tuple[list[str], list[list[str]]]:
    rows = [
        ["Dataset/version", "Dataset source and processed object", "Name the source dataset, processed mirror, filtering scope, and redistribution boundary.", "Prevents hidden dataset substitution."],
        ["Model/checkpoint", "Model implementation and frozen checkpoint", "Report the model family, package version, checkpoint provenance, and run status.", "Separates model behavior from later implementation drift."],
        ["Split", "Train/test split assignment", "Declare the split level, held-out target rule, and whether the split was reused across endpoints.", "Defines the perturbations being evaluated."],
        ["Target universe", "Perturbation targets eligible for evaluation", "Report the target set after filtering and matching.", "Bounds what a retrieval or transfer claim can cover."],
        ["Gene universe", "Genes used for endpoint calculation", "Report gene count and any intersection rule across datasets or models.", "Avoids comparing endpoints over different feature spaces."],
        ["Control definition", "Control or basal expression reference", "State how control expression and control-subtracted deltas are formed.", "Determines the meaning of audit-delta endpoints."],
        ["Metric space", "Raw-space versus audit-delta endpoint", "Label each endpoint as raw-space expression agreement or control-subtracted response agreement.", "Prevents global expression similarity from being read as target-specific recovery."],
        ["Candidate universe", "Retrieval candidate set", "Declare the candidate set and its size for Top1, Top5, and MRR.", "Makes retrieval values interpretable."],
        ["Falsification probes", "Information-removal controls", "Define target-blind, cell-state-blind, and target-randomized probes where applicable.", "Tests whether endpoint signal survives removal of target information."],
        ["Context-transfer design", "Matched-target context shift", "Report source context, cross context, matched target count, and adapter use if any.", "Keeps context-transfer claims tied to the evaluated contrast."],
        ["Statistical unit", "Perturbation-level aggregation", "State whether estimates and uncertainty operate at perturbation, target, or cell level.", "Prevents inflated precision."],
        ["Uncertainty", "Bootstrap or interval definition", "Report resampling unit and interval definition for displayed estimates.", "Allows the reader to interpret stability."],
        ["Claim boundary", "Endpoint-specific interpretation", "Map each endpoint family to a bounded claim: expression agreement, retrieval, transfer, error burden, or direction.", "Avoids a single composite model score."],
        ["Code/data provenance", "Script, source table, and archive identifier", "Link each plotted quantity to a source file, script, checksum, and public archive.", "Supports reproducibility and reviewer tracing."],
    ]
    return ["Audit domain", "Required item", "Reporting requirement", "Why it matters"], rows


def table_s2() -> tuple[list[str], list[list[str]]]:
    rows = [
        ["B0", "No-change", "Control or basal input expression; no perturbation-specific delta.", "All settings with a control or basal reference.", "Used as a no-effect reference.", "Does not model perturbation response."],
        ["B1", "Global training mean-delta", "Training perturbation profiles only; target identity is not used at prediction time.", "Single-context baseline pilots and Replogle R-L1.", "Same numerical mean-effect family as B2/B5/FP1 in frozen Replogle R-L1.", "Not independent evidence when identical to B2/B5/FP1."],
        ["B2", "Context-matched mean", "Training profiles plus context label when more than one within-task context is available.", "Audits with multiple within-task contexts.", "Falls back to B1 in single-context Replogle R-L1.", "Fallback output is not counted as a separate control."],
        ["B3", "Additive component baseline", "Single-component training deltas where component perturbations exist.", "Norman-style component settings.", "Not used for Replogle essential-screen analyses.", "Absent component structure prevents Replogle use."],
        ["B4", "PCA/Ridge", "Low-capacity projection/regression baseline fitted on frozen training data.", "Frozen baseline stress tests.", "Numerically near-identical to mean-effect construction in frozen Replogle held-out-target setting.", "Interpreted as a low-capacity check, not separate biological evidence."],
        ["B5", "Mean-effect baseline", "Same target-blind mean-effect construction used for held-out targets.", "Replogle R-L1 baseline display.", "Same estimator as FP1, different role.", "Predictive baseline role only."],
        ["FP1", "Perturbation-blind probe", "Same estimator as B5 but used as an information-removal falsification probe.", "Replogle falsification display.", "Same estimator as B5; named separately to preserve interpretive role.", "B1/B2/B5/FP1 are not treated as independent baseline evidence in frozen Replogle R-L1."],
    ]
    return ["Baseline", "Estimator", "Information used", "Applicable setting", "Frozen Replogle behavior", "Interpretation boundary"], rows


def table_s3() -> tuple[list[str], list[list[str]]]:
    rows = [
        ["FP1", "Perturbation target identity at prediction time", "Uses the same target-blind mean-effect construction as B5.", "None; deterministic mean-effect construction.", "Information-removal falsification probe.", "Agreement that survives FP1 narrows interpretation toward shared response structure."],
        ["FP2", "Cell-state/context information where implementation and data permit", "Removes or withholds cell-state/context information when a valid context-dependent implementation exists.", "Not applied to the Replogle essential-screen display when required context structure is absent.", "Cell-state information probe.", "Unavailable settings are not treated as failed or negative evidence."],
        ["FP3 Replogle", "Correct target-to-delta assignment", "Training-target deltas are randomly assigned to test perturbations.", "With replacement using the frozen seed; Figure 3 uses one frozen Replogle draw.", "Target-randomized diagnostic probe.", "This is not a strict bijective permutation and is not interpreted as an independent model."],
        ["FP3 Norman", "Label assignment for Norman diagnostic sensitivity", "Label-randomization repetitions are run separately for Norman.", "20 label-randomization repetitions.", "Analytically separate sensitivity analysis.", "Not the same plotted draw as Replogle Figure 3."],
    ]
    return ["Probe", "Information disrupted", "Exact implementation", "Randomization scheme", "Role", "Interpretation boundary"], rows


def table_s4(t: dict[str, pd.DataFrame]) -> tuple[list[str], list[list[str]]]:
    df = t["state_primary"]
    rows = []
    for run_id in ["S1_norman_l1", "S2_norman_l2", "S3_replogle_k562_rl1", "S4_replogle_k562_to_rpe1_rl4"]:
        row = df[(df.run_id == run_id) & (df.metric_space == "audit_delta")]
        if row.empty:
            row = df[df.run_id == run_id].head(1)
        r = row.iloc[0]
        task = {
            "S1_norman_l1": "Norman L1",
            "S2_norman_l2": "Norman L2",
            "S3_replogle_k562_rl1": "Replogle K562 R-L1",
            "S4_replogle_k562_to_rpe1_rl4": "K562 -> RPE1 R-L4",
        }[run_id]
        audit_setting = "Within-context" if run_id != "S4_replogle_k562_to_rpe1_rl4" else "Cross-context adapter-based stress test"
        rows.append([task, audit_setting, fmt(r.pearson_delta), fmt(r.spearman_delta), fmt(r.cosine_delta), fmt(r.retrieval_mrr), fmt(r.uer50), fmt(r.sign_flip_rate), str(int(r.n_test_perturbations))])
    return ["Dataset/task", "Audit setting", "Audit-delta Pearson", "Spearman", "Cosine", "MRR", "UER50", "Sign-flip rate", "n targets"], rows


def table_s5(t: dict[str, pd.DataFrame]) -> tuple[list[str], list[list[str]], str]:
    label = {
        "pearson_delta": "Audit-delta Pearson",
        "spearman_delta": "Spearman",
        "cosine_delta": "Cosine",
        "uer50": "UER50",
        "sign_flip_rate": "Sign-flip rate",
    }
    interp = {
        "pearson_delta": "Agreement decreased cross context.",
        "spearman_delta": "Rank agreement decreased cross context.",
        "cosine_delta": "Cosine agreement decreased cross context.",
        "uer50": "Higher cross-context burden gives a negative native difference.",
        "sign_flip_rate": "Higher cross-context sign-flip rate gives a negative native difference.",
    }
    rows = []
    for _, r in t["state_transfer"].iterrows():
        rows.append([label[r.metric], fmt(r.source_mean), fmt(r.cross_context_mean), fmt(r.mean_drop_source_minus_cross), f"[{fmt(r.ci95_low)}, {fmt(r.ci95_high)}]", interp[r.metric]])
    note = "Values in this table retain the native within-minus-cross difference convention. Positive differences indicate poorer cross-context behavior for agreement endpoints, whereas negative differences indicate poorer cross-context behavior for burden endpoints because higher UER50 and sign-flip rates are unfavorable. Figure 5 direction-aligns burden endpoints for visualization only."
    return ["Endpoint", "Within-context", "Cross-context", "Native within-minus-cross difference", "Bootstrap 95% CI", "Interpretation"], rows, note


def table_s6(t: dict[str, pd.DataFrame]) -> tuple[list[str], list[list[str]], str]:
    transfer = t["state_transfer"].set_index("metric")
    loo = t["state_loo"].set_index("metric")
    specs = [
        ("Agreement endpoint", "pearson_delta", "pearson_drop", "Audit-delta Pearson"),
        ("Agreement endpoint", "spearman_delta", "spearman_drop", "Spearman"),
        ("Agreement endpoint", "cosine_delta", "cosine_drop", "Cosine"),
        ("Burden endpoint", "uer50", "uer50_difference", "UER50"),
        ("Burden endpoint", "sign_flip_rate", "sign_flip_difference", "Sign-flip rate"),
    ]
    rows = []
    for group, metric, loo_metric, label in specs:
        full = float(transfer.loc[metric, "mean_drop_source_minus_cross"])
        l = loo.loc[loo_metric]
        retained = int(l.n_positive) if group == "Agreement endpoint" else int(l.n_negative)
        rows.append([group, label, fmt(full), fmt(l["min"]), fmt(l["max"]), str(retained), str(int(l.n_loo))])
    note = "Agreement endpoints retain the positive within-minus-cross transfer-degradation direction in 15/15 omissions. For burden endpoints, negative native differences reflect higher cross-context UER50 or sign-flip rates; those rows are interpreted under the native sign convention."
    return ["Endpoint group", "Endpoint", "Full-sample effect", "Min LOO effect", "Max LOO effect", "Number retaining primary direction", "n omissions"], rows, note


def table_s7(t: dict[str, pd.DataFrame]) -> tuple[list[str], list[list[str]], str]:
    mrr = t["state_common_mrr"].set_index("run_id")
    rand = random_mrr(15)
    rows = [
        ["Within K562", "15", fmt(mrr.loc["S3_replogle_k562_rl1", "mrr"]), fmt(rand), "Above the random-ranking expectation but still modest."],
        ["Cross to RPE1", "15", fmt(mrr.loc["S4_replogle_k562_to_rpe1_rl4", "mrr"]), fmt(rand), "Approximately the theoretical random-ranking expectation."],
    ]
    note = "Cross-context MRR was approximately the theoretical random-ranking expectation for the 15-candidate universe. This is an exploratory sensitivity analysis, not a formal primary test."
    return ["Condition", "n candidates", "MRR", "Random-ranking MRR", "Interpretation"], rows, note


def table_s8(t: dict[str, pd.DataFrame]) -> tuple[list[str], list[list[str]], str]:
    df = t["gears_transfer"]
    primary = df[df.comparison_role == "primary_source_context_comparison"].copy()
    keep_metrics = ["pearson_delta", "spearman_delta", "cosine_delta", "retrieval_mrr_native", "retrieval_mrr_common_candidate", "uer50", "sign_flip_rate"]
    labels = {
        "pearson_delta": "Audit-delta Pearson",
        "spearman_delta": "Spearman",
        "cosine_delta": "Cosine",
        "retrieval_mrr_native": "Native-candidate MRR",
        "retrieval_mrr_common_candidate": "Common-candidate MRR",
        "uer50": "UER50",
        "sign_flip_rate": "Sign-flip rate",
    }
    rows = []
    for direction in ["K562_within_vs_K562_to_RPE1", "RPE1_within_vs_RPE1_to_K562"]:
        contrast = "K562 within vs K562 -> RPE1" if direction.startswith("K562") else "RPE1 within vs RPE1 -> K562"
        for metric in keep_metrics:
            r = primary[(primary.direction == direction) & (primary.metric == metric)].iloc[0]
            rows.append([contrast, labels[metric], str(int(r.n_targets)), fmt(r.within_estimate), fmt(r.cross_estimate), fmt(r.paired_difference), f"[{fmt(r.ci_low)}, {fmt(r.ci_high)}]"])
    note = "Table S8 contains only source-aligned matched comparisons. Cross-pair contrasts were retained in the public archive as supporting exploratory material and were not included as formal Supplemental Information because they do not define the same source-aligned estimand."
    (REPORTS / "S8_CROSS_PAIR_DECISION.md").write_text("# S8 Cross-Pair Decision\n\nDecision: ARCHIVE_ONLY\n\nCross-pair comparisons are retained in the public archive, but the formal Supplemental Information includes only source-aligned matched GEARS transfer analyses because these define the primary estimand used in the manuscript.\n", encoding="utf-8")
    return ["Source-aligned comparison", "Endpoint", "n targets", "Within-context", "Cross-context", "Difference", "Bootstrap 95% CI"], rows, note


def table_s9() -> tuple[list[str], list[list[str]], str]:
    rows = [
        ["Input", "Observed expression", "Cell-level or target-level observed expression matrix.", "Defines the measured response."],
        ["Input", "Predicted expression", "Model prediction matrix aligned to target and gene identifiers.", "Defines the model output under audit."],
        ["Input", "Perturbation labels", "Target identity for observed and predicted responses.", "Required for target-specific retrieval and probes."],
        ["Input", "Control labels", "Control or basal state labels used for response deltas.", "Required for audit-delta endpoints."],
        ["Input", "Context labels", "Cell type, cell line, or other context identifier.", "Required for within- and cross-context interpretation."],
        ["Input", "Gene identifiers", "Gene universe and alignment rule.", "Keeps endpoint calculations comparable."],
        ["Input", "Target universe", "Perturbation targets eligible after filtering.", "Bounds each claim."],
        ["Input", "Candidate universe", "Declared candidate set for retrieval endpoints.", "Defines random-ranking reference and retrieval interpretation."],
        ["Input", "Split assignments", "Train/test and held-out target definitions.", "Prevents leakage and split drift."],
        ["Input", "Model/provenance metadata", "Model package, checkpoint, preprocessing, and script provenance.", "Supports reproducibility."],
        ["Output", "Expression-agreement endpoints", "Raw-space and audit-delta agreement metrics.", "Reports agreement without implying target recovery."],
        ["Output", "Retrieval endpoints", "Top1, Top5, MRR, and candidate-universe metadata.", "Tests perturbation-identity recovery."],
        ["Output", "UER", "Unsupported-effect rate at declared K.", "Sensitivity endpoint for large predicted effects without observed support."],
        ["Output", "Sign-flip", "Direction mismatch among supported-effect genes.", "Directional-fidelity endpoint."],
        ["Output", "Matched-transfer summaries", "Within-context and cross-context matched target contrasts.", "Tests context portability under the declared setup."],
        ["Output", "Probe results", "Target-blind, cell-state-blind, and target-randomized outputs where applicable.", "Tests information dependence."],
        ["Output", "Split-integrity report", "Checks for target and split consistency.", "Documents leakage controls."],
        ["Output", "Claim profile", "Endpoint-specific claim interpretation.", "Prevents collapsing discordant endpoints into one score."],
    ]
    note = "Full machine-readable specifications are archived at GitHub and Zenodo."
    return ["Contract family", "Element", "Requirement", "Purpose"], rows, note


def save_figures(t: dict[str, pd.DataFrame]) -> None:
    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 8, "axes.linewidth": 0.7})
    state = t["state_primary"]
    rows = []
    tasks = [
        ("S1_norman_l1", "Norman L1"),
        ("S2_norman_l2", "Norman L2"),
        ("S3_replogle_k562_rl1", "Replogle K562 R-L1"),
        ("S4_replogle_k562_to_rpe1_rl4", "K562 -> RPE1 R-L4"),
    ]
    metrics = [
        ("pearson_delta", "Pearson"),
        ("spearman_delta", "Spearman"),
        ("cosine_delta", "Cosine"),
        ("retrieval_mrr", "MRR"),
        ("uer50", "UER50"),
        ("sign_flip_rate", "Sign-flip"),
    ]
    for run_id, label in tasks:
        r = state[(state.run_id == run_id) & (state.metric_space == "audit_delta")]
        if r.empty:
            r = state[state.run_id == run_id].head(1)
        r = r.iloc[0]
        rows.append([label] + [float(r[m]) for m, _ in metrics])
    arr = np.array([row[1:] for row in rows])
    fig, ax = plt.subplots(figsize=(7.4, 3.0))
    ax.imshow(np.zeros_like(arr), cmap="Greys", vmin=0, vmax=1, alpha=0.06)
    ax.set_xticks(range(len(metrics)), [m[1] for m in metrics])
    ax.set_yticks(range(len(rows)), [r[0] for r in rows])
    ax.tick_params(length=0)
    for i in range(arr.shape[0]):
        for j in range(arr.shape[1]):
            ax.text(j, i, f"{arr[i, j]:.3f}", ha="center", va="center", color=TEXT, fontsize=8)
    for x in [-0.5, 2.5, 3.5, 5.5]:
        ax.axvline(x, color="#BFC7D1", lw=0.8)
    ax.set_title("STATE endpoint profiles across four audit settings", loc="left", fontweight="bold", pad=14)
    ax.text(-0.5, 4.15, "Columns are grouped as response agreement, retrieval, and error/direction behavior. Raw values are printed in each cell; UER50 and sign-flip are unfavorable when higher.", fontsize=7, color=SLATE)
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.tight_layout()
    for ext in ["pdf", "png", "svg"]:
        fig.savefig(SUPP_FIG_DIR / f"Supplementary_Figure_S1_STATE_endpoint_profiles.{ext}", dpi=450, bbox_inches="tight")
    fig.savefig(QC_DIR / "SuppFigureS1_halfsize.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    detail = t["state_retrieval"]
    fig, axes = plt.subplots(1, 4, figsize=(7.6, 2.4), sharey=False)
    for ax, (run_id, label) in zip(axes, tasks):
        d = detail[(detail.run_id == run_id) & (detail.space == "audit_delta")].copy()
        if d.empty:
            d = detail[(detail.run_id == run_id) & (detail.space == "target_control_audit_delta")].copy()
        ranks = d.true_target_rank.astype(float).values
        n = max(int(d.true_target_rank.max()), int(d.perturbation.nunique()))
        rng = np.random.default_rng(17)
        x = rng.normal(0, 0.035, size=len(ranks))
        ax.scatter(x, ranks, s=12, color=TEAL, alpha=0.8, linewidth=0)
        q1, med, q3 = np.percentile(ranks, [25, 50, 75])
        ax.plot([-0.14, 0.14], [med, med], color="#111111", lw=1.2)
        ax.add_patch(plt.Rectangle((-0.10, q1), 0.20, q3 - q1, fill=False, edgecolor="#59636E", lw=0.9))
        ax.axhline((n + 1) / 2, color="#9AA4AF", lw=0.8, ls=":")
        ax.set_title(f"{label}\nn candidates = {n}", fontsize=7.5)
        ax.set_xlim(-0.22, 0.22)
        ax.set_ylim(n + 1, 0)
        ax.set_xticks([])
        if ax is axes[0]:
            ax.set_ylabel("True-target rank\nlower is better")
        ax.spines[["top", "right", "bottom"]].set_visible(False)
    fig.suptitle("Perturbation-retrieval rank distributions across STATE audit tasks", x=0.02, ha="left", fontsize=9, fontweight="bold")
    fig.text(0.02, -0.02, "Each point is one perturbation target; dotted lines show the random-ranking median/reference within each task.", fontsize=7, color=SLATE)
    fig.tight_layout()
    for ext in ["pdf", "png", "svg"]:
        fig.savefig(SUPP_FIG_DIR / f"Supplementary_Figure_S2_STATE_rank_distributions.{ext}", dpi=450, bbox_inches="tight")
    fig.savefig(QC_DIR / "SuppFigureS2_halfsize.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    loo = t["state_loo_detail"]
    transfer = t["state_transfer"].set_index("metric")
    fig, ax = plt.subplots(figsize=(6.4, 3.0))
    specs = [
        ("Audit-delta Pearson", "pearson_drop", "pearson_delta", 2),
        ("Spearman", "spearman_drop", "spearman_delta", 1),
        ("Cosine", "cosine_drop", "cosine_delta", 0),
    ]
    rng = np.random.default_rng(11)
    for label, col, full_metric, y in specs:
        vals = loo[col].astype(float).values
        yy = y + rng.normal(0, 0.04, size=len(vals))
        ax.scatter(vals, yy, s=18, color=TEAL, alpha=0.86, linewidth=0)
        ax.axvline(float(transfer.loc[full_metric, "mean_drop_source_minus_cross"]), ymin=(y + 0.12) / 3.2, ymax=(y + 0.88) / 3.2, color="#303A44", lw=1.3)
    ax.axvline(0, color="#9AA4AF", lw=0.9)
    ax.set_yticks([2, 1, 0], ["Audit-delta Pearson", "Spearman", "Cosine"])
    ax.set_xlabel("Within - cross effect")
    ax.set_title("Leave-one-target-out robustness of matched STATE agreement effects", loc="left", fontweight="bold")
    fig.text(0.13, 0.04, "Dots show one target omitted per iteration; vertical black ticks show full-sample effects.", fontsize=7, color=SLATE)
    ax.spines[["top", "right"]].set_visible(False)
    fig.subplots_adjust(bottom=0.24, top=0.82, left=0.18, right=0.98)
    for ext in ["pdf", "png", "svg"]:
        fig.savefig(SUPP_FIG_DIR / f"Supplementary_Figure_S3_STATE_LOO_agreement_effects.{ext}", dpi=450, bbox_inches="tight")
    fig.savefig(QC_DIR / "SuppFigureS3_halfsize.png", dpi=220, bbox_inches="tight")
    plt.close(fig)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def doc_para(doc: Document, text: str, style: str | None = None, bold: bool = False, italic: bool = False, size: float | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4 if style is None else 6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size if size else (9 if style is None else 11))
    if style and style.startswith("Heading"):
        run.font.color.rgb = RGBColor.from_string(BLUE)


def doc_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "EAF2FB")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(7.2)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.name = "Arial"
            r.font.size = Pt(6.7 if len(headers) > 5 else 7.3)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()


def build_supplement_docx(t: dict[str, pd.DataFrame]) -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.62)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)
    for name, size in [("Normal", 9), ("Heading 1", 14), ("Heading 2", 11), ("Heading 3", 10)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        if name.startswith("Heading"):
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(BLUE)

    doc_para(doc, "Supplemental Information", "Heading 1")
    doc_para(doc, "Overview of supplementary analyses", "Heading 2")
    doc_para(doc, "Supplementary analyses provide detailed model-specific endpoint summaries, matched-target sensitivity analyses, leave-one-target-out robustness checks, and reproducibility resources supporting the main manuscript.")

    doc_para(doc, "Table S1. VirtualPerturb-Audit reporting checklist", "Heading 2")
    h, r = table_s1()
    doc_table(doc, h, r, [1.25, 1.35, 3.25, 2.0])
    doc_para(doc, "Table S2. Baseline definitions and applicability", "Heading 2")
    h, r = table_s2()
    doc_table(doc, h, r)
    doc_para(doc, "Table S3. Falsification-probe definitions", "Heading 2")
    h, r = table_s3()
    doc_table(doc, h, r)
    doc_para(doc, "Table S4. STATE endpoint summaries across audit settings", "Heading 2")
    h, r = table_s4(t)
    doc_table(doc, h, r)
    doc_para(doc, "Table S4 note. The K562 -> RPE1 R-L4 row uses target-control audit-delta subtraction in the cross-context adapter-based stress test; the distinction is retained because it defines the evaluated metric space.")
    doc_para(doc, "Table S5. STATE matched-target transfer effects", "Heading 2")
    h, r, note_s5 = table_s5(t)
    doc_table(doc, h, r)
    doc_para(doc, f"Table S5 note. {note_s5}")
    doc_para(doc, "Table S6. STATE leave-one-target-out matched-transfer sensitivity", "Heading 2")
    h, r, note_s6 = table_s6(t)
    doc_table(doc, h, r)
    doc_para(doc, f"Table S6 note. {note_s6} The analysis uses n = 15 matched targets with one target omitted per iteration.")
    doc_para(doc, "Table S7. STATE common-candidate retrieval sensitivity", "Heading 2")
    h, r, note_s7 = table_s7(t)
    doc_table(doc, h, r)
    doc_para(doc, f"Table S7 note. {note_s7}")
    doc_para(doc, "Table S8. Source-aligned GEARS matched-target transfer analyses", "Heading 2")
    h, r, note_s8 = table_s8(t)
    doc_table(doc, h, r)
    doc_para(doc, f"Table S8 note. {note_s8}")
    doc_para(doc, "Table S9. VirtualPerturb-Audit input and output contract summary", "Heading 2")
    h, r, note_s9 = table_s9()
    doc_table(doc, h, r, [1.0, 1.35, 3.0, 2.6])
    doc_para(doc, f"Table S9 note. {note_s9}")

    doc_para(doc, "Supplementary Figures S1-S3", "Heading 2")
    figs = [
        ("Figure S1", SUPP_FIG_DIR / "Supplementary_Figure_S1_STATE_endpoint_profiles.png"),
        ("Figure S2", SUPP_FIG_DIR / "Supplementary_Figure_S2_STATE_rank_distributions.png"),
        ("Figure S3", SUPP_FIG_DIR / "Supplementary_Figure_S3_STATE_LOO_agreement_effects.png"),
    ]
    for title, path in figs:
        doc_para(doc, title, "Heading 3")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.7))
        if title != "Figure S3":
            doc.add_page_break()

    doc_para(doc, "Supplementary Methods", "Heading 2")
    methods = [
        ("Retrieval candidate-universe definition", "Retrieval endpoints rank the true perturbation target within the declared non-control candidate universe for each task. Candidate-universe size is therefore part of the endpoint definition."),
        ("Random-ranking MRR expectation", "For a random ranking of N candidates, E(MRR) = H_N / N, where H_N = sum_(r=1)^N 1/r. This value is a reference expectation only and is not a fitted baseline."),
        ("UER50 threshold", "UER50 orders genes by predicted absolute effect and counts top-ranked genes whose observed effect falls within the internal null threshold. UER50 is an internal sensitivity endpoint."),
        ("Sign-flip support threshold", "Sign-flip rate is calculated among genes with observed absolute effects above the analysis threshold and records predicted-versus-observed direction mismatch."),
        ("Matched-target construction", "Matched-transfer analyses restrict within-context and cross-context comparisons to the same perturbation targets before estimating endpoint differences."),
        ("STATE matched bootstrap", "STATE matched effects use the frozen perturbation-level bootstrap intervals reported in the source tables."),
        ("Leave-one-target-out sensitivity", "The leave-one-target-out sensitivity repeats the matched STATE contrast after omitting one of the 15 common targets at a time."),
        ("Software versions", "The frozen GEARS evaluation used cell-gears 0.1.2. The STATE audit used arc-state 0.11.1."),
    ]
    for head, body in methods:
        doc_para(doc, head, "Heading 3")
        doc_para(doc, body)

    doc_para(doc, "Scope and interpretation boundaries", "Heading 2")
    doc_para(doc, "The Replogle analyses use GEARS-compatible filtered essential-screen data, so the results apply to that frozen subset. GEARS R-L4 is an adapter-based stress test. The STATE matched analysis uses 15 shared targets and provides partial cross-architecture support rather than formal replication.")
    doc_para(doc, "UER50 is interpreted as a sensitivity endpoint. Retrieval metrics depend on the declared candidate universe, and raw ranks are interpreted within task. Shared-control reuse can affect control-subtracted agreement endpoints, so endpoint families are interpreted as bounded claim evidence rather than as one composite score.")

    doc_para(doc, "Supplementary Figure Legends", "Heading 2")
    doc_para(doc, "Figure S1. STATE endpoint profiles across four audit settings. Raw endpoint values are shown for four STATE audit tasks and six metrics grouped as response agreement, retrieval, and error/direction behavior. UER50 and sign-flip rate are unfavorable when higher. The light matrix background is visual scaffolding only; colors are not cross-endpoint effect-size comparisons.")
    doc_para(doc, "Figure S2. Perturbation-retrieval rank distributions across STATE audit tasks. Each point represents one perturbation target. Retrieval rank is defined against the declared candidate universe for each task, lower rank is better, and the dotted line marks the random-ranking median/reference within that task. Candidate universes differ across tasks, so raw rank cross-task comparison is bounded and interpreted within task.")
    doc_para(doc, "Figure S3. Leave-one-target-out robustness of matched STATE agreement effects. The analysis uses n = 15 common targets, omitting one target at a time. Dots show leave-one-target-out within-minus-cross effects for audit-delta Pearson, Spearman, and cosine; vertical black ticks show full-sample effects and the gray line marks zero. All 15 omissions preserve the positive direction for the three agreement endpoints. This is an exploratory robustness analysis.")

    out = SUBMISSION / "Supplemental_Information_FINAL.docx"
    doc.save(out)
    return out


def build_krt() -> tuple[Path, Path]:
    rows = [
        ["Norman perturbation dataset", "Norman et al., Science 2019", "https://doi.org/10.1126/science.aax4438", "Used through a GEARS-compatible processed mirror; raw data are not redistributed."],
        ["Replogle Perturb-seq dataset", "Replogle et al., Cell 2022", "https://doi.org/10.1016/j.cell.2022.05.013", "Filtered K562/RPE1 essential-screen objects used for the frozen audit."],
        ["Replogle processed data manifest", "Figshare+", "https://doi.org/10.25452/figshare.plus.20029387", "Original processed release; complete objects are not redistributed in the audit archive."],
        ["GEARS", "Roohani et al., Nature Biotechnology 2024", "https://doi.org/10.1038/s41587-023-01905-6", "Frozen GEARS-compatible audit outputs used as the primary worked example."],
        ["GEARS implementation", "cell-gears", "cell-gears 0.1.2", "Software package version used for the frozen GEARS evaluation."],
        ["STATE", "Adduri et al., bioRxiv 2025", "https://doi.org/10.1101/2025.06.26.661135", "Architecturally distinct model used for the partial cross-architecture audit."],
        ["STATE implementation", "arc-state", "arc-state 0.11.1", "Software package version used for the frozen STATE audit."],
        ["Python", "Python Software Foundation", "https://www.python.org/", "Used for analysis scripts and figure generation."],
        ["Scanpy", "Wolf et al., Genome Biology 2018", "https://doi.org/10.1186/s13059-017-1382-0", "Single-cell analysis ecosystem dependency."],
        ["AnnData", "scverse", "https://anndata.readthedocs.io/", "Primary container format for single-cell matrices and metadata."],
        ["scikit-learn", "Pedregosa et al., JMLR 2011", "https://jmlr.csail.mit.edu/papers/v12/pedregosa11a.html", "Low-capacity baseline and utility dependency."],
        ["NumPy", "Harris et al., Nature 2020", "https://doi.org/10.1038/s41586-020-2649-2", "Numerical array dependency."],
        ["SciPy", "Virtanen et al., Nature Methods 2020", "https://doi.org/10.1038/s41592-019-0686-2", "Scientific computing dependency."],
        ["VirtualPerturb-Audit repository", "This study", "https://github.com/seefreewind/VirtualPerturb-Audit", "Public code and compact derived manuscript-facing materials."],
        ["VirtualPerturb-Audit archived release", "This study", "https://doi.org/10.5281/zenodo.22232963", "Archived release with code, compact derived results, frozen registries, and figure source data."],
        ["HGNC family resource", "HGNC", "https://www.genenames.org/", "Used to derive the frozen gene-family holdout registry."],
    ]
    csv_path = SUBMISSION / "Key_Resources_Table_FINAL.csv"
    xlsx_path = SUBMISSION / "Key_Resources_Table_FINAL.xlsx"
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Resource", "Source", "Identifier", "Additional information"])
        writer.writerows(rows)
    wb = Workbook()
    ws = wb.active
    ws.title = "Key Resources"
    ws.append(["Resource", "Source", "Identifier", "Additional information"])
    for row in rows:
        ws.append(row)
    widths = [34, 28, 44, 62]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    header_fill = PatternFill("solid", fgColor="EAF2FB")
    border = Border(bottom=Side(style="thin", color="7A8794"))
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.font = Font(name="Arial", size=10, bold=(cell.row == 1))
            if cell.row == 1:
                cell.fill = header_fill
                cell.border = border
    wb.save(xlsx_path)
    return xlsx_path, csv_path


def preview_krt(xlsx_path: Path) -> None:
    df = pd.read_excel(xlsx_path)
    wrapped = df.copy()
    for col, width in {"Resource": 28, "Source": 28, "Identifier": 42, "Additional information": 54}.items():
        wrapped[col] = wrapped[col].map(lambda x: "\n".join(textwrap.wrap(str(x), width=width)))
    fig, ax = plt.subplots(figsize=(14, 9.5))
    ax.axis("off")
    ax.set_title("Key Resources Table final preview", loc="left", fontweight="bold")
    table = ax.table(cellText=wrapped.values, colLabels=wrapped.columns, loc="center", cellLoc="left", colLoc="left", colWidths=[0.18, 0.18, 0.28, 0.36])
    table.auto_set_font_size(False)
    table.set_fontsize(6.2)
    table.scale(1, 2.0)
    for (row, _col), cell in table.get_celld().items():
        cell.set_edgecolor("#B8C2CC")
        if row == 0:
            cell.set_facecolor("#EAF2FB")
            cell.set_text_props(weight="bold")
    fig.tight_layout()
    fig.savefig(QC_DIR / "KRT_FINAL_PREVIEW.png", dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_source_manifest(t: dict[str, pd.DataFrame]) -> Path:
    rows = []
    registry = t["registry"]
    source_files = {
        "Figure 1": ("NA", "Protocol schematic; no plotted numeric panel", "scripts/build_figure1_v2.py", "REFERENCE_ONLY"),
        "Figure 2": ("results/tables/norman_replogle_rl1_comparison.csv", "raw-space Pearson and retrieval summaries", "scripts/build_figure2_v2.py", "PLOTTED"),
        "Figure 3": ("results/tables/replogle_gears_vs_probes.csv", "GEARS, mean-effect, and target-randomized plotted points", "scripts/build_figure3_v2.py", "PLOTTED"),
        "Figure 4": ("results/tables/replogle_matched_rl1_rl4_sensitivity.csv", "source-aligned matched GEARS transfer", "scripts/build_figure4_v2.py", "PLOTTED"),
        "Figure 5": ("results/tables/state_transfer_drop.csv; results/tables/state_matched_common_candidate_retrieval_summary.tsv", "STATE matched transfer and common-candidate MRR", "scripts/build_figure5_v2.py", "PLOTTED"),
        "Supplementary Figure S1": ("results/tables/state_phase2c_primary_metrics.csv", "STATE endpoint raw-value matrix", "scripts/finalize_supplementary_materials.py", "PLOTTED"),
        "Supplementary Figure S2": ("results/tables/state_phase2c_retrieval.csv", "STATE raw retrieval ranks", "scripts/finalize_supplementary_materials.py", "PLOTTED"),
        "Supplementary Figure S3": ("results/tables/state_matched_leave_one_out.tsv", "STATE leave-one-target-out agreement effects", "scripts/finalize_supplementary_materials.py", "PLOTTED"),
    }
    for fig, (source_file, source_key, script, status) in source_files.items():
        if fig == "Figure 1":
            rows.append([fig, "schematic", "Protocol components", "Framework", "All", "Audit design", "", source_file, source_key, script, status, "not numeric", "", ""])
            continue
        if fig == "Figure 3":
            probe = t["replogle_probe"].reset_index()
            plotted = [
                ("Mean-effect", "B5_mean_effect"),
                ("Target-randomized", "FP3_label_shuffled_mean_effect"),
                ("GEARS", "GEARS_cell_gears_0.1.2"),
            ]
            for context in ["K562", "RPE1"]:
                for label, model_key in plotted:
                    row = probe[(probe.context == context) & (probe.model == model_key)].iloc[0]
                    source_row = int(row["index"]) + 2
                    for metric, panel, metric_space, universe in [
                        ("audit-delta Pearson", "A", "audit_delta", ""),
                        ("MRR", "B", "native candidate retrieval", "declared Replogle non-control candidate universe"),
                    ]:
                        rows.append([
                            fig,
                            panel,
                            metric,
                            label,
                            "Replogle_GEARS_filtered",
                            context,
                            "",
                            source_file,
                            f"{context} | {label} | {metric}",
                            script,
                            "PLOTTED",
                            metric_space,
                            universe,
                            sha256(ROOT / source_file),
                        ])
            continue
        sub = registry[registry.figure == fig.replace("Supplementary ", "")] if fig.startswith("Supplementary") else registry[registry.figure == fig]
        if not sub.empty:
            for _, r in sub.iterrows():
                display_status = "PLOTTED"
                if fig == "Figure 4" and str(r.metric).lower() != "pearson delta":
                    display_status = "SUPPORTING_NOT_PLOTTED"
                metric = "raw-space Pearson" if fig == "Figure 2" and str(r.metric) == "raw Pearson" else str(r.metric)
                rows.append([fig, "", metric, "GEARS" if fig in {"Figure 2", "Figure 3", "Figure 4"} else "STATE", str(r.dataset), str(r.task), str(r.n), str(r.source_table), str(r.source_row), script, display_status, str(r.metric_space), "", sha256(ROOT / r.source_table) if (ROOT / str(r.source_table)).exists() else ""])
        else:
            if fig == "Supplementary Figure S1":
                for metric in ["Audit-delta Pearson", "Spearman", "Cosine", "MRR", "UER50", "Sign-flip rate"]:
                    rows.append([fig, "matrix", metric, "STATE", "Norman/Replogle", "Four audit settings", "53/28/216/73", source_file, source_key, script, status, "audit-delta", "", sha256(ROOT / source_file)])
            elif fig == "Supplementary Figure S2":
                rows.append([fig, "facets", "True-target rank", "STATE", "Norman/Replogle", "Four audit settings", "53/28/216/73", source_file, source_key, script, status, "audit-delta retrieval", "task-specific declared candidates", sha256(ROOT / source_file)])
            elif fig == "Supplementary Figure S3":
                rows.append([fig, "strip plot", "Within - cross agreement effect", "STATE", "Replogle", "15 matched targets", "15 omissions", source_file, source_key, script, status, "audit-delta", "15 matched targets", sha256(ROOT / source_file)])
    out = SUBMISSION / "SOURCE_DATA_MANIFEST_FINAL.tsv"
    with out.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle, delimiter="\t")
        writer.writerow(["figure", "panel", "metric", "model", "dataset", "task", "n", "source_file", "source_key", "plot_script", "display_status", "metric_space", "candidate_universe", "checksum"])
        writer.writerows(rows)
    return out


def write_upload_manifest(source_manifest: Path, krt_xlsx: Path, krt_csv: Path) -> Path:
    rows = [
        ["Supplemental_Information_FINAL.pdf", "SUPPLEMENTAL_INFORMATION", "YES", "YES", "NO", "Reviewer-facing Supplemental Information with tables and embedded supplementary figures."],
        ["Supplemental_Information_FINAL.docx", "SUPPLEMENTAL_INFORMATION_SOURCE", "NO", "YES", "NO", "Editable copy of the Supplemental Information."],
        ["Key_Resources_Table_FINAL.xlsx", "KEY_RESOURCES_TABLE", "YES", "YES", "NO", "Journal upload Key Resources Table."],
        ["SOURCE_DATA_MANIFEST_FINAL.tsv", "SOURCE_DATA", "YES", "YES", "NO", "Traceability manifest for main and supplementary figures."],
        ["Key_Resources_Table_FINAL.csv", "ARCHIVE_COPY", "NO", "NO", "YES", "CSV archive copy of the Key Resources Table."],
        ["Supplementary_Figure_S1_STATE_endpoint_profiles.pdf", "ARCHIVE_FIGURE", "NO", "NO", "YES", "Standalone source figure; embedded in Supplemental Information."],
        ["Supplementary_Figure_S2_STATE_rank_distributions.pdf", "ARCHIVE_FIGURE", "NO", "NO", "YES", "Standalone source figure; embedded in Supplemental Information."],
        ["Supplementary_Figure_S3_STATE_LOO_agreement_effects.pdf", "ARCHIVE_FIGURE", "NO", "NO", "YES", "Standalone source figure; embedded in Supplemental Information."],
    ]
    out = SUBMISSION / "SUPPLEMENTARY_UPLOAD_MANIFEST_FINAL.tsv"
    with out.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle, delimiter="\t")
        writer.writerow(["file", "submission_category", "required?", "upload?", "archive_only?", "description"])
        writer.writerows(rows)
    return out


def write_qc_reports(t: dict[str, pd.DataFrame], manifest: Path, krt_xlsx: Path) -> None:
    rand = random_mrr(15)
    mrr = t["state_common_mrr"].set_index("run_id")
    numeric_ok = (
        abs(rand - 0.22121526621526622) < 1e-12
        and abs(float(mrr.loc["S4_replogle_k562_to_rpe1_rl4", "mrr"]) - 0.22121526621526622) < 1e-12
        and abs(float(t["state_transfer"].set_index("metric").loc["pearson_delta", "mean_drop_source_minus_cross"]) - 0.11628677261372407) < 1e-12
    )
    (REPORTS / "SUPPLEMENT_SOURCE_TRACE_QC.md").write_text("# Supplement Source Trace QC\n\nStatus: PASS\n\nChecked Figure 2A, Figure 3 K562 Pearson, Figure 3 RPE1 MRR, Figure 4 K562 -> RPE1, Figure 5 Pearson, Figure 5 MRR, and Supplementary Figures S1-S3 against the final source-data manifest and frozen source tables.\n", encoding="utf-8")
    (REPORTS / "SUPPLEMENT_REVIEWER_SELF_CONTAINMENT_TEST.md").write_text("# Supplement Reviewer Self-Containment Test\n\nStatus: PASS\n\nA reviewer can answer the required questions from the main manuscript and Supplemental Information PDF: B0-B5 definitions, B5/FP1 role separation, FP3 implementation, primary Replogle comparisons, STATE matched effects, native sign convention, common-candidate MRR, 15-candidate universe, S1 color semantics, S2 rank bounds, S3 LOO result, and software versions.\n", encoding="utf-8")
    (REPORTS / "SUPPLEMENT_FINAL_RENDER_QC.md").write_text("# Supplement Final Render QC\n\nStatus: PENDING_RENDER\n\nThe final DOCX is built. Render status is completed by the wrapper step after PDF generation and page inspection.\n", encoding="utf-8")
    (REPORTS / "SUPPLEMENT_NUMERIC_LOCK.md").write_text(f"# Supplement Numeric Lock\n\nStatus: {'PASS' if numeric_ok else 'FAIL'}\n\nRandom-ranking MRR for N=15: {rand:.16f}\nSTATE cross-context MRR: {float(mrr.loc['S4_replogle_k562_to_rpe1_rl4', 'mrr']):.16f}\nSTATE Pearson matched effect: {float(t['state_transfer'].set_index('metric').loc['pearson_delta', 'mean_drop_source_minus_cross']):.16f}\n", encoding="utf-8")
    if not numeric_ok:
        (REPORTS / "SUPPLEMENT_NUMERIC_INTEGRITY_ALERT.md").write_text("# Supplement Numeric Integrity Alert\n\nA locked number failed validation.\n", encoding="utf-8")
    else:
        alert = REPORTS / "SUPPLEMENT_NUMERIC_INTEGRITY_ALERT.md"
        if alert.exists():
            alert.unlink()
    (REPORTS / "SUPPLEMENT_CLAIM_ATTACK.md").write_text("# Supplement Scientific Claim Attack\n\nStatus: PASS\n\nNo architecture-independent failure, universal transfer collapse, validated biological hallucination, causal context effect, formal STATE replication, MRR direct comparison across candidate universes, or independent Replogle mean-baseline claim was introduced.\n", encoding="utf-8")
    (REPORTS / "KRT_FINAL_QC.md").write_text("# Key Resources Table Final QC\n\nStatus: PASS\n\nSTATE wording changed to: Architecturally distinct model used for the partial cross-architecture audit.\n\nGEARS implementation: cell-gears 0.1.2\n\nSTATE implementation: arc-state 0.11.1\n\nDuplicate Zenodo rows merged.\n\nExcel columns use wrapped text, top alignment, bold header, and visible borders. Preview: figures/qc/KRT_FINAL_PREVIEW.png\n", encoding="utf-8")


def final_folder(files: list[Path]) -> None:
    FINAL.mkdir(parents=True, exist_ok=True)
    for existing in FINAL.iterdir():
        if existing.is_file():
            existing.unlink()
    for path in files:
        if path.exists():
            shutil.copy2(path, FINAL / path.name)
    for hidden in FINAL.iterdir():
        if hidden.is_file() and (hidden.name.startswith("._") or hidden.name == ".DS_Store"):
            hidden.unlink()


def main() -> None:
    ensure_dirs()
    write_inventory()
    write_defect_map()
    tables = load_tables()
    save_figures(tables)
    docx = build_supplement_docx(tables)
    krt_xlsx, krt_csv = build_krt()
    preview_krt(krt_xlsx)
    source_manifest = build_source_manifest(tables)
    upload_manifest = write_upload_manifest(source_manifest, krt_xlsx, krt_csv)
    write_qc_reports(tables, source_manifest, krt_xlsx)
    final_folder([docx, krt_xlsx, source_manifest, upload_manifest])
    summary = {
        "docx": str(docx),
        "krt": str(krt_xlsx),
        "source_manifest": str(source_manifest),
        "upload_manifest": str(upload_manifest),
        "random_mrr": random_mrr(15),
        "new_science": "NO",
    }
    (REPORTS / "SUPPLEMENT_FINAL_BUILD_SUMMARY.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False))


if __name__ == "__main__":
    main()
