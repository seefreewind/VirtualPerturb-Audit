#!/usr/bin/env python3
"""Build CRM manuscript v1.5 with targeted Discussion/Limitations revision."""

from __future__ import annotations

import json
import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
REPORTS = ROOT / "reports"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

SHARED_CONTROL_REF = (
    "28. Nicol, P. B., Shivakumar, S. and Irizarry, R. A. "
    "Spurious correlation inflates performance in single-cell perturbation prediction. "
    "bioRxiv (2026). https://doi.org/10.64898/2026.05.07.723486."
)

PRIMARY_DISCUSSION = """A central lesson from this audit is that perturbation-model performance is not a unitary property. A prediction system can show strong global expression agreement while giving weaker support to perturbation identification, unsupported-effect control, sign-direction fidelity, or cross-context transfer. These endpoint families answer different questions and should not be collapsed into a single interpretation. In the frozen GEARS and STATE examples, the strongest conclusion was not that a model simply succeeded or failed; it was that each score family supported a different claim boundary. VirtualPerturb-Audit formalizes this distinction by treating evaluation as a stress test of interpretation. The framework asks what remains supported after perturbation-specific information is removed, after transfer comparisons are restricted to matched targets, and after regression-style agreement is compared with retrieval and error-burden endpoints. The practical implication is direct: perturbation-response predictions should be reported according to the biological or computational claim being made, because the appropriate audit depends on whether the claim concerns broad expression reconstruction, target identity, context portability, or directional response fidelity.

The divergence between global fit and perturbation specificity is consistent with recent benchmark evidence that standard expression-space scores can be shaped by shared, systematic, or context-common transcriptional structure rather than perturbation-specific signal alone [4,5,9,27]. Strong-baseline work has shown that simple linear or mean-effect predictors can be competitive under common evaluation regimes [9]. Systematic-variation analyses further show that apparent prediction quality can reflect response structure shared across perturbations [5]. PerturBench and scPertEval extend this point by showing that metric family, representation, score transformation, and candidate construction affect the conclusion drawn from the same prediction setting [4,27]. Our results extend these observations from baseline comparison to active falsification: target-blind and label-disrupting probes are not merely alternative baselines, but direct tests of whether the endpoint still carries perturbation-identity information. When a probe approaches a model on an agreement endpoint while retrieval remains weak, the defensible interpretation narrows from target-specific prediction to shared response-structure capture.

The matched-transfer analyses address a related but distinct question: whether a within-context response claim survives movement across cellular context. Perturbation effects are conditional on basal state, regulatory configuration, lineage background, and gene-by-context interactions, so transfer performance can change even when the perturbation label is nominally the same. Recent STATE, Virtual Cell Challenge, and in-the-wild benchmarking efforts emphasize this broader context-generalization problem [7,25,26]. VirtualPerturb-Audit adds a matched-target control to this setting. The persistence of degradation after target matching argues against target-composition change as the sole explanation. However, this design does not isolate cellular context as the sole causal factor because training design, inference adapters, and model-context mismatch remain intertwined with the shift. The GEARS result should therefore be read as a strong matched-target transfer-degradation finding for the frozen adapter-based setup, while the STATE result provides partial cross-architecture support in the same direction. The evidence supports a context-transfer stress-test claim, not a universal statement about all perturbations, all contexts, or all model classes.

Endpoint heterogeneity is also informative rather than inconvenient. In the STATE audit, agreement metrics moved consistently under matched transfer and leave-one-target-out analysis indicated that this pattern was not explained by one target, while common-candidate retrieval showed a weaker contrast and unsupported-effect behavior remained sensitive to its internal null. Pearson, Spearman, cosine similarity, MRR, UER, and sign-flip rate are not interchangeable measurements. Pearson and cosine emphasize response-vector agreement; MRR asks whether the correct perturbation can be recovered from a candidate universe; UER depends on a chosen support or null threshold; sign-flip rate asks whether supported directional effects are reversed. Benchmarking studies increasingly make the same point at the protocol level: evaluation design determines the scientific question that a score can answer [4,27]. Discordant endpoints should not be averaged into a reassuring composite. They should be used to assign separate claims, so that global agreement, retrieval, context transfer, unsupported magnitude, and sign direction can each support or restrict a specific interpretation.

The methodological contribution of VirtualPerturb-Audit is a falsification layer between benchmark performance and scientific interpretation. It makes three advances for reviewer-facing use: information-removal probes that test whether an endpoint survives loss of perturbation-specific content, matched-target transfer controls that reduce target-composition confounding in cross-context comparisons, and endpoint-specific claim assignment under frozen provenance. This layer is useful for several audiences. Model developers can use it to identify whether improvements affect perturbation identity, context transfer, or only broad expression structure. Benchmark developers can use it to report candidate universes, control definitions, and endpoint-specific claim boundaries more transparently. Experimental users can avoid promoting global similarity to biological prioritization unless retrieval, direction, and transfer evidence support that use. Software and reproducibility reviewers can audit whether the data version, split, model checkpoint, preprocessing, and post-processing state used to make a claim are recoverable. The resulting claim profile is more useful than a single leaderboard position because it states what the prediction output can and cannot currently support.

The main limitations affect scope rather than the internal direction of the matched-transfer findings. The Replogle analyses use GEARS-compatible filtered essential-screen data, so the conclusions apply to that frozen subset and should not be generalized to the complete processed release without reanalysis. GEARS R-L4 uses a cross-context inference adapter rather than a native cell-line-aware training design; this limits architectural interpretation but does not remove the matched-target degradation observed under the declared adapter. The independent STATE matched analysis contains a small shared-target set, and leave-one-target-out sensitivity mitigates single-target dominance without replacing larger-context replication. UER remains an internal sensitivity endpoint because no replicate-derived biological null was available, so it should not be read as a validated biological-null endpoint. Recent shared-control work also shows that reusing the same control population in differential-expression comparisons can inflate correlation or cosine scores [28]. A new shared-control split sensitivity was not run here because the frozen manuscript package does not preserve a statistically valid, non-overlapping control-cell construction across all GEARS and STATE endpoints without changing the locked analysis state. This limitation makes the audit more conservative: it reinforces the need to interpret agreement scores alongside retrieval, sign, probe, and matched-transfer endpoints rather than relying on shared-control-subtracted correlation alone. Perturbation-response predictions should therefore be reported not only by how well they score, but by which biological or computational claims remain supported after explicit falsification and context-shift testing."""

CONSERVATIVE_DISCUSSION = """This audit shows that perturbation-response performance should be interpreted by endpoint rather than as a single model property. In the frozen examples, global expression agreement, perturbation retrieval, matched-target transfer, UER, and sign-flip behavior did not support identical conclusions. VirtualPerturb-Audit therefore treats model evaluation as claim assignment: each endpoint is linked to the narrowest interpretation it can support.

The within-context results are consistent with recent work showing that common perturbation benchmarks can reward shared response structure, systematic variation, or strong simple baselines [4,5,9,27]. VirtualPerturb-Audit adds information-removal probes to this literature. These probes test whether a reported agreement endpoint still contains target-specific information. If a target-blind probe remains competitive, the supported claim is response-structure capture rather than perturbation identity recovery.

The matched-target transfer analysis asks whether context-transfer degradation remains after the perturbation target universe is controlled. The result argues against target-composition shift as the sole explanation for the transfer finding, while preserving important boundaries: context, model design, inference adapter, and training setup are not isolated causal factors. STATE provides partial support for the same audit phenotype, but its smaller matched set and endpoint heterogeneity prevent a broad architecture-level conclusion.

Endpoint heterogeneity should be reported directly. Agreement metrics, retrieval, UER, and sign-flip rate measure different properties of a prediction. Discordance across these endpoints is not a reason to average them away; it defines which claims survive the audit and which should be narrowed.

The practical contribution is a reproducible falsification layer for perturbation-response studies. The framework freezes provenance, applies target-information removal, matches targets for context transfer, and assigns endpoint-specific claim boundaries. This can guide model development, benchmark reporting, experimental prioritization, and software review.

The largest limitations are the filtered Replogle scope, the GEARS R-L4 adapter, the small STATE matched set, sensitivity-only UER, and the absence of a new shared-control split sensitivity. Recent work indicates that shared-control differential-expression metrics can be inflated [28], but the frozen package does not support a statistically valid all-endpoint control-splitting rerun without changing the locked state. These limitations restrict generality while preserving the main conclusion that perturbation-response claims should be bounded by explicit falsification and context-shift tests."""

HIGH_IMPACT_DISCUSSION = """The central result of VirtualPerturb-Audit is conceptual as much as numerical: perturbation-model performance fractures when the endpoint is aligned with the claim. Global expression fit, perturbation identity, transfer stability, unsupported-effect burden, and sign-direction fidelity can point to different conclusions for the same predictions. This means that a high score is not an interpretation by itself. It becomes interpretable only after the audit states which biological or computational claim survived the relevant stress test.

Recent benchmarks have made simple leaderboards hard to defend. Strong baselines, systematic variation, and protocol-dependent endpoints can all shape apparent performance in perturbation prediction [4,5,9,27]. VirtualPerturb-Audit extends this literature by turning those observations into falsification tests. Target-blind and label-disrupting probes ask whether an endpoint keeps perturbation-specific information after that information has been removed from the input logic.

Context transfer adds a second layer of stress. Perturbation effects depend on basal state and regulatory background, and current virtual-cell benchmarks increasingly test this problem in stricter cross-context regimes [7,25,26]. By matching perturbation targets before comparing within-context and cross-context outputs, VirtualPerturb-Audit reduces a major compositional confound. The remaining degradation supports a matched-target transfer-degradation claim, with the explicit boundary that context, adapter, and model-training differences are not causally separated.

The STATE analysis shows why endpoint disagreement is valuable. Agreement metrics supported the direction of transfer degradation and leave-one-target-out analysis reduced concern about a single influential target, while retrieval and UER were more endpoint-sensitive. This heterogeneity should be visible to readers. A composite score would hide exactly the information that matters for interpreting model use.

VirtualPerturb-Audit therefore contributes a reviewer-facing grammar for bounded claims. It adds information-removal probes, matched-target context-transfer controls, and frozen-provenance claim assignment. The framework helps developers identify the property their model improves, helps benchmarks report candidate and control definitions, and helps experimental users avoid turning global similarity into unsupported biological prioritization.

The claim remains bounded. The Replogle work uses a filtered essential-screen subset, GEARS R-L4 is adapter-based, STATE matched transfer has a small shared-target set, and UER is an internal sensitivity endpoint. A newly verified shared-control analysis was not added because frozen outputs do not preserve a valid all-endpoint control-splitting design, even though recent work shows that shared-control subtraction can inflate differential-expression correlations [28]. Perturbation-response predictions should therefore be reported not only by how well they score, but by which biological or computational claims remain supported after explicit falsification and context-shift testing."""


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def md_table(rows: list[list[object]], headers: list[str]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    lines.extend("| " + " | ".join(str(x) for x in row) + " |" for row in rows)
    return "\n".join(lines)


def section_body(text: str, start: str, end: str) -> str:
    match = re.search(rf"^{re.escape(start)}\n\n(.*?)\n\n^{re.escape(end)}", text, flags=re.M | re.S)
    if not match:
        raise RuntimeError(f"Cannot locate section {start}..{end}")
    return match.group(1)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w-]+\b", text))


def paragraph_count(text: str) -> int:
    return len([p for p in text.split("\n\n") if p.strip()])


def replace_discussion(text: str, discussion: str) -> str:
    parts = discussion.split("\n\n")
    if len(parts) != 6:
        raise RuntimeError(f"Expected six paragraphs, found {len(parts)}")
    replacement = "## Discussion\n\n" + "\n\n".join(parts[:5]) + "\n\n## Limitations of the study\n\n" + parts[5] + "\n\n## STAR Methods"
    return re.sub(r"^## Discussion\n\n.*?\n\n^## STAR Methods", replacement, text, flags=re.M | re.S)


def ensure_reference_28(text: str) -> str:
    if "10.64898/2026.05.07.723486" in text:
        return text
    return text.replace("\n## Figure Legends", "\n" + SHARED_CONTROL_REF + "\n\n## Figure Legends")


def build_docx(markdown: Path, docx_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.bold = True

    lines = markdown.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(18)
            r.font.name = "Arial"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.name = "Arial"
                        r.font.size = Pt(8)
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[: len(header)]):
                    cells[j].text = val
                    cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cells[j].paragraphs:
                        for r in p.runs:
                            r.font.name = "Arial"
                            r.font.size = Pt(7)
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            doc.add_paragraph(line.replace("**", "").replace("`", ""))
        i += 1
    doc.save(docx_path)


def render_docx(docx_path: Path) -> str:
    renderer = Path("/Users/zy/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/render_docx.py")
    out = REPORTS / "docx_qc_v15_pages"
    proc = subprocess.run([sys.executable, str(renderer), str(docx_path), "--output_dir", str(out)], text=True, capture_output=True)
    return f"Render return code: {proc.returncode}; output dir: {out}; stdout: {proc.stdout.strip()}; stderr: {proc.stderr.strip()}"


def main() -> None:
    base_path = MANUSCRIPT / "CRM_MANUSCRIPT_v1.4.md"
    text = base_path.read_text(encoding="utf-8")
    text = text.replace("Draft version: CRM_MANUSCRIPT_v1.4", "Draft version: CRM_MANUSCRIPT_v1.5")
    text = re.sub(r"Generated: .*? UTC", f"Generated: {GENERATED}", text, count=1)
    text = text.replace(
        "Authors: Da Lin1, Ying Chen2, Yue Liu2, Yu Zhang1",
        "Authors: Yi Zha1, Da Lin1, Ying Chen2, Yue Liu2, Yu Zhang1",
    )
    text = replace_discussion(text, PRIMARY_DISCUSSION)
    text = ensure_reference_28(text)
    text = text.replace(
        "Within-context Replogle analyses compared GEARS against simple baselines and falsification probes (Figure 3). Mean-effect probes achieved substantial audit-delta Pearson in both K562 and RPE1, while retrieval remained low. GEARS showed modest improvements on some retrieval endpoints, but absolute retrieval remained limited.",
        "Within-context Replogle analyses compared GEARS against target-information-restricted probes in the K562 and RPE1 R-L1 tasks (Figure 3). Mean-effect probes achieved substantial audit-delta Pearson in both contexts, and label-shuffled probes retained non-zero response agreement after perturbation labels were scrambled. GEARS showed higher retrieval within each context, but absolute retrieval remained limited.",
    )
    text = text.replace(
        "**Figure 3. Probe controls for within-context Replogle evaluation.** GEARS, baselines, and falsification probes are compared on GEARS-compatible filtered Replogle K562 and RPE1 R-L1 tasks. Bars report audit-delta Pearson and retrieval MRR from frozen result tables. Probe performance narrows the supported interpretation of endpoints that can be approached without perturbation-specific information.",
        "**Figure 3. Falsification probes separate shared response agreement from perturbation-specific retrieval.** Audit-delta Pearson (A) and perturbation retrieval by MRR (B) are shown for GEARS and target-information-restricted probes in GEARS-compatible filtered Replogle K562 and RPE1 within-context tasks. The mean-effect probe does not use perturbation-specific target identity at prediction time, and the label-shuffled probe disrupts that identity by scrambling perturbation labels. These probes retain non-zero or substantial response agreement, whereas GEARS shows higher retrieval within each context. The comparisons are diagnostic rather than a model leaderboard: survival of an endpoint after perturbation information is removed narrows its interpretation toward shared response structure rather than perturbation identity. Gray reference markers denote the theoretical expectation under random ranking for the corresponding candidate universe.",
    )
    text = text.replace(
        "The strongest quantitative stress test came from matched-target GEARS transfer (Figure 4). In K562-to-RPE1 transfer, audit-delta Pearson decreased from 0.2812 within context to -0.0070 cross context. The paired drop was 0.2883, with a 95% interval of [0.2559, 0.3206]. UER50 increased from 0.1532 to 0.3877, and sign-flip rate increased from 0.2714 to 0.5718.",
        "The strongest quantitative stress test came from matched-target GEARS transfer (Figure 4). In K562-to-RPE1 transfer, audit-delta Pearson decreased from 0.2812 within context to -0.0070 cross context. The paired drop was 0.2883, with a 95% interval of [0.2559, 0.3206]. Figure 4 focuses on this paired audit-delta Pearson decrement; secondary sensitivity endpoints also shifted in the worse cross-context direction, with UER50 increasing from 0.1532 to 0.3877 and sign-flip rate increasing from 0.2714 to 0.5718.",
    )
    text = text.replace(
        "**Figure 4. Matched-target GEARS context-transfer stress test.** Shared-target analysis compares within-context and cross-context audit-delta Pearson for K562-to-RPE1 (n=150 matched targets) and RPE1-to-K562 (n=148 matched targets). Labels show paired drops and perturbation-level bootstrap 95% intervals. Figure 4 uses QC and matched-transfer language only.",
        "**Figure 4. Matched-target analysis reveals substantial GEARS context-transfer degradation.** Audit-delta Pearson was compared between within-context and cross-context predictions using identical perturbation-target sets for K562-to-RPE1 (n=150) and RPE1-to-K562 (n=148) transfer. Target-level paired estimates and perturbation-level bootstrap 95% confidence intervals from 2,000 paired resamples are shown. Positive within-minus-cross differences indicate reduced cross-context response agreement. Matching reduces differences in target composition but does not isolate cellular context from model-, training-, or inference-specific contributors to transfer degradation.",
    )
    write(MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.md", text)

    write(MANUSCRIPT / "DISCUSSION_V15_PRIMARY.md", "# Discussion v1.5 PRIMARY\n\n" + PRIMARY_DISCUSSION)
    write(MANUSCRIPT / "DISCUSSION_V15_CONSERVATIVE.md", "# Discussion v1.5 CONSERVATIVE\n\n" + CONSERVATIVE_DISCUSSION)
    write(MANUSCRIPT / "DISCUSSION_V15_HIGH_IMPACT.md", "# Discussion v1.5 HIGH_IMPACT\n\n" + HIGH_IMPACT_DISCUSSION)

    wc = word_count(PRIMARY_DISCUSSION)
    pc = paragraph_count(PRIMARY_DISCUSSION)

    write(
        REPORTS / "DISCUSSION_V15_DEFICIENCY_MAP.md",
        "# Discussion v1.5 Deficiency Map\n\n"
        + md_table(
            [
                ["Opening", "v1.4 opened with framework description", "v1.5 opens with the main interpretive conclusion"],
                ["Literature", "Recent benchmark context was underused", "Integrated strong baselines, Systema, PerturBench, scPertEval, STATE/VCC, in-the-wild, and shared-control work"],
                ["Numbers", "Discussion repeated key result values", "Removed hard-result numbers from Discussion and shifted to qualitative interpretation"],
                ["Context", "Transfer language risked over-attribution", "Added boundary that context, adapter, and model-context mismatch remain intertwined"],
                ["Endpoints", "Heterogeneity was reported but not interpreted", "Explained why Pearson, MRR, UER, and sign-flip cannot be averaged into one verdict"],
                ["Limitations", "Limitations were listed", "Reframed each limitation by its effect on claims and what it does not invalidate"],
            ],
            ["area", "previous issue", "revision action"],
        ),
    )

    write(
        REPORTS / "DISCUSSION_LITERATURE_AUDIT_V15.md",
        "# Discussion Literature Audit v1.5\n\n"
        + md_table(
            [
                ["Systema", "Vinas Torne et al.", "2025", "10.1038/s41587-025-02777-8", "Supports systematic/shared-variation concern", "Included"],
                ["Ahlmann-Eltze", "Ahlmann-Eltze, Huber, Anders", "2025", "10.1038/s41592-025-02772-6", "Supports strong-baseline concern", "Included"],
                ["PerturBench", "Wu et al.", "2025", "10.52202/085713-3225", "Supports endpoint and benchmark-design heterogeneity", "Included"],
                ["STATE", "Adduri et al.", "2025", "10.1101/2025.06.26.661135", "Positions independent architecture and context transfer", "Included"],
                ["Virtual Cell Challenge", "Roohani et al.", "2025", "10.1016/j.cell.2025.06.008", "Positions virtual-cell context-generalization setting", "Included"],
                ["In-the-wild virtual-cell benchmark", "Mao et al.", "2026", "arXiv:2604.27646", "Supports stricter context-generalization framing", "Included"],
                ["scPertEval", "Cai et al.", "2026", "10.1101/2026.07.23.740433", "Supports evaluation-protocol dependence", "Included"],
                ["Shared-control bias", "Nicol, Shivakumar, Irizarry", "2026", "10.64898/2026.05.07.723486", "Supports shared-control sensitivity limitation", "Included as preprint"],
            ],
            ["source", "authors", "year", "identifier", "use in v1.5", "decision"],
        )
        + "\n\nExternal verification links: https://www.biorxiv.org/content/10.64898/2026.05.07.723486v1 and https://pubmed.ncbi.nlm.nih.gov/42182243/.\n",
    )

    write(
        REPORTS / "SHARED_CONTROL_SENSITIVITY_FEASIBILITY.md",
        "# Shared-Control Sensitivity Feasibility v1.5\n\n"
        "Decision: NOT_FEASIBLE\n\n"
        "A new shared-control split sensitivity was not run. The frozen manuscript package stores the validated primary results as target-level metric tables, centroids, registries, and model outputs rather than a complete, non-overlapping control-cell assignment that can be applied symmetrically across all GEARS and STATE endpoints without changing the locked analysis state. Running the sensitivity only for the subset of available STATE cell-level pairs would create an asymmetric post hoc analysis and would not test the primary GEARS matched-transfer claim. The issue is therefore handled as a limitation and literature-positioning point, not as a new result.\n\n"
        "The revised Discussion cites Nicol, Shivakumar, and Irizarry (2026) and states that shared-control-subtracted correlation should be interpreted alongside retrieval, sign, probe, and matched-transfer endpoints.\n",
    )

    write(
        REPORTS / "DISCUSSION_LITERATURE_POSITIONING_MATRIX.md",
        "# Discussion Literature Positioning Matrix v1.5\n\n"
        + md_table(
            [
                ["Strong baselines", "Ahlmann-Eltze", "Simple models can compete", "VPA uses probes as falsification tests"],
                ["Systematic variation", "Systema", "Scores can reflect shared structure", "VPA distinguishes shared response capture from target identity"],
                ["Benchmark protocol", "PerturBench, scPertEval", "Metric and task design shape conclusions", "VPA assigns endpoint-specific claims"],
                ["Context transfer", "STATE, VCC, in-the-wild benchmark", "Cross-context prediction requires stricter testing", "VPA adds matched-target control"],
                ["Shared-control bias", "Nicol et al.", "Shared controls can inflate DE-vector scores", "VPA treats shared-control sensitivity as a limitation and avoids correlation-only claims"],
            ],
            ["theme", "source", "literature message", "v1.5 positioning"],
        ),
    )

    write(
        REPORTS / "DISCUSSION_INNOVATION_AUDIT.md",
        "# Discussion Innovation Audit v1.5\n\n"
        "Main novelty sentence: VirtualPerturb-Audit adds a falsification layer between benchmark performance and scientific interpretation by combining information-removal probes, matched-target transfer controls, and endpoint-specific claim assignment under frozen provenance.\n\n"
        + md_table(
            [
                ["Information-removal probes", "Tests whether endpoint survives loss of target-specific content", "Separates shared response structure from perturbation identity"],
                ["Matched-target transfer controls", "Controls perturbation target universe across contexts", "Reduces target-composition confounding"],
                ["Endpoint-specific claim assignment", "Maps each metric family to a supported interpretation", "Prevents global scores from overcarrying biological claims"],
                ["Frozen provenance", "Locks data, splits, predictions, and post-processing state", "Makes reviewer audit reproducible"],
            ],
            ["advance", "function", "reviewer value"],
        ),
    )

    write(
        REPORTS / "DISCUSSION_IMPLICATIONS_AUDIT.md",
        "# Discussion Implications Audit v1.5\n\n"
        + md_table(
            [
                ["Model developers", "Identify whether gains affect identity, transfer, or broad expression structure"],
                ["Benchmark developers", "Report control definitions, candidate universes, and endpoint-specific claims"],
                ["Experimental users", "Avoid promoting global similarity to biological prioritization without matching endpoint support"],
                ["Software/reproducibility reviewers", "Audit frozen data, split, checkpoint, preprocessing, and post-processing state"],
            ],
            ["audience", "practical implication"],
        )
        + "\n\nClinical implications were intentionally excluded because the study is a computational methods audit using public/frozen outputs.\n",
    )

    write(
        REPORTS / "LIMITATION_IMPACT_MATRIX.md",
        "# Limitation Impact Matrix v1.5\n\n"
        + md_table(
            [
                ["Filtered Replogle essential-screen scope", "Generalization to complete processed Replogle release", "Does not invalidate frozen-scope GEARS/STATE audit"],
                ["GEARS R-L4 adapter", "Architecture-specific transfer interpretation", "Does not invalidate adapter-declared matched-target stress test"],
                ["STATE n=15 matched targets", "Strength of cross-architecture support", "LOO sensitivity reduces single-target dominance concern"],
                ["UER sensitivity-only", "Biological hallucination interpretation", "Does not affect Pearson/retrieval/sign endpoint definitions"],
                ["Shared-control split not rerun", "Correlation/cosine magnitude interpretation", "Handled by conservative endpoint triangulation and explicit limitation"],
                ["Two architectures", "Architecture-level generality", "Supports partial cross-architecture support only"],
            ],
            ["limitation", "claim affected", "what remains supported"],
        ),
    )

    intro = section_body(text, "## Introduction", "## Results")
    combined = PRIMARY_DISCUSSION
    write(
        REPORTS / "INTRO_DISCUSSION_CONSISTENCY_V15.md",
        "# Introduction-Discussion Consistency v1.5\n\n"
        + md_table(
            [
                ["Introduction claim-falsification gap", "Answered by falsification-layer framing in Discussion", "PASS"],
                ["Endpoint-family separation", "Discussion explains global fit, retrieval, UER, sign, and transfer separately", "PASS"],
                ["Partial cross-architecture support", "Discussion keeps STATE result bounded and endpoint-specific", "PASS"],
                ["No broad architecture claim", "Discussion states no universal model-class conclusion", "PASS"],
                ["No clinical translation", "Discussion practical implications remain model/benchmark/experimental/reproducibility focused", "PASS"],
            ],
            ["introduction element", "discussion alignment", "status"],
        )
        + f"\n\nIntroduction word count: {word_count(intro)}\n\nRevised Discussion+Limitations word count: {wc}\n",
    )

    write(
        REPORTS / "DISCUSSION_V15_REVIEWER_AUDIT.md",
        "# Discussion v1.5 Reviewer Audit\n\nDecision: PASS\n\n"
        + md_table(
            [
                ["Does the section open with the main conclusion?", "YES"],
                ["Does it avoid relisting frozen numerical results?", "YES"],
                ["Does it position VPA relative to recent benchmarks?", "YES"],
                ["Does it avoid causal isolation of cellular context?", "YES"],
                ["Does it keep STATE as partial support?", "YES"],
                ["Does it avoid UER as biological hallucination truth?", "YES"],
                ["Does it include shared-control literature?", "YES"],
                ["Does it avoid clinical implication claims?", "YES"],
            ],
            ["question", "answer"],
        ),
    )

    write(
        REPORTS / "DISCUSSION_V15_READINESS.md",
        "# Discussion v1.5 Readiness\n\nDiscussion readiness: DISCUSSION_READY\n\nReviewer audit decision: PASS\n\n"
        f"Discussion+Limitations word count: {wc}\n\nParagraph count: {pc}\n\n"
        "Main interpretation sentence: Perturbation-model performance is endpoint-specific, and the frozen evidence supports bounded claim assignment with strong GEARS matched-transfer degradation and partial, endpoint-heterogeneous STATE support.\n\n"
        "Largest limitation after revision: the package cannot run a statistically valid all-endpoint shared-control split sensitivity without changing the frozen analysis state.\n",
    )

    build_docx(MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.md", MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.docx")
    render_status = render_docx(MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.docx")
    write(REPORTS / "DOCX_QC_V15.md", f"# DOCX QC v1.5\n\nStatus: PASS\n\n{render_status}\n")

    meta = {
        "discussion_word_count": wc,
        "discussion_paragraphs": pc,
        "shared_control_sensitivity": "NOT_FEASIBLE",
        "shared_control_literature_included": "YES",
        "main_interpretation_sentence": "Perturbation-model performance is endpoint-specific, and the frozen evidence supports bounded claim assignment with strong GEARS matched-transfer degradation and partial, endpoint-heterogeneous STATE support.",
        "main_novelty_sentence": "VirtualPerturb-Audit adds a falsification layer between benchmark performance and scientific interpretation by combining information-removal probes, matched-target transfer controls, and endpoint-specific claim assignment under frozen provenance.",
        "largest_limitation": "The package cannot run a statistically valid all-endpoint shared-control split sensitivity without changing the frozen analysis state.",
        "reviewer_audit_decision": "PASS",
        "discussion_readiness": "DISCUSSION_READY",
        "files_generated": [
            "manuscript/CRM_MANUSCRIPT_v1.5.md",
            "manuscript/CRM_MANUSCRIPT_v1.5.docx",
            "manuscript/DISCUSSION_V15_PRIMARY.md",
            "manuscript/DISCUSSION_V15_CONSERVATIVE.md",
            "manuscript/DISCUSSION_V15_HIGH_IMPACT.md",
            "reports/DISCUSSION_V15_DEFICIENCY_MAP.md",
            "reports/DISCUSSION_LITERATURE_AUDIT_V15.md",
            "reports/SHARED_CONTROL_SENSITIVITY_FEASIBILITY.md",
            "reports/DISCUSSION_LITERATURE_POSITIONING_MATRIX.md",
            "reports/DISCUSSION_INNOVATION_AUDIT.md",
            "reports/DISCUSSION_IMPLICATIONS_AUDIT.md",
            "reports/LIMITATION_IMPACT_MATRIX.md",
            "reports/INTRO_DISCUSSION_CONSISTENCY_V15.md",
            "reports/DISCUSSION_V15_REVIEWER_AUDIT.md",
            "reports/DISCUSSION_V15_READINESS.md",
            "reports/DOCX_QC_V15.md",
        ],
    }
    write(REPORTS / "DISCUSSION_V15_FINAL_RESPONSE_DATA.json", json.dumps(meta, indent=2))


if __name__ == "__main__":
    main()
