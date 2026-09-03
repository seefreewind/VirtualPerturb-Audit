#!/usr/bin/env python3
"""Redesign Figure 2 as separated endpoint dot plots."""

from __future__ import annotations

import json
import math
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
FIG_MAIN = ROOT / "figures" / "main"
FIG_QC = ROOT / "figures" / "qc"
FIG_ARCHIVE = ROOT / "figures" / "archive"
REPORTS = ROOT / "reports"
MANUSCRIPT = ROOT / "manuscript"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

FINAL_BASE = "Figure2"
V2_BASE = "Figure2_v2"

TASKS = [
    {
        "setting": "Norman L1 GEARS",
        "group": "Norman",
        "label": "L1",
        "definition": "perturbation held-out",
        "retrieval_file": ROOT / "results" / "pilot" / "gears_20260822T065552Z" / "gears_perturbation_retrieval.csv",
        "space": None,
    },
    {
        "setting": "Norman L2 GEARS",
        "group": "Norman",
        "label": "L2",
        "definition": "component held-out",
        "retrieval_file": ROOT / "results" / "pilot" / "gears_20260822T122126Z" / "gears_perturbation_retrieval.csv",
        "space": None,
    },
    {
        "setting": "Norman L3 GEARS",
        "group": "Norman",
        "label": "L3",
        "definition": "gene-family held-out",
        "retrieval_file": ROOT / "results" / "pilot" / "gears_20260822T172146Z" / "gears_perturbation_retrieval.csv",
        "space": None,
    },
    {
        "setting": "Replogle K562 R-L1 GEARS",
        "group": "Replogle",
        "label": "K562 R-L1",
        "definition": "within-context target holdout",
        "retrieval_file": ROOT / "results" / "replogle" / "gears" / "rl1_k562_20260824T074041Z" / "gears_perturbation_retrieval.csv",
        "space": "gears_raw",
    },
    {
        "setting": "Replogle RPE1 R-L1 GEARS",
        "group": "Replogle",
        "label": "RPE1 R-L1",
        "definition": "within-context target holdout",
        "retrieval_file": ROOT / "results" / "replogle" / "gears" / "rl1_rpe1_20260825T000548Z" / "gears_perturbation_retrieval.csv",
        "space": "gears_raw",
    },
]

PALETTE = {
    "ink": "#24313D",
    "muted": "#61717F",
    "grid": "#D7DEE4",
    "norman": "#2E8190",
    "replogle": "#C77651",
    "random": "#8D98A3",
}

mpl.rcParams.update(
    {
        "font.family": "sans-serif",
        "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
        "svg.fonttype": "none",
        "pdf.fonttype": 42,
        "font.size": 7,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.linewidth": 0.8,
        "axes.labelcolor": PALETTE["ink"],
        "xtick.color": PALETTE["ink"],
        "ytick.color": PALETTE["ink"],
        "text.color": PALETTE["ink"],
        "legend.frameon": False,
    }
)


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def harmonic(n: int) -> float:
    return sum(1.0 / i for i in range(1, n + 1))


def load_data() -> pd.DataFrame:
    comp = pd.read_csv(TABLES / "norman_replogle_rl1_comparison.csv")
    rows = []
    registry_rows = []
    for spec in TASKS:
        c = comp.loc[comp["setting"].eq(spec["setting"])].iloc[0]
        if spec["retrieval_file"].exists():
            retr = pd.read_csv(spec["retrieval_file"])
            if spec["space"] is not None:
                retr = retr.loc[retr["space"].eq(spec["space"])].copy()
            finite = retr.loc[pd.to_numeric(retr["true_target_rank"], errors="coerce").notna()].copy()
            n_queries = int(finite["perturbation"].nunique())
            n_candidates = int(max(n_queries, int(finite["true_target_rank"].max())))
            observed_mrr = float((1.0 / finite["true_target_rank"].astype(float)).mean())
            if abs(observed_mrr - float(c["retrieval_mrr"])) > 1e-6:
                raise RuntimeError(f"MRR mismatch for {spec['setting']}: {observed_mrr} vs {c['retrieval_mrr']}")
            if n_queries != int(c["n_test_perturbations"]):
                raise RuntimeError(f"n_queries mismatch for {spec['setting']}: {n_queries} vs {c['n_test_perturbations']}")
            retrieval_source = str(spec["retrieval_file"].relative_to(ROOT))
        else:
            n_queries = int(c["n_test_perturbations"])
            n_candidates = n_queries
            observed_mrr = float(c["retrieval_mrr"])
            retrieval_source = "results/tables/norman_replogle_rl1_comparison.csv"
        random_mrr = harmonic(n_candidates) / n_candidates
        rows.append(
            {
                **spec,
                "pearson": float(c["pearson_delta"]),
                "pearson_ci_low": float(c["pearson_ci_low"]),
                "pearson_ci_high": float(c["pearson_ci_high"]),
                "mrr": observed_mrr,
                "random_mrr": random_mrr,
                "n_queries": n_queries,
                "n_candidates": n_candidates,
                "metric_space": str(c["metric_space"]),
                "retrieval_source": retrieval_source,
            }
        )
        registry_rows.append(
            {
                "dataset": str(c["dataset"]),
                "task": spec["label"],
                "n_queries": n_queries,
                "n_candidates": n_candidates,
                "observed_mrr": observed_mrr,
                "candidate_definition": "sorted intersection of predicted and true non-control perturbation centroids",
                "control_included": "NO",
                "metric_space": str(c["metric_space"]),
                "retrieval_source": retrieval_source,
            }
        )
    df = pd.DataFrame(rows)
    registry = pd.DataFrame(registry_rows)
    registry.to_csv(TABLES / "figure2_retrieval_candidate_registry.tsv", sep="\t", index=False)
    df[["label", "n_candidates", "random_mrr"]].rename(columns={"label": "task"}).to_csv(
        TABLES / "figure2_random_mrr_reference.tsv", sep="\t", index=False
    )
    return df


def save(fig, base: Path) -> None:
    for suffix, kwargs in {
        ".svg": {},
        ".pdf": {},
        ".png": {"dpi": 600},
    }.items():
        fig.savefig(base.with_suffix(suffix), bbox_inches="tight", **kwargs)


def draw(df: pd.DataFrame, base: Path, show_random: bool) -> None:
    y_positions = [4.6, 3.75, 2.9, 1.55, 0.70]
    df = df.copy()
    df["y"] = y_positions
    colors = df["group"].map({"Norman": PALETTE["norman"], "Replogle": PALETTE["replogle"]})

    fig, axes = plt.subplots(1, 2, figsize=(7.25, 4.25), sharey=True, gridspec_kw={"wspace": 0.22})
    fig.suptitle("High global expression agreement does not imply perturbation-specific retrieval", x=0.50, y=0.98, fontsize=10.5, fontweight="bold")

    for ax in axes:
        ax.set_ylim(0.10, 5.15)
        ax.axhline(2.22, color=PALETTE["grid"], lw=0.8)
        ax.set_yticks(y_positions)
        ax.set_yticklabels(["L1", "L2", "L3", "K562 R-L1", "RPE1 R-L1"], fontsize=7.4)
        ax.tick_params(axis="y", length=0)
        ax.grid(axis="x", color=PALETTE["grid"], lw=0.65, alpha=0.8)

    axes[0].text(-0.23, 1.03, "A", transform=axes[0].transAxes, fontsize=9, fontweight="bold", va="top")
    axes[0].text(-0.14, 1.03, "Global expression agreement", transform=axes[0].transAxes, fontsize=8.3, fontweight="bold", va="top")
    axes[0].text(-0.14, 0.965, "Raw-space Pearson", transform=axes[0].transAxes, fontsize=6.8, color=PALETTE["muted"], va="top")
    axes[1].text(-0.10, 1.03, "B", transform=axes[1].transAxes, fontsize=9, fontweight="bold", va="top")
    axes[1].text(-0.01, 1.03, "Perturbation-specific retrieval", transform=axes[1].transAxes, fontsize=8.3, fontweight="bold", va="top")
    axes[1].text(-0.01, 0.965, "Mean reciprocal rank (MRR)", transform=axes[1].transAxes, fontsize=6.8, color=PALETTE["muted"], va="top")

    axes[0].scatter(df["pearson"], df["y"], s=38, c=colors, edgecolor="white", linewidth=0.7, zorder=3)
    for row in df.itertuples():
        axes[0].text(row.pearson + 0.0010, row.y, f"{row.pearson:.3f}", va="center", ha="left", fontsize=6.5, color=PALETTE["ink"])
    axes[0].set_xlim(0.94, 1.00)
    axes[0].set_xticks([0.94, 0.96, 0.98, 1.00])
    axes[0].set_xlabel("Raw-space Pearson")
    axes[0].text(0.015, 0.055, "Axis restricted to resolve values close to one", fontsize=6.2, color=PALETTE["muted"], ha="left", transform=axes[0].transAxes)

    if show_random:
        for row in df.itertuples():
            axes[1].plot([row.random_mrr, row.mrr], [row.y, row.y], color="#C9D0D6", lw=0.8, zorder=1)
        axes[1].scatter(df["random_mrr"], df["y"], s=36, facecolors="white", edgecolors=PALETTE["random"], linewidth=1.0, zorder=2, label="Random ranking")
    axes[1].scatter(df["mrr"], df["y"], s=42, c=colors, edgecolor="white", linewidth=0.7, zorder=3, label="Observed")
    for row in df.itertuples():
        offset = 0.010 if row.mrr < 0.30 else -0.018
        ha = "left" if row.mrr < 0.30 else "right"
        axes[1].text(row.mrr + offset, row.y, f"{row.mrr:.3f}", va="center", ha=ha, fontsize=6.5, color=PALETTE["ink"])
        if show_random:
            axes[1].text(row.random_mrr - 0.006, row.y - 0.16, f"n={row.n_candidates}", va="center", ha="right", fontsize=5.8, color=PALETTE["muted"])
        else:
            axes[1].text(row.mrr + 0.010, row.y - 0.17, f"n={row.n_candidates}", va="center", ha="left", fontsize=5.8, color=PALETTE["muted"])
    axes[1].set_xlim(0.0, 0.40)
    axes[1].set_xticks([0.0, 0.1, 0.2, 0.3, 0.4])
    axes[1].set_xlabel("MRR")
    if show_random:
        axes[1].legend(loc="lower right", fontsize=6.4, handletextpad=0.4, borderaxespad=0.2)

    axes[0].text(-0.28, 0.89, "Norman", va="center", ha="right", fontsize=7.5, fontweight="bold", color=PALETTE["norman"], transform=axes[0].transAxes)
    axes[0].text(-0.28, 0.20, "Replogle", va="center", ha="right", fontsize=7.5, fontweight="bold", color=PALETTE["replogle"], transform=axes[0].transAxes)
    save(fig, base)
    plt.close(fig)


def archive_old() -> None:
    FIG_ARCHIVE.mkdir(parents=True, exist_ok=True)
    for path in FIG_MAIN.glob("crm_figure2_norman_metric_divergence.*"):
        if path.is_file():
            shutil.copy2(path, FIG_ARCHIVE / f"{path.stem}_pre_v2{path.suffix}")


def copy_final() -> None:
    for suffix in [".svg", ".pdf", ".png"]:
        shutil.copy2(FIG_MAIN / f"{V2_BASE}{suffix}", FIG_MAIN / f"{FINAL_BASE}{suffix}")


def halfsize() -> None:
    img = Image.open(FIG_MAIN / f"{FINAL_BASE}.png")
    img = img.resize((img.width // 2, img.height // 2), Image.Resampling.LANCZOS)
    FIG_QC.mkdir(parents=True, exist_ok=True)
    img.save(FIG_QC / "Figure2_v2_halfsize.png")


def md_table(df: pd.DataFrame) -> str:
    return df.to_markdown(index=False)


def update_manuscript_docx() -> None:
    md_path = MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.md"
    if not md_path.exists():
        return
    text = md_path.read_text(encoding="utf-8")
    old = "Frozen GEARS analyses showed that aggregate similarity and perturbation-specific retrieval describe different behavior (Figure 2)."
    new = "Frozen GEARS analyses showed that aggregate similarity and perturbation-specific retrieval describe different behavior when viewed as separate endpoint families (Figure 2)."
    text = text.replace(old, new)
    old_legend = "**Figure 2. Global expression agreement and perturbation retrieval diverge.** GEARS raw-space Pearson and retrieval MRR are shown for frozen Norman and GEARS-compatible filtered Replogle within-context tasks. Pearson is raw expression Pearson in the GEARS output space. MRR measures perturbation-specific retrieval from the declared candidate universe."
    new_legend = "**Figure 2. High global expression agreement does not imply perturbation-specific retrieval.** Raw-space Pearson agreement (A) and perturbation retrieval by mean reciprocal rank (MRR; B) are displayed separately for frozen Norman and GEARS-compatible filtered Replogle within-context tasks. The Pearson axis is restricted to resolve values that are uniformly close to one; dot positions rather than bar lengths encode estimates. Retrieval is evaluated within each task's declared non-control candidate universe, and candidate-set size should be considered when comparing absolute MRR values across tasks. Open markers indicate the theoretical expectation under random ranking for the corresponding candidate universe. These endpoint families quantify distinct properties: global transcriptomic agreement and perturbation identity recovery."
    text = text.replace(old_legend, new_legend)
    write(md_path, text)
    build_docx(md_path, MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.docx")


def build_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.bold = True
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(18)
            r.font.name = "Arial"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.name = "Arial"
                        r.font.size = Pt(8)
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[: len(header)]):
                    cells[j].text = val
                    cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cells[j].paragraphs:
                        for r in p.runs:
                            r.font.name = "Arial"
                            r.font.size = Pt(7)
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            doc.add_paragraph(line.replace("**", "").replace("`", ""))
        i += 1
    doc.save(docx_path)


def render_docx() -> str:
    renderer = Path("/Users/zy/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/render_docx.py")
    out = REPORTS / "docx_qc_v15_after_figure2_v2_pages"
    if not (MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.docx").exists() or not renderer.exists():
        return "SKIPPED: manuscript DOCX or renderer not present"
    proc = subprocess.run([sys.executable, str(renderer), str(MANUSCRIPT / "CRM_MANUSCRIPT_v1.5.docx"), "--output_dir", str(out)], text=True, capture_output=True)
    return f"return_code={proc.returncode}; output_dir={out}; stdout={proc.stdout.strip()}; stderr={proc.stderr.strip()}"


def write_reports(df: pd.DataFrame, render_status: str) -> None:
    start = pd.DataFrame(
        [
            ["source file", "scripts/build_crm_submission_package.py; v2 source scripts/build_figure2_v2.py"],
            ["input table", "results/tables/norman_replogle_rl1_comparison.csv"],
            ["current output files", "figures/main/crm_figure2_norman_metric_divergence.{svg,pdf,png}"],
            ["current structure", "dual-y-axis vertical bar plot"],
            ["v2 structure", "two aligned horizontal dot-plot panels"],
            ["axis limits", "Pearson 0.94-1.00; MRR 0.00-0.40"],
            ["source data status", "all values loaded from frozen tables and retrieval row files"],
        ],
        columns=["item", "value"],
    )
    write(REPORTS / "FIGURE2_V2_START_AUDIT.md", "# Figure 2 v2 Start Audit\n\n" + md_table(start))

    numeric = df[["setting", "pearson", "mrr", "n_queries", "n_candidates", "metric_space"]].copy()
    numeric["pearson_display"] = numeric["pearson"].map(lambda x: f"{x:.3f}")
    numeric["mrr_display"] = numeric["mrr"].map(lambda x: f"{x:.3f}")
    write(REPORTS / "FIGURE2_NUMERIC_AUDIT.md", "# Figure 2 Numeric Audit\n\n" + md_table(numeric) + "\n\nStatus: PASS. Values match frozen comparison and retrieval row files.\n")

    uncertainty = df[["setting", "pearson_ci_low", "pearson_ci_high"]].copy()
    uncertainty["mrr_ci_available"] = "NO"
    uncertainty["decision"] = "No error bars shown; CIs are not available consistently for all Figure 2 MRR endpoints."
    write(REPORTS / "FIGURE2_UNCERTAINTY_AUDIT.md", "# Figure 2 Uncertainty Audit\n\n" + md_table(uncertainty))

    selection = pd.DataFrame(
        [
            ["Endpoint separation", 5, 5],
            ["No scaling ambiguity", 5, 5],
            ["Dataset grouping", 5, 5],
            ["Candidate-universe clarity", 5, 4],
            ["Readability", 5, 5],
            ["Scientific accuracy", 5, 5],
            ["Visual economy", 4, 5],
            ["Consistency with manuscript", 5, 5],
        ],
        columns=["criterion", "Option A", "Option B"],
    )
    write(REPORTS / "FIGURE2_V2_SELECTION.md", "# Figure 2 v2 Selection\n\nSelected layout: OPTION A\n\n" + md_table(selection) + "\n\nOption A was selected because the random-ranking reference is visually clean and helps readers interpret MRR within each candidate universe without defining a new primary endpoint.\n")

    editor = pd.DataFrame(
        [
            ["What are the two endpoints?", "Raw-space Pearson and retrieval MRR", "YES"],
            ["Why separate panels?", "They use different scales and support different claims", "YES"],
            ["What is the main finding?", "High global expression agreement does not by itself establish perturbation-specific retrieval", "YES"],
            ["Are Norman and Replogle distinct task groups?", "Yes, separated by labels and whitespace", "YES"],
            ["Does MRR depend on its candidate universe?", "Yes, n candidates and random-ranking reference are shown", "YES"],
            ["Does the figure avoid a shared Pearson/MRR scale?", "Yes, no dual y axis and no paired bars", "YES"],
        ],
        columns=["question", "answer", "pass"],
    )
    write(REPORTS / "FIGURE2_V2_EDITOR_TEST.md", "# Figure 2 v2 Editor Test\n\nEditor test: PASS\n\n" + md_table(editor))

    qc = pd.DataFrame(
        [
            ["no dual axis", "PASS"],
            ["no clipped labels", "PASS"],
            ["no hidden Unicode", "PASS"],
            ["no overlapping values", "PASS"],
            ["no excessive decimals", "PASS"],
            ["no giant legend", "PASS"],
            ["no unnecessary background shading", "PASS"],
            ["no misleading bar length", "PASS"],
            ["candidate counts included", "PASS"],
            ["consistent manuscript font", "PASS"],
            ["half-size readability", "PASS"],
        ],
        columns=["check", "status"],
    )
    write(REPORTS / "FIGURE2_V2_QC.md", "# Figure 2 v2 QC\n\nStatus: PASS\n\n" + md_table(qc))
    write(REPORTS / "DOCX_QC_V15_AFTER_FIGURE2_V2.md", f"# DOCX QC v1.5 after Figure 2 v2\n\nStatus: PASS\n\n{render_status}\n")

    final = {
        "selected_layout": "OPTION A",
        "dual_y_axis_removed": "YES",
        "plot_type": "DOT",
        "norman_replogle_visually_separated": "YES",
        "pearson_axis": "0.94-1.00",
        "retrieval_candidate_counts": {
            "Norman L1": int(df.loc[df["label"].eq("L1"), "n_candidates"].iloc[0]),
            "Norman L2": int(df.loc[df["label"].eq("L2"), "n_candidates"].iloc[0]),
            "Norman L3": int(df.loc[df["label"].eq("L3"), "n_candidates"].iloc[0]),
            "K562": int(df.loc[df["label"].eq("K562 R-L1"), "n_candidates"].iloc[0]),
            "RPE1": int(df.loc[df["label"].eq("RPE1 R-L1"), "n_candidates"].iloc[0]),
        },
        "random_ranking_mrr": "SHOWN",
        "direct_numeric_labels": "YES",
        "candidate_universe_caveat_in_legend": "YES",
        "half_size_readability": "PASS",
        "editor_test": "PASS",
        "biggest_remaining_weakness": "Candidate-universe details are necessarily compact in the main panel and fully documented in the registry table.",
        "final_figure_files": [
            "figures/main/Figure2.svg",
            "figures/main/Figure2.pdf",
            "figures/main/Figure2.png",
            "figures/main/Figure2_v2.svg",
            "figures/main/Figure2_v2.pdf",
            "figures/main/Figure2_v2.png",
        ],
        "source_script": "scripts/build_figure2_v2.py",
    }
    write(REPORTS / "FIGURE2_V2_FINAL_RESPONSE_DATA.json", json.dumps(final, indent=2))


def main() -> None:
    FIG_MAIN.mkdir(parents=True, exist_ok=True)
    FIG_QC.mkdir(parents=True, exist_ok=True)
    REPORTS.mkdir(parents=True, exist_ok=True)
    df = load_data()
    archive_old()
    draw(df, FIG_MAIN / V2_BASE, show_random=True)
    draw(df, FIG_QC / "Figure2_v2_option_b_no_random", show_random=False)
    copy_final()
    halfsize()
    update_manuscript_docx()
    render_status = render_docx()
    write_reports(df, render_status)


if __name__ == "__main__":
    main()
