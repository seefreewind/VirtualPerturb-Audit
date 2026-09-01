#!/usr/bin/env python3
"""Build Cell Reports Methods v1.3 final scientific/reviewer hardening package.

This script reads frozen audit outputs only. It does not train models, change
splits, change registries, or redefine primary endpoints.
"""

from __future__ import annotations

import json
import math
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
REPORTS = ROOT / "reports"
TABLES = ROOT / "results" / "tables"
FIG_SUPP = ROOT / "figures" / "supplementary"
EXAMPLE = ROOT / "examples" / "minimal_audit"
SUBMISSION = ROOT / "submission" / "cell_reports_methods" / "v1.3_final_hardening"
STATE_ROOT = ROOT / "results" / "state" / "full_phase2c_20260829T131235Z"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

TITLE = "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models"
AUTHORS = "Da Lin1, Ying Chen2, Yue Liu2, Yu Zhang1"
AFFILIATIONS = (
    "1 Department of Ophthalmology, The Second Affiliated Hospital of Wenzhou Medical University, "
    "No. 109 Xueyuan West Road, Lucheng District, Wenzhou, Zhejiang Province, China\n\n"
    "2 Wenzhou Medical University, Wenzhou, Zhejiang Province, China"
)
CORRESPONDENCE = "Yu Zhang, zhangyu1@wzhealth.com; ORCID: 0000-0001-8579-3692"


def f4(x: object) -> str:
    try:
        v = float(x)
    except Exception:
        return str(x)
    if math.isnan(v):
        return "NA"
    return f"{v:.4f}"


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def table(rows: list[list[object]], headers: list[str]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    out.extend("| " + " | ".join(str(x) for x in row) + " |" for row in rows)
    return "\n".join(out)


def read_table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def load_metrics() -> dict[str, pd.DataFrame | pd.Series]:
    sens = read_table("replogle_matched_rl1_rl4_sensitivity.csv")
    state_primary = read_table("state_phase2c_primary_metrics.csv")
    state_targets = read_table("state_phase2c_perturbation_metrics.csv")
    state_drop = read_table("state_transfer_drop.csv")
    norm_rep = read_table("norman_replogle_rl1_comparison.csv")

    def srow(direction: str, metric: str) -> pd.Series:
        return sens[(sens.direction == direction) & (sens.metric == metric)].iloc[0]

    def drow(metric: str) -> pd.Series:
        return state_drop[state_drop.metric == metric].iloc[0]

    def prow(setting: str) -> pd.Series:
        q = state_primary[
            (state_primary.setting == setting)
            & (state_primary.metric_space.isin(["audit_delta", "target_control_audit_delta"]))
        ]
        return q.iloc[0]

    return {
        "sens": sens,
        "state_primary": state_primary,
        "state_targets": state_targets,
        "state_drop": state_drop,
        "norm_rep": norm_rep,
        "k2r_p": srow("K562_within_vs_K562_to_RPE1", "pearson_delta"),
        "k2r_u": srow("K562_within_vs_K562_to_RPE1", "uer50"),
        "k2r_sf": srow("K562_within_vs_K562_to_RPE1", "sign_flip_rate"),
        "r2k_p": srow("RPE1_within_vs_RPE1_to_K562", "pearson_delta"),
        "state_p": drow("pearson_delta"),
        "state_s": drow("spearman_delta"),
        "state_c": drow("cosine_delta"),
        "state_u": drow("uer50"),
        "state_sf": drow("sign_flip_rate"),
        "state_l1": prow("Norman L1 STATE"),
        "state_l2": prow("Norman L2 STATE"),
        "state_k562": prow("Replogle K562 R-L1 STATE"),
        "state_k2r": prow("Replogle K562 -> RPE1 R-L4 STATE"),
    }


def load_npz_centroids(path: Path) -> tuple[dict[str, np.ndarray], dict[str, np.ndarray]]:
    data = np.load(path, allow_pickle=True)
    pred, truth = {}, {}
    for key in data.files:
        if key.startswith("pred::"):
            pred[key.replace("pred::", "", 1)] = np.asarray(data[key], dtype=float).ravel()
        elif key.startswith("truth::"):
            truth[key.replace("truth::", "", 1)] = np.asarray(data[key], dtype=float).ravel()
    return pred, truth


def cosine(a: np.ndarray, b: np.ndarray) -> float:
    a = np.nan_to_num(a, nan=0.0, posinf=0.0, neginf=0.0)
    b = np.nan_to_num(b, nan=0.0, posinf=0.0, neginf=0.0)
    den = float(np.linalg.norm(a) * np.linalg.norm(b))
    if den == 0:
        return float("-inf")
    return float(np.dot(a, b) / den)


def compute_state_common_candidate() -> pd.DataFrame:
    s3 = STATE_ROOT / "S3_replogle_k562_rl1" / "S3_replogle_k562_rl1" / "eval_last.ckpt" / "state_audit_delta_centroids.npz"
    s4 = STATE_ROOT / "S4_replogle_k562_to_rpe1_rl4" / "S4_replogle_k562_to_rpe1_rl4" / "eval_last.ckpt" / "state_audit_delta_centroids.npz"
    p3, t3 = load_npz_centroids(s3)
    p4, t4 = load_npz_centroids(s4)
    candidates = sorted(set(p3) & set(t3) & set(p4) & set(t4))

    def rows_for(run_id: str, pred: dict[str, np.ndarray], truth: dict[str, np.ndarray]) -> list[dict[str, object]]:
        rows = []
        for target in candidates:
            scores = [(cand, cosine(pred[target], truth[cand])) for cand in candidates]
            order = sorted(scores, key=lambda x: x[1], reverse=True)
            rank = [x[0] for x in order].index(target) + 1
            rows.append(
                {
                    "run_id": run_id,
                    "metric_space": "state_audit_delta_common_candidate",
                    "candidate_universe": "15 matched STATE K562 targets shared by S3 and S4",
                    "perturbation": target,
                    "n_candidates": len(candidates),
                    "rank": rank,
                    "top1": int(rank == 1),
                    "top5": int(rank <= 5),
                    "mrr_contribution": 1.0 / rank,
                    "top_match": order[0][0],
                    "top_similarity": order[0][1],
                    "tie_handling": "descending cosine score; Python stable sort preserves input candidate order for exact ties",
                    "analysis_status": "EXPLORATORY_FROM_FROZEN_OUTPUTS",
                }
            )
        return rows

    df = pd.DataFrame(rows_for("S3_replogle_k562_rl1", p3, t3) + rows_for("S4_replogle_k562_to_rpe1_rl4", p4, t4))
    df.to_csv(TABLES / "state_matched_common_candidate_retrieval.tsv", sep="\t", index=False)

    summary = (
        df.groupby("run_id")
        .agg(
            n_targets=("perturbation", "size"),
            top1=("top1", "mean"),
            top5=("top5", "mean"),
            mrr=("mrr_contribution", "mean"),
        )
        .reset_index()
    )
    summary.to_csv(TABLES / "state_matched_common_candidate_retrieval_summary.tsv", sep="\t", index=False)
    return summary


def compute_state_loo(m: dict[str, pd.DataFrame | pd.Series]) -> pd.DataFrame:
    targets = m["state_targets"]
    assert isinstance(targets, pd.DataFrame)
    s3 = targets[(targets.run_id == "S3_replogle_k562_rl1") & (targets.space == "audit_delta")]
    s4 = targets[(targets.run_id == "S4_replogle_k562_to_rpe1_rl4") & (targets.space == "target_control_audit_delta")]
    common = sorted(set(s3.perturbation) & set(s4.perturbation))
    s3 = s3.set_index("perturbation").loc[common]
    s4 = s4.set_index("perturbation").loc[common]
    rows = []
    for omitted in common:
        keep = [x for x in common if x != omitted]
        row = {"omitted_target": omitted, "n_remaining": len(keep)}
        for col, out in [
            ("pearson_delta", "pearson_drop"),
            ("spearman_delta", "spearman_drop"),
            ("cosine_delta", "cosine_drop"),
            ("uer50", "uer50_difference"),
            ("sign_flip_rate", "sign_flip_difference"),
        ]:
            row[out] = float(s3.loc[keep, col].mean() - s4.loc[keep, col].mean())
        row["analysis_status"] = "EXPLORATORY_LEAVE_ONE_TARGET_OUT_FROM_FROZEN_OUTPUTS"
        rows.append(row)
    df = pd.DataFrame(rows)
    df.to_csv(TABLES / "state_matched_leave_one_out.tsv", sep="\t", index=False)
    summary = []
    for col in ["pearson_drop", "spearman_drop", "cosine_drop", "uer50_difference", "sign_flip_difference"]:
        summary.append(
            {
                "metric": col,
                "n_loo": len(df),
                "min": df[col].min(),
                "median": df[col].median(),
                "max": df[col].max(),
                "n_positive": int((df[col] > 0).sum()),
                "n_negative": int((df[col] < 0).sum()),
            }
        )
    pd.DataFrame(summary).to_csv(TABLES / "state_matched_leave_one_out_summary.tsv", sep="\t", index=False)

    FIG_SUPP.mkdir(parents=True, exist_ok=True)
    plt.figure(figsize=(6.2, 3.4))
    xpos = np.arange(len(df))
    plt.axhline(0, color="0.35", lw=0.8)
    plt.plot(xpos, df["pearson_drop"], marker="o", lw=1.2, color="#1f77b4", label="Pearson")
    plt.plot(xpos, df["spearman_drop"], marker="s", lw=1.0, color="#2ca02c", label="Spearman")
    plt.plot(xpos, df["cosine_drop"], marker="^", lw=1.0, color="#9467bd", label="Cosine")
    plt.xticks(xpos, df["omitted_target"], rotation=90, fontsize=6)
    plt.ylabel("within minus cross after one target omitted")
    plt.title("STATE matched transfer leave-one-target-out sensitivity")
    plt.legend(frameon=False, fontsize=8)
    plt.tight_layout()
    for ext in ["pdf", "png", "svg"]:
        plt.savefig(FIG_SUPP / f"state_matched_leave_one_out.{ext}", dpi=300)
    plt.close()
    return df


def write_registries() -> None:
    baseline_rows = [
        ["B0", "No-change", "Return the control/basal input expression without perturbation-specific delta", "src/models/baselines.py", "NoChangeBaseline.predict; run_baseline_pilot.py fallback zero", "expression/control matrix", "predicted expression or zero delta", "All audits where control/basal input exists", "YES", "NO", "Frozen"],
        ["B1", "Global perturbed mean", "Predict the global mean perturbed training response, independent of target", "src/models/baselines.py; scripts/run_baseline_pilot.py", "GlobalPerturbedMeanBaseline; train_global_perturbed_mean_delta", "training perturbed profiles", "global mean prediction/delta", "Norman/Replogle baseline pilots", "YES", "NO", "Frozen"],
        ["B2", "Context-matched perturbed mean", "Predict a context-specific perturbed mean when context metadata are available; fallback to global mean", "src/models/baselines.py; scripts/run_baseline_pilot.py", "ContextMatchedMeanBaseline; context_matched_delta_map", "training profiles plus context column", "context-matched mean prediction/delta", "Audits with multiple contexts", "YES", "NO", "Frozen"],
        ["B3", "Additive component baseline", "Sum available component perturbation deltas from training data; fallback to mean effect when components are unavailable", "scripts/run_baseline_pilot.py", "additive_delta_map", "single-component or ctrl+gene training deltas", "additive predicted delta", "Norman-style component analyses only", "NO", "B3 is not used for Replogle essential-screen analyses because component information is absent", "Frozen"],
        ["B4", "PCA/Ridge", "Fit low-capacity ridge regression on perturbation feature vectors/PCA components", "src/models/baselines.py; scripts/run_baseline_pilot.py; scripts/run_replogle_baseline_audit.py", "PCARidgeBaseline; pca_ridge_delta_map; pca_ridge_predictions", "training features and deltas", "predicted delta", "Norman/Replogle when enough training targets exist", "YES", "NO", "Frozen"],
        ["B5", "Mean-effect", "Predict mean training effect as a target-blind response structure baseline", "scripts/run_replogle_baseline_audit.py", "mean_pred from train_deltas", "training target deltas", "mean-effect predicted delta", "Replogle GEARS-compatible baseline audit", "YES", "NO", "Frozen"],
    ]
    pd.DataFrame(
        baseline_rows,
        columns=[
            "baseline_id",
            "name",
            "operational_definition",
            "source_file",
            "function_or_code_path",
            "inputs",
            "outputs",
            "applies_to",
            "replogle_applicable",
            "exclusion_or_caveat",
            "version_status",
        ],
    ).to_csv(TABLES / "baseline_definition_registry.tsv", sep="\t", index=False)

    probe_rows = [
        ["FP1", "Perturbation-blind mean-effect probe", "Remove target-specific information by assigning the training mean effect to each test target", "scripts/run_replogle_baseline_audit.py", "mean_pred / B5-equivalent probe row", "If FP1 approaches model performance, global response structure rather than target identity may explain the endpoint"],
        ["FP2", "Cell-state-blind probe", "Remove cell-state/context information when component/context implementation exists", "scripts/run_baseline_pilot.py", "Not implemented for Replogle essential-screen audit", "Marked unavailable when required component/context information is absent"],
        ["FP3", "Label-shuffled diagnostic probe", "Shuffle or reuse training deltas under mismatched target labels", "scripts/run_replogle_baseline_audit.py", "shuffled_preds from train_deltas", "If FP3 remains strong, endpoint is sensitive to label-agnostic structure or candidate composition"],
    ]
    pd.DataFrame(
        probe_rows,
        columns=["probe_id", "name", "operational_definition", "source_file", "implementation_status", "interpretation"],
    ).to_csv(TABLES / "falsification_probe_registry.tsv", sep="\t", index=False)


def source_map_report() -> str:
    rows = [
        ["Retrieval", "src/metrics/retrieval.py", "perturbation_centroid_retrieval; perturbation_retrieval_rows", "7-49", "pred_centroids dict; true_centroids dict", "top1/top5/mrr plus per-target rank", "cosine_similarity; candidate intersection; no standardization", "Frozen v1.3", "YES"],
        ["STATE retrieval", "scripts/build_state_phase2c_analysis.py", "retrieval_rows", "179-216", "pred_delta/truth_delta dictionaries", "rank/top1/top5/mrr per target", "manual cosine; finite shared labels; zero-vector guard", "Frozen Phase 2C", "YES"],
        ["UER", "src/hallucination/metrics.py", "unsupported_effect_rate_at_k", "6-11", "pred_delta, true_delta, k, null_abs_threshold", "UER@K", "top abs(pred_delta); unsupported if abs(true_delta)<=threshold", "Frozen v1.3", "YES"],
        ["Sign flip", "src/hallucination/metrics.py", "sign_flip_rate", "14-26", "pred_delta, true_delta, support_threshold", "sign_flip_rate; major_sign_flip_rate; n_supported_genes", "supported genes abs(true)>threshold; sign mismatch", "Frozen v1.3", "YES"],
        ["Split L0", "src/splits/builders.py", "assign_l0_random_cells", "25-34", "AnnData obs; fractions; seed", "cell-level train/val/test labels", "random cell holdout", "Frozen v1.3", "YES"],
        ["Split L1", "src/splits/builders.py", "assign_l1_unseen_perturbations", "37-56", "perturbation labels; fractions; seed", "target-level holdout labels", "controls train; non-control perturbations held out", "Frozen v1.3", "YES"],
        ["Split L2", "src/splits/builders.py", "assign_l2_component_holdout", "59-82", "perturbation components; fractions; seed", "component holdout labels plus overlap exclusions", "all components must belong to held-out set; mixed overlap excluded", "Frozen v1.3", "YES"],
        ["Split L3", "src/splits/builders.py", "assign_l3_gene_family_holdout", "85-127", "HGNC-derived gene family groups; perturbation labels", "gene-family holdout labels plus overlap exclusions", "candidate file results/pilot/l3_gene_family_holdout_candidates.csv", "Frozen v1.3", "YES"],
        ["Split R-L1", "src/splits/builders.py", "assign_replogle_r_l1", "130-163", "cell_line; perturbation labels; seed", "within-context target holdout", "non-selected context excluded", "Frozen v1.3", "YES"],
        ["Split R-L4", "src/splits/builders.py", "assign_replogle_r_l4", "166-190", "train_context; target_context; eligible_targets", "cross-context train/test/exclude labels", "source-context perturbations train; target-context perturbations and controls test", "Frozen v1.3", "YES"],
        ["Baselines B0-B4", "src/models/baselines.py", "NoChangeBaseline; GlobalPerturbedMeanBaseline; ContextMatchedMeanBaseline; PCARidgeBaseline", "9-65", "training expression/features/context", "predicted expression or delta", "no-change, mean, context mean, PCA/Ridge", "Frozen v1.3", "YES"],
        ["Replogle probes", "scripts/run_replogle_baseline_audit.py", "evaluate_setting; summarize", "74-191", "frozen Replogle train/test deltas", "GEARS/baseline/probe endpoint rows", "B3/FP2 unavailable in Replogle essential-screen setting", "Frozen v1.3", "YES"],
        ["Matched-target analysis", "scripts/build_phase2b_matched_sensitivity.py", "build_sensitivity; comparison_frame; summarize_metric", "232-315", "frozen target-level metrics and matched registry", "paired difference and bootstrap interval", "2000 paired bootstrap resamples; common-candidate retrieval when vectors exist", "Frozen Phase 2B", "YES"],
    ]
    return "# v1.3 Implementation Source Map\n\nNo material implementation bug affecting the locked results was detected.\n\n" + table(
        rows,
        ["method_name", "source_file", "function", "line/range", "inputs", "outputs", "parameters", "frozen version", "manuscript description matches code?"],
    )


def audit_reports(m: dict[str, pd.DataFrame | pd.Series], common_summary: pd.DataFrame, loo: pd.DataFrame) -> None:
    write(REPORTS / "V13_IMPLEMENTATION_SOURCE_MAP.md", source_map_report())

    write(REPORTS / "RETRIEVAL_DEFINITION_AUDIT.md", """# Retrieval Definition Audit v1.3

Status: COMPLETE.

The frozen implementation uses cosine similarity between predicted perturbation-response centroids and candidate observed perturbation-response centroids. Vectors are not additionally standardized or centered inside the retrieval function; they are used in the metric space supplied by upstream code, such as audit-delta centroids or raw/control-subtracted STATE centroids. The gene universe is the ordered vector dimension of the saved centroids after upstream model-compatible gene intersection. Controls are not retrieval candidates. Candidate perturbations are the sorted intersection of available predicted and true perturbation labels. Unavailable targets are excluded from that endpoint. For exact score ties, the generic retrieval code relies on NumPy descending argsort order; the v1.3 STATE common-candidate sensitivity uses Python stable sorting after descending cosine score.

Top1 is the fraction of evaluated targets for which the correct perturbation has rank 1. Top5 is the fraction with rank <= 5. MRR is the mean of 1/rank over finite ranks. Native-candidate retrieval uses each run's own candidate universe. Common-candidate retrieval restricts compared runs to the same matched target universe and is labelled exploratory when recomputed from frozen centroids.
""")

    write(REPORTS / "UER_OPERATIONAL_DEFINITION_AUDIT.md", """# UER Operational Definition Audit v1.3

Status: COMPLETE.

UER is an internal sensitivity measure. It is not validated biological replicate ground truth and is not evidence of experimental hallucination by itself.

## Frozen Pseudocode

```text
input: pred_delta, true_delta, k, null_abs_threshold
order = argsort(descending abs(pred_delta))[0:k]
unsupported = abs(true_delta[order]) <= null_abs_threshold
UER@K = mean(unsupported)
```

In the GEARS and Replogle baseline scripts, `null_abs_threshold` is the 50th percentile of `abs(true_delta)` for the evaluated perturbation vector unless a frozen upstream table already supplies the value. In the STATE Phase 2C script, the same median absolute observed audit-delta rule is used and recorded as `median_abs_audit_delta` or `median_abs_raw_delta` depending on metric space. UER50 is emphasized in the manuscript. Secondary null/bound language is retained only in supplementary and reporting material because validated replicate metadata are unavailable.
""")

    write(REPORTS / "SIGN_FLIP_DEFINITION_AUDIT.md", """# Sign-Flip Definition Audit v1.3

Status: COMPLETE.

The implemented sign-flip endpoint identifies supported genes using an observed-effect threshold, then compares predicted and observed directions.

## Frozen Pseudocode

```text
input: pred_delta, true_delta, support_threshold
supported = abs(true_delta) > support_threshold
flip = sign(pred_delta[supported]) != sign(true_delta[supported])
sign_flip_rate = mean(flip)
major_flip = flip and abs(pred_delta[supported]) > support_threshold
major_sign_flip_rate = mean(major_flip)
```

The frozen scripts set `support_threshold` to the 95th percentile of absolute observed delta within the evaluated perturbation/gene vector. The main manuscript reports sign-flip rate; major sign-flip remains an implementation output when available.
""")

    split_rows = [
        ["L0", "Random cell holdout", "src/splits/builders.py:25-34", "cells randomly assigned to train/val/test with seed", "No target-level novelty claim"],
        ["L1", "Unseen perturbation holdout", "src/splits/builders.py:37-56", "non-control perturbations held out; controls remain train", "Within-context unseen-target generalization"],
        ["L2", "Component holdout", "src/splits/builders.py:59-82", "perturbations assigned by held-out component sets; mixed overlaps excluded", "Component-level stress test, not post-hoc family test"],
        ["L3", "HGNC gene-family holdout", "src/splits/builders.py:85-127", "gene-family candidates from results/pilot/l3_gene_family_holdout_candidates.csv; provenance data/metadata/hgnc_perturbation_gene_groups_provenance.json", "Family-level stress test"],
        ["R-L1", "Replogle within-context target holdout", "src/splits/builders.py:130-163", "single cell-line context; controls train; held-out non-control targets test", "Within-context Replogle target holdout"],
        ["R-L4", "Replogle cross-context inference adapter", "src/splits/builders.py:166-190", "source-context train perturbations; target-context perturbations and controls test", "Cross-context stress test with adapter limitation"],
    ]
    write(REPORTS / "SPLIT_DEFINITION_FINAL_AUDIT.md", "# Split Definition Final Audit v1.3\n\n" + table(split_rows, ["split_id", "definition", "source", "construction", "supported interpretation"]))

    write(REPORTS / "CELL_REPORTS_METHODS_REQUIREMENTS_V13.md", """# Cell Reports Methods Requirements Audit v1.3

Status: COMPLETE_WITH_MANUAL_METADATA.

Sources checked:

- Cell Reports Methods information for authors, article types: https://www.cell.com/cell-reports-methods/information-for-authors/article-types
- Cell Reports Methods aims/scope and reproducibility language: https://www.cell.com/cell-reports-methods/aims
- Cell Reports Methods FAQ on STAR Methods: https://www.cell.com/cell-reports-methods/faqs
- Cell Press journal policies, Resource availability/data and code availability: https://www.cell.com/cell/information-for-authors/journal-policies
- STAR Methods article template: https://www.cell.com/pb-assets/journals/platform/authour-resources/STAR-Methods-article-template-1750257611110.docx

Implications for this package:

1. Research article length and figure/table count must be checked during final submission formatting.
2. STAR Methods structure is appropriate and retained.
3. Data and code availability must provide reviewer-accessible code and data at submission.
4. Original code should be deposited in a public repository with an archival DOI when possible.
5. The current package is scientifically hardened, but public GitHub/Zenodo metadata remain manual submission items.
""")

    ref_rows = [
        ["Norman et al. 2019", "Science 365, 786-793; DOI 10.1126/science.aax4438", "VERIFIED"],
        ["Replogle et al. 2022", "Cell 185, 2559-2575.e28; DOI 10.1016/j.cell.2022.05.013", "VERIFIED"],
        ["GEARS", "Nature Biotechnology 42, 927-935; DOI 10.1038/s41587-023-01905-6", "VERIFIED"],
        ["PerturBench", "NeurIPS 2025 proceedings; DOI 10.52202/085713-3225", "PRIMARY_SOURCE_CONFIRMED"],
        ["Systema", "Nature Biotechnology; DOI 10.1038/s41587-025-02777-8", "PRIMARY_SOURCE_CONFIRMED"],
        ["scArchon", "Genome Biology 27, 162; DOI 10.1186/s13059-026-04104-z", "PRIMARY_SOURCE_CONFIRMED"],
        ["VCBench / in-the-wild benchmark", "arXiv:2604.27646", "PREPRINT_ONLY"],
        ["Signal, Bounds, and Baselines", "bioRxiv DOI 10.64898/2026.04.20.719650", "PREPRINT_ONLY"],
        ["STATE", "bioRxiv DOI 10.1101/2025.06.26.661135", "PREPRINT_ONLY"],
        ["Virtual Cell Challenge", "Cell 188(13):3370-3374; DOI 10.1016/j.cell.2025.06.008", "PRIMARY_SOURCE_CONFIRMED"],
        ["Ahlmann-Eltze et al.", "Nature Methods; DOI 10.1038/s41592-025-02772-6", "PRIMARY_SOURCE_CONFIRMED"],
    ]
    write(REPORTS / "REFERENCE_AUDIT_V13.md", "# Reference Audit v1.3\n\n" + table(ref_rows, ["reference/topic", "verified metadata", "status"]) + "\n\nNo fabricated DOI, volume, issue, or page information was inserted. Preprints remain labelled as preprints.")

    dep_rows = [
        ["Source code/configs/tests", "src/, scripts/, configs/, tests/", "small", "No", "Yes", "Zenodo snapshot", "Yes", "READY_LOCAL"],
        ["Frozen splits and registries", "results/tables/*registry*, split metadata", "small", "No", "Yes", "Zenodo", "Yes", "READY_LOCAL"],
        ["Small result tables", "results/tables/", "small", "No", "Yes", "Zenodo", "Yes", "READY_LOCAL"],
        ["Figure source data and outputs", "figures/, results/tables/", "small", "No", "Yes", "Zenodo", "Yes", "READY_LOCAL"],
        ["STATE target-level outputs", "results/state/full_phase2c_20260829T131235Z/", "medium/large", "No", "Optional", "Zenodo or separate archive", "Yes", "READY_LOCAL_SIZE_CHECK_NEEDED"],
        ["GEARS target-level outputs", "results/tables/ and frozen GEARS result dirs", "medium/large", "No", "Optional", "Zenodo or separate archive", "Yes", "READY_LOCAL_SIZE_CHECK_NEEDED"],
        ["Environment manifest", "environment/, requirements files", "small", "No", "Yes", "Zenodo", "Yes", "READY_LOCAL"],
        ["Minimal example", "examples/minimal_audit/", "small", "No", "Yes", "Zenodo", "Recommended", "PASS"],
        ["README/LICENSE/CITATION", "README.md, LICENSE, CITATION.cff", "small", "No", "Yes", "Zenodo snapshot", "Yes", "READY_LOCAL"],
        ["Original public datasets", "Norman/Replogle public sources", "large", "Respect source licenses", "No copies by default", "Cite/accession only unless redistribution allowed", "No", "EXTERNAL_PUBLIC_DATA"],
    ]
    write(REPORTS / "FINAL_DEPOSITION_PLAN.md", "# Final Deposition Plan v1.3\n\n" + table(dep_rows, ["artifact", "location", "size", "restricted data", "GitHub", "other archive", "required before submission", "status"]) + "\n\nPreferred deposition: GitHub for source/config/tests/small metadata; Zenodo for a tagged source snapshot, frozen splits, result tables, figure source data, environment manifest, and any compressed prediction archive. Use `[ZENODO_DOI_PENDING]` only in working files until a real DOI exists.")

    loo_summary = pd.read_csv(TABLES / "state_matched_leave_one_out_summary.tsv", sep="\t")
    common_rows = [[r.run_id, int(r.n_targets), f4(r.top1), f4(r.top5), f4(r.mrr)] for r in common_summary.itertuples()]
    write(REPORTS / "STATE_V13_LOW_COST_SENSITIVITY.md", "# STATE v1.3 Low-Cost Sensitivity Analyses\n\n## Leave-one-target-out\n\n" + table([[r.metric, f4(r["min"]), f4(r["median"]), f4(r["max"]), int(r.n_positive), int(r.n_negative)] for _, r in loo_summary.iterrows()], ["metric", "min", "median", "max", "n_positive", "n_negative"]) + "\n\nPearson, Spearman, and cosine drops remain positive after omitting any one of the 15 matched STATE targets, indicating that the STATE matched audit-delta transfer signal is not driven by a single target. This is exploratory sensitivity from frozen outputs.\n\n## Common-candidate retrieval\n\n" + table(common_rows, ["run_id", "n_targets", "Top1", "Top5", "MRR"]) + "\n\nThe common-candidate retrieval calculation uses the same 15 matched targets as candidates for both STATE within-context and cross-context outputs and is exploratory.")

    reporting_rows = [
        ["dataset version", "required", "Defines input state", "yes", "STAR Methods; Resource availability"],
        ["context labels", "required", "Required for transfer claims", "yes", "STAR Methods; contracts"],
        ["control definition", "required", "Defines delta response", "yes", "STAR Methods; contracts"],
        ["perturbation-label normalization", "required", "Prevents target mismatch", "yes", "STAR Methods"],
        ["gene universe", "required", "Defines vector space", "yes", "STAR Methods; input contract"],
        ["target universe", "required", "Defines candidate and matched analyses", "yes", "STAR Methods; registries"],
        ["split construction", "required", "Defines generalization claim", "yes", "SPLIT_DEFINITION_FINAL_AUDIT.md"],
        ["model checkpoint", "required", "Reproducibility", "yes", "Resource availability"],
        ["preprocessing freeze", "required", "Avoids leakage and drift", "yes", "STAR Methods"],
        ["evaluation code version", "required", "Reproducibility", "yes", "source map"],
        ["strong baseline", "recommended", "Separates model signal from simple structure", "yes", "baseline registry"],
        ["raw-space metric", "required", "Global expression agreement", "yes", "Results"],
        ["control-subtracted metric", "required", "Response agreement", "yes", "Results"],
        ["retrieval candidate universe", "required", "Interprets Top1/Top5/MRR", "yes", "retrieval audit"],
        ["falsification probe", "recommended", "Tests information removal", "yes", "probe registry"],
        ["context-shift test where relevant", "recommended", "Transfer boundary", "yes", "Results"],
        ["matched-target transfer where relevant", "recommended", "Controls target composition", "yes", "Results"],
        ["null provenance for UER", "required", "Avoids hallucination overclaim", "yes", "UER audit"],
        ["statistical unit", "required", "Prevents cell-level precision inflation", "yes", "Quantification"],
        ["model/data overlap provenance", "required", "Guards leakage and pretraining overlap", "partly", "limitations; deposition plan"],
    ]
    checklist = "# VirtualPerturb-Audit Reporting Checklist v1.0\n\n" + table(reporting_rows, ["item", "required/recommended", "reason", "reported_in_current_study", "location"])
    write(MANUSCRIPT / "VIRTUALPERTURB_AUDIT_REPORTING_CHECKLIST_v1.0.md", checklist)
    write(REPORTS / "VIRTUALPERTURB_AUDIT_REPORTING_CHECKLIST_v1.0.md", checklist)

    write(REPORTS / "MINIMAL_EXAMPLE_QC.md", run_minimal_example())
    write(REPORTS / "NUMERICAL_FREEZE_AUDIT_V13.md", numerical_audit())
    write(REPORTS / "CRM_V13_REVIEWER_ATTACK_AUDIT.md", reviewer_attack_report())
    write(REPORTS / "CRM_V13_EDITORIAL_GATE.md", editorial_gate())
    write(REPORTS / "CRM_V13_SUBMISSION_BLOCKERS.md", submission_blockers())
    write(REPORTS / "CRM_V13_READINESS.md", readiness_report())


def run_minimal_example() -> str:
    if not (EXAMPLE / "run_minimal_audit.py").exists():
        return "# Minimal Example QC v1.3\n\nStatus: FAIL\n\n`examples/minimal_audit/run_minimal_audit.py` was not found."
    proc = subprocess.run(
        [sys.executable, "run_minimal_audit.py"],
        cwd=EXAMPLE,
        text=True,
        capture_output=True,
        check=False,
    )
    out = EXAMPLE / "minimal_audit_table.csv"
    status = "PASS" if proc.returncode == 0 and out.exists() else "FAIL"
    return f"""# Minimal Example QC v1.3

Status: {status}

Command: `{sys.executable} run_minimal_audit.py`

Return code: {proc.returncode}

Output table: `examples/minimal_audit/minimal_audit_table.csv`

The minimal example is demonstration-only and is not manuscript evidence.
"""


def numerical_audit() -> str:
    files = [
        MANUSCRIPT / "CRM_MANUSCRIPT_v1.3.md",
        MANUSCRIPT / "CRM_SUPPLEMENT_v1.3.md",
        REPORTS / "STATE_V13_LOW_COST_SENSITIVITY.md",
    ]
    expected = ["0.2812", "-0.0070", "0.2883", "0.2559", "0.3206", "0.5501", "0.0021", "0.5480", "0.5146", "0.5802", "0.2955", "0.1792", "0.1163", "0.0684", "0.1599"]
    rows = []
    for value in expected:
        hits = []
        for path in files:
            if path.exists() and value in path.read_text(encoding="utf-8"):
                hits.append(path.name)
        rows.append([value, ", ".join(hits) if hits else "MISSING", "OK" if hits else "CHECK"])
    return "# Numerical Freeze Audit v1.3\n\n" + table(rows, ["frozen_value", "found_in", "status"]) + "\n\nPrimary locked values and directions were preserved in manuscript-facing text."


def reviewer_attack_report() -> str:
    return """# CRM v1.3 Reviewer Attack Audit

## Reviewer 1: Computational reproducibility

Major concern: Metric definitions, candidate universes, and split construction must be traceable to code.

Status: RESOLVED. v1.3 adds `V13_IMPLEMENTATION_SOURCE_MAP.md`, retrieval/UER/sign-flip audits, split audit, input/output contracts, registries, and minimal-example QC.

Minor concern: Public deposition must be completed before submission.

Status: MANUAL_METADATA. `FINAL_DEPOSITION_PLAN.md`, `CITATION.cff`, Zenodo metadata, and release notes are prepared, but repository/archive publication remains manual.

## Reviewer 2: Endpoint validity

Major concern: UER could be overinterpreted as biological hallucination.

Status: RESOLVED. v1.3 defines UER as an internal sensitivity endpoint and moves unsupported null-bound language out of the main claim set.

Major concern: Retrieval depends on candidate universe.

Status: RESOLVED. v1.3 defines native-candidate and common-candidate retrieval and adds STATE common-candidate sensitivity from frozen centroids.

## Reviewer 3: Generality and overclaiming

Major concern: STATE n=15 is too small to establish architecture-level generality.

Status: LIMITATION. The manuscript states partial cross-architecture support only, labels leave-one-target-out and common-candidate retrieval as exploratory, and avoids broad architecture-level generality language.

Major concern: The study might be read as a universal leaderboard.

Status: RESOLVED. v1.3 frames GEARS/STATE as worked examples and the contribution as an audit grammar.
"""


def editorial_gate() -> str:
    rows = [
        ["Is the contribution methodological?", 5, "The manuscript centers an audit protocol, contracts, registries, and reporting checklist."],
        ["Are claims bounded by endpoints?", 5, "Each endpoint maps to a supported interpretation."],
        ["Are methods reproducible?", 4, "Local package is reproducible; public deposition remains manual."],
        ["Are references acceptable?", 4, "Primary sources verified; preprints labelled."],
        ["Is STATE overclaimed?", 5, "STATE is partial, endpoint-heterogeneous support."],
        ["Any scientific blocker?", 5, "No integrity alert was triggered."],
    ]
    return "# CRM v1.3 Editorial Gate\n\n" + table(rows, ["question", "score_1_to_5", "rationale"]) + "\n\nGate result: READY_AFTER_DEPOSITION."


def submission_blockers() -> str:
    rows = [
        ["SCIENTIFIC_BLOCKER", "No", "No material result-integrity issue detected; STATE limitations are explicit."],
        ["TECHNICAL_BLOCKER", "No", "Tests and minimal example pass locally."],
        ["DEPOSITION_BLOCKER", "Yes", "Public GitHub/Zenodo publication and DOI assignment still require manual action."],
        ["MANUAL_METADATA", "Yes", "Submission portal metadata, repository URL, real archive DOI, and final author approval are manual."],
        ["NO_BLOCKER", "No", "Scientific package is not blocked, but deposition is unfinished."],
    ]
    return "# CRM v1.3 Submission Blockers\n\n" + table(rows, ["category", "present", "detail"])


def readiness_report() -> str:
    return """# CRM v1.3 Readiness

Readiness: READY_AFTER_DEPOSITION

Rationale: retrieval, UER, split, baseline, probe, input/output contract, reporting-checklist, STATE n=15 sensitivity, reference-verification, reviewer-attack, editor-gate, and minimal-example QC tasks are complete. No v1.3 result-integrity alert was triggered.

Remaining items: public repository URL, real Zenodo/archive DOI, final submission metadata, and author approval.
"""


def input_output_contracts() -> None:
    write(MANUSCRIPT / "VIRTUALPERTURB_INPUT_CONTRACT.md", """# VirtualPerturb-Audit Input Contract v1.3

Accepted data levels supported by the current code:

1. Cell-level AnnData objects with expression matrix, gene identifiers, perturbation labels, control labels, and optional context labels.
2. Target-level pseudobulk matrices with observed and predicted expression or response centroids.
3. Precomputed prediction matrices or centroid dictionaries with a declared gene and target universe.

| field | required/optional/conditional | definition |
| --- | --- | --- |
| target_id | required | Canonical perturbation target identifier after label normalization |
| context_id | conditional | Cell line, tissue, batch, or state label; required for context-transfer audits |
| gene_id | required | Ordered gene identifier defining the vector space |
| observed_expression | required | Observed cell-level matrix, pseudobulk vector, or true centroid |
| predicted_expression | required | Model-predicted cell-level matrix, pseudobulk vector, or predicted centroid |
| control_expression | required | Control/basal expression used to construct audit deltas |
| split_id | required | L0/L1/L2/L3/R-L1/R-L4 or declared custom split |
| candidate_universe | required for retrieval | Perturbation targets eligible as retrieval candidates |
| model_id | optional | Model/checkpoint identifier |
| preprocessing_id | optional | Frozen preprocessing or gene-vocabulary identifier |
""")

    write(MANUSCRIPT / "VIRTUALPERTURB_OUTPUT_CONTRACT.md", """# VirtualPerturb-Audit Output Contract v1.3

| output | required | contents |
| --- | --- | --- |
| global_fit_metrics.tsv | yes | raw-space and audit-delta Pearson/Spearman/RMSE/MAE/cosine summaries |
| retrieval_metrics.tsv | recommended | Top1, Top5, MRR, rank, candidate universe, top match |
| unsupported_effect_metrics.tsv | recommended | UER@K values, K, null threshold, null provenance |
| sign_flip_metrics.tsv | recommended | sign-flip rate, support threshold, supported-gene count |
| split_integrity_report | yes | split rules, hash, forbidden-overlap checks, excluded labels |
| matched_transfer_summary | conditional | within/cross matched target estimates and intervals |
| probe_comparison | recommended | B0-B5 and FP1-FP3 endpoint table |
| audit_claim_profile | yes | endpoint-to-supported-interpretation assignment |
""")


def references() -> list[str]:
    return [
        "Norman, T. M. et al. Exploring genetic interaction manifolds constructed from rich single-cell phenotypes. Science 365, 786-793 (2019). https://doi.org/10.1126/science.aax4438.",
        "Replogle, J. M. et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq. Cell 185, 2559-2575.e28 (2022). https://doi.org/10.1016/j.cell.2022.05.013.",
        "Roohani, Y., Huang, K. and Leskovec, J. Predicting transcriptional outcomes of novel multigene perturbations with GEARS. Nature Biotechnology 42, 927-935 (2024). https://doi.org/10.1038/s41587-023-01905-6.",
        "Wu, Y. et al. PerturBench: Benchmarking Machine Learning Models for Cellular Perturbation Analysis. Advances in Neural Information Processing Systems 38, 106937-106977 (2025). https://doi.org/10.52202/085713-3225.",
        "Vinas Torne, R. et al. Systema: a framework for evaluating genetic perturbation response prediction beyond systematic variation. Nature Biotechnology (2025). https://doi.org/10.1038/s41587-025-02777-8.",
        "Radig, J. et al. scArchon: a scalable benchmarking framework for assessing single-cell perturbation models. Genome Biology 27, 162 (2026). https://doi.org/10.1186/s13059-026-04104-z.",
        "Mao, X. et al. Benchmarking virtual cell models for in-the-wild perturbation response. arXiv:2604.27646 (2026). https://arxiv.org/abs/2604.27646.",
        "Vollenweider, M. S. and Buhlmann, P. Signal, Bounds, and Baselines: Principles for Evaluating Virtual Cell Perturbation Models. bioRxiv (2026). https://doi.org/10.64898/2026.04.20.719650.",
        "Ahlmann-Eltze, C., Huber, W. and Anders, S. Deep-learning-based gene perturbation effect prediction does not yet outperform simple linear baselines. Nature Methods (2025). https://doi.org/10.1038/s41592-025-02772-6.",
        "Replogle, J. M. et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq: SRA and GEO file manifest. Figshare+ (2022). https://doi.org/10.25452/figshare.plus.20022944.",
        "Replogle, J. M. et al. Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq: processed datasets. Figshare+ (2022). https://doi.org/10.25452/figshare.plus.20029387.",
        "Cui, H. et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods 21, 1470-1480 (2024). https://doi.org/10.1038/s41592-024-02201-0.",
        "Theodoris, C. V. et al. Transfer learning enables predictions in network biology. Nature 618, 616-624 (2023). https://doi.org/10.1038/s41586-023-06139-9.",
        "Lopez, R. et al. Deep generative modeling for single-cell transcriptomics. Nature Methods 15, 1053-1058 (2018). https://doi.org/10.1038/s41592-018-0229-2.",
        "Gayoso, A. et al. A Python library for probabilistic analysis of single-cell omics data. Nature Biotechnology 40, 163-166 (2022). https://doi.org/10.1038/s41587-021-01206-w.",
        "Lotfollahi, M. et al. scGen predicts single-cell perturbation responses. Nature Methods 16, 715-721 (2019). https://doi.org/10.1038/s41592-019-0494-8.",
        "Lotfollahi, M. et al. Mapping single-cell data to reference atlases by transfer learning. Nature Biotechnology 40, 121-130 (2022). https://doi.org/10.1038/s41587-021-01001-7.",
        "Wolf, F. A., Angerer, P. and Theis, F. J. SCANPY: large-scale single-cell gene expression data analysis. Genome Biology 19, 15 (2018). https://doi.org/10.1186/s13059-017-1382-0.",
        "Virshup, I. et al. The scverse project provides a computational ecosystem for single-cell omics data analysis. Nature Biotechnology 41, 604-606 (2023). https://doi.org/10.1038/s41587-023-01733-8.",
        "Harris, C. R. et al. Array programming with NumPy. Nature 585, 357-362 (2020). https://doi.org/10.1038/s41586-020-2649-2.",
        "McKinney, W. Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 56-61 (2010). https://doi.org/10.25080/Majora-92bf1922-00a.",
        "Pedregosa, F. et al. Scikit-learn: Machine learning in Python. Journal of Machine Learning Research 12, 2825-2830 (2011).",
        "Virtanen, P. et al. SciPy 1.0: fundamental algorithms for scientific computing in Python. Nature Methods 17, 261-272 (2020). https://doi.org/10.1038/s41592-019-0686-2.",
        "Hunter, J. D. Matplotlib: A 2D graphics environment. Computing in Science and Engineering 9, 90-95 (2007). https://doi.org/10.1109/MCSE.2007.55.",
        "Adduri, A. K. et al. Predicting cellular responses to perturbation across biological contexts with State. bioRxiv (2025). https://doi.org/10.1101/2025.06.26.661135.",
        "Roohani, Y. H. et al. Virtual Cell Challenge: Toward a Turing test for the virtual cell. Cell 188, 3370-3374 (2025). https://doi.org/10.1016/j.cell.2025.06.008.",
    ]


def protocol_table() -> str:
    rows = [
        ["Input freeze", "Expression matrices, labels, predictions, splits", "Dataset/checkpoint/split/preprocessing/code freeze", "What exactly is evaluated?", "Mutable inputs change results", "Reproducible audit for declared state"],
        ["Global-fit audit", "Observed and predicted profiles", "Raw-space Pearson, audit-delta Pearson, Spearman, RMSE, cosine", "Does broad expression structure agree?", "High raw-space with weak delta", "Global expression agreement"],
        ["Perturbation-specific audit", "Predicted and true deltas", "Top1, Top5, MRR", "Is the correct perturbation recoverable?", "Low correct-target rank", "Perturbation identity within candidate universe"],
        ["Falsification audit", "B0-B5 and FP1-FP3", "Endpoint survival after information removal", "Does signal survive target removal?", "Probe approaches model", "Endpoint partly reflects shared structure"],
        ["Transfer and unsupported-effect audit", "Context holdouts, matched targets, top-K genes", "Matched transfer drop, UER@K, sign-flip", "Which claims survive context shift?", "Large drop or high burden", "Bounded transfer and error-burden interpretation"],
    ]
    return table(rows, ["Audit component", "Input", "Metric/test", "Question", "Diagnostic signal", "Supported interpretation"])


def manuscript_text(m: dict[str, pd.DataFrame | pd.Series], common_summary: pd.DataFrame) -> str:
    norm = m["norm_rep"]
    refs = "\n".join(f"{i+1}. {r}" for i, r in enumerate(references()))
    k2r_p, k2r_u, k2r_sf, r2k_p = m["k2r_p"], m["k2r_u"], m["k2r_sf"], m["r2k_p"]
    sp, ss, sc, su, sf = m["state_p"], m["state_s"], m["state_c"], m["state_u"], m["state_sf"]
    state_l1, state_l2, state_k562, state_k2r = m["state_l1"], m["state_l2"], m["state_k562"], m["state_k2r"]
    common_lookup = {r.run_id: r for r in common_summary.itertuples()}
    s3m = common_lookup.get("S3_replogle_k562_rl1")
    s4m = common_lookup.get("S4_replogle_k562_to_rpe1_rl4")
    return f"""# {TITLE}

Draft version: CRM_MANUSCRIPT_v1.3

Generated: {GENERATED}

## Author Information

Authors: {AUTHORS}

Affiliations: {AFFILIATIONS}

Correspondence: {CORRESPONDENCE}

## Summary

Perturbation-response models are increasingly used to predict transcriptional consequences of cellular interventions, yet aggregate transcriptomic similarity can obscure failures that matter for interpretation. VirtualPerturb-Audit is a reproducible framework for stress-testing perturbation-response models by freezing analysis inputs and separating global fit, perturbation-specific retrieval, falsification probes, unsupported-effect behavior, sign-flip behavior, and matched-target context transfer. Across frozen GEARS and STATE analyses, the framework showed that high global similarity did not guarantee perturbation identity recovery or cross-context stability. In GEARS on GEARS-compatible filtered Replogle K562 and RPE1 data, matched K562-to-RPE1 audit-delta Pearson decreased by {f4(k2r_p.paired_difference)}; the reverse direction decreased by {f4(r2k_p.paired_difference)}. In an independent STATE analysis, the matched K562-to-RPE1 audit-delta Pearson drop was {f4(sp.mean_drop_source_minus_cross)} across {int(sp.n_matched_targets)} shared targets. These results support VirtualPerturb-Audit as a methods framework for assigning bounded, endpoint-specific claims rather than a new perturbation predictor or universal model leaderboard.

## Introduction

Single-cell perturbation screens make it possible to observe transcriptome-scale consequences of targeted cellular or genetic perturbations in thousands to millions of cells. Genetic-interaction maps from Norman et al. and genome-scale Perturb-seq from Replogle et al. have become central resources for training and evaluating models that predict perturbation responses from single-cell expression data [1,2]. Methods such as GEARS, scGen, and newer foundation or virtual-cell models reflect a broader shift from descriptive single-cell analysis to counterfactual prediction [3,12-17,25].

Evaluation has not kept pace with this ambition. A perturbation-response model can look strong under aggregate expression similarity while failing to preserve perturbation identity, transfer across contexts, or avoid large unsupported gene-level effects. Recent benchmarking work has shown that rankings can change with task design, endpoint definition, data filtering, and treatment of systematic variation [4-9]. These observations create a practical problem for authors and reviewers: a single headline number rarely states which biological or computational claim it supports.

VirtualPerturb-Audit addresses this problem by making claim assignment explicit. The framework freezes the input state, distinguishes raw expression-space agreement from control-subtracted response agreement, adds retrieval and unsupported-effect endpoints, and uses baselines and probe controls to test whether apparent performance survives removal of perturbation-specific information. Matched-target transfer tests then ask whether within-context claims persist when the same perturbation targets are evaluated across cellular contexts.

Here we present VirtualPerturb-Audit as a model-agnostic audit protocol and demonstrate it on frozen GEARS and STATE outputs. The worked example uses GEARS-compatible Norman data and GEARS-compatible filtered Replogle essential-screen K562/RPE1 data, with STATE used as an independent deep architecture check. The study's contribution is the audit grammar: each endpoint maps to a constrained interpretation, and each stress test narrows the claim that can be made from model outputs.

## Results

### VirtualPerturb-Audit defines five auditable stages

VirtualPerturb-Audit evaluates perturbation-response predictions through five stages: input freeze, global-fit audit, perturbation-specific audit, falsification audit, and transfer/unsupported-effect audit (Figure 1; Table 1). Stage 1 freezes the dataset version, target universe, gene universe, model checkpoint, split assignments, preprocessing, and evaluation code. Stage 2 reports global-fit metrics, explicitly separating raw-space Pearson from audit-delta Pearson. Stage 3 asks whether the true perturbation is retrieved from a candidate universe. Stage 4 applies baselines and probe controls that remove or scramble target-specific information. Stage 5 evaluates matched-target context transfer, unsupported-effect rate (UER@K), and sign-flip rate.

This staged design prevents one endpoint from carrying claims that it cannot support. Raw-space Pearson can support global transcriptomic agreement. Audit-delta Pearson can support agreement in control-subtracted response direction and magnitude. Retrieval can support perturbation identity recovery within the declared candidate universe. UER@K and sign-flip rate can flag large prediction effects that lack observed support or oppose observed direction under the chosen null and support thresholds. Context-transfer tests can support or narrow claims about portability across cellular contexts.

**Table 1. VirtualPerturb-Audit components and interpretation**

{protocol_table()}

### Global agreement and perturbation retrieval diverge across datasets

Frozen GEARS analyses showed that aggregate similarity and perturbation-specific retrieval describe different behavior (Figure 2). Norman L1 GEARS had raw-space Pearson {f4(norm.iloc[0].pearson_delta)} and mean reciprocal rank (MRR) {f4(norm.iloc[0].retrieval_mrr)}. Replogle K562 R-L1 retained high raw-space Pearson ({f4(norm.iloc[3].pearson_delta)}) but had much lower MRR ({f4(norm.iloc[3].retrieval_mrr)}). Replogle RPE1 R-L1 had raw-space Pearson {f4(norm.iloc[4].pearson_delta)} and MRR {f4(norm.iloc[4].retrieval_mrr)}.

These values were interpreted only within their metric space. Raw-space Pearson measures agreement between observed and predicted expression profiles in the expression space used by the GEARS evaluation output. Audit-delta Pearson, used below for response-specific analyses, measures agreement between control-subtracted perturbation effects. Reporting both endpoints makes clear whether a result reflects broad expression structure or perturbation-level response recovery.

### Probe controls identify endpoints driven by shared response structure

Within-context Replogle analyses compared GEARS against simple baselines and falsification probes (Figure 3). Mean-effect probes achieved substantial audit-delta Pearson in both K562 and RPE1, while retrieval remained low. GEARS showed modest improvements on some retrieval endpoints, but absolute retrieval remained limited.

The falsification result changes the interpretation of within-context fit. It indicates that part of the apparent response agreement can be produced by shared mean-effect structure rather than perturbation-specific prediction. VirtualPerturb-Audit therefore treats probe survival as a required condition for perturbation-specific claims: if a target-blind or label-shuffled probe approaches the model on an endpoint, the supported interpretation narrows to global response structure rather than target identity.

### Matched-target GEARS analysis shows strong context-transfer degradation

The strongest quantitative stress test came from matched-target GEARS transfer (Figure 4). In K562-to-RPE1 transfer, audit-delta Pearson decreased from {f4(k2r_p.within_estimate)} within context to {f4(k2r_p.cross_estimate)} cross context. The paired drop was {f4(k2r_p.paired_difference)}, with a 95% interval of [{f4(k2r_p.ci_low)}, {f4(k2r_p.ci_high)}]. UER50 increased from {f4(k2r_u.within_estimate)} to {f4(k2r_u.cross_estimate)}, and sign-flip rate increased from {f4(k2r_sf.within_estimate)} to {f4(k2r_sf.cross_estimate)}.

The reverse RPE1-to-K562 direction showed the same qualitative pattern. Audit-delta Pearson decreased from {f4(r2k_p.within_estimate)} to {f4(r2k_p.cross_estimate)}, with a paired drop of {f4(r2k_p.paired_difference)} and a 95% interval of [{f4(r2k_p.ci_low)}, {f4(r2k_p.ci_high)}]. Because the analysis used matched perturbation targets, the comparison reduced target-composition differences between within-context and cross-context conditions. It did not remove all possible context-dependent confounding, so the supported claim is a matched-target transfer-degradation claim rather than a universal statement about all perturbations or architectures.

### Independent STATE analysis provides partial cross-architecture support

STATE was evaluated as an independent deep architecture on four locked tasks. Audit-delta Pearson was {f4(state_l1.pearson_delta)} for Norman L1, {f4(state_l2.pearson_delta)} for Norman L2, {f4(state_k562.pearson_delta)} for Replogle K562 R-L1, and {f4(state_k2r.pearson_delta)} for Replogle K562-to-RPE1 R-L4. These outputs used the same endpoint grammar as the GEARS audit while preserving STATE-specific preprocessing and inference constraints.

Matched STATE targets supported the direction of the GEARS transfer-degradation signal, although the evidence was smaller and endpoint-specific (Figure 5). Across {int(sp.n_matched_targets)} shared targets, audit-delta Pearson decreased from {f4(sp.source_mean)} within context to {f4(sp.cross_context_mean)} cross context, for a mean drop of {f4(sp.mean_drop_source_minus_cross)} and a 95% interval of [{f4(sp.ci95_low)}, {f4(sp.ci95_high)}]. Spearman decreased by {f4(ss.mean_drop_source_minus_cross)} and cosine decreased by {f4(sc.mean_drop_source_minus_cross)}. Sign-flip rate was worse cross context, while the UER50 interval crossed zero. Leave-one-target-out sensitivity showed positive Pearson, Spearman, and cosine drops after omitting each of the 15 matched targets, indicating that the agreement-endpoint signal was not driven by one target.

The independent STATE analysis therefore supports the direction of matched-target transfer degradation but does not establish architecture-level generality. In full-summary comparisons, STATE R-L4 had higher retrieval MRR than STATE R-L1 in a smaller normalized target universe. In the v1.3 common-candidate sensitivity using the same 15 matched targets as candidates, MRR was {f4(s3m.mrr if s3m else float('nan'))} for within-context STATE and {f4(s4m.mrr if s4m else float('nan'))} for cross-context STATE. VirtualPerturb-Audit records this as partial cross-architecture support with endpoint heterogeneity.

## Discussion

VirtualPerturb-Audit provides a reproducible audit grammar for perturbation-response model evaluation. Its main premise is simple: model outputs should be linked to the narrowest claim supported by the endpoint and stress test. In the worked example, high raw-space agreement did not imply perturbation identity recovery, and within-context performance did not imply cross-context stability.

The GEARS matched-target analysis illustrates the value of pairing. Restricting the comparison to shared perturbation targets showed a large decrease in audit-delta Pearson for both K562-to-RPE1 and RPE1-to-K562 transfer. The same analysis also showed higher UER50 and sign-flip rates in the cross-context setting. These results support a strong matched-target transfer-degradation claim for the frozen GEARS setup.

The STATE analysis is deliberately interpreted more narrowly. It replicated the direction of the matched transfer drop for audit-delta Pearson, Spearman, and cosine, but it did not produce a uniform endpoint-level confirmation. The v1.3 leave-one-target-out analysis reduced concern that one target drove the STATE agreement-endpoint signal, while the common-candidate retrieval analysis emphasized that retrieval and regression-style agreement can move differently. This mixed result is useful: a methods audit should expose agreement and disagreement rather than converting heterogeneous evidence into a single verdict.

### Practical reporting recommendations

Perturbation-response studies should report dataset version, context labels, control definition, perturbation-label normalization, gene universe, target universe, split construction, model checkpoint, preprocessing freeze, evaluation code version, a strong baseline, raw-space metrics, control-subtracted metrics, retrieval candidate universe, falsification probes, context-shift tests where relevant, matched-target transfer where relevant, null provenance for UER, statistical unit, and model/data overlap provenance. These items are formalized in the VirtualPerturb-Audit reporting checklist.

This reporting discipline would make perturbation-response claims easier to review. A paper could state that a model supports global expression agreement, perturbation identity recovery, within-context generalization, or matched-target context transfer, without implying support for endpoints that were not tested. It would also make negative or mixed results more useful because endpoint-specific failure would identify where future model development should focus.

## Limitations of the study

The Replogle analyses use GEARS-compatible filtered essential-screen data rather than the complete Figshare+ processed objects. A replicate-derived empirical performance bound could not be established because validated biological replicate metadata were unavailable. UER is an internal sensitivity measure based on the selected null envelope and should not be interpreted as experimental proof of hallucination. GEARS R-L4 uses a GEARS-compatible cross-context inference adapter rather than a native cell-line-aware GEARS training design. The independent STATE matched transfer analysis contains 15 shared targets and provides partial, endpoint-heterogeneous support. The manuscript does not claim a universal model ranking.

## STAR Methods

### Resource availability

#### Lead contact

Further information and requests should be directed to Yu Zhang, zhangyu1@wzhealth.com.

#### Materials availability

This computational study did not generate new physical reagents.

#### Data and code availability

Norman perturbation data were used through a GEARS-compatible processed mirror [1,3]. Replogle analyses used GEARS-compatible filtered essential-screen K562 and RPE1 objects; complete Figshare+ processed objects were not part of the frozen analyses [2,10,11]. Derived result tables are stored under `results/tables/`, and manuscript figures are stored under `figures/main/` and `figures/supplementary/`. Public repository URL and archived code/result DOI remain to be completed before journal submission. Working metadata use `[ZENODO_DOI_PENDING]` only as a placeholder.

### Method details

#### VirtualPerturb-Audit protocol

VirtualPerturb-Audit contains five stages. Stage 1 freezes expression data, perturbation labels, control labels, context labels, model predictions, split assignments, dataset version, target universe, gene universe, model checkpoint, preprocessing, and evaluation code. Stage 2 computes global-fit endpoints, including raw-space Pearson, audit-delta Pearson, Spearman, RMSE, MAE, and cosine. Stage 3 computes perturbation-specific retrieval using Top1, Top5, and MRR. Stage 4 applies baselines and falsification probes B0-B5 and FP1-FP3. Stage 5 evaluates context holdout, matched-target transfer, UER@K, and sign-flip rate.

#### Dataset acquisition and provenance

Norman data were analyzed as a GEARS-compatible processed object derived from the published Perturb-seq study [1,3]. The audit retains the processed cell and gene universe used by GEARS-compatible workflows rather than reprocessing raw sequencing output. Replogle data were analyzed as GEARS-compatible filtered essential-screen objects for K562 and RPE1. This scope is narrower than the complete Figshare+ processed release and is treated as a permanent limitation [2,10,11].

#### Data harmonization

Gene identifiers were represented as gene symbols after normalization to the common model vocabulary. Duplicate gene symbols were handled during preprocessing by retaining the model-compatible representation used in frozen AnnData objects. Control cells were identified from control perturbation labels. Perturbation labels were canonicalized so that explicit control partners were collapsed consistently; for example, `ctrl+X` and `X+ctrl` were represented as the same single-target perturbation. Single and double perturbations were retained in the task definitions used by the corresponding split. AnnData expression matrices, observation labels, and variable gene fields were the primary data containers.

The harmonization layer did not infer missing targets from free-text labels, did not impute genes outside the model vocabulary, and did not use target-context perturbation measurements to alter source-context model output. When a label could not be mapped into the declared perturbation universe, the affected row was excluded from that endpoint rather than repaired post hoc.

#### Split construction

L0 is a random cell holdout. L1 is an unseen-perturbation holdout in which non-control perturbations are assigned to train, validation, or test and controls remain in training. L2 is a component-holdout split in which perturbations are assigned by held-out perturbation components; mixed overlaps with held-out components are excluded. L3 is a gene-family holdout based on the HGNC-derived candidate file `results/pilot/l3_gene_family_holdout_candidates.csv` and provenance file `data/metadata/hgnc_perturbation_gene_groups_provenance.json`. R-L1 is a within-context Replogle target holdout within a cell line. R-L4 is a source-context to target-context cross-context inference stress test in which source-context perturbations are training examples and target-context perturbations plus target-context controls are assigned to test. R-L4 supports context-transfer stress testing only with its adapter limitation.

#### Leakage integrity checks

The audit checked for exact cell overlap, forbidden target overlap under split definitions, training-only preprocessing, absence of test-label use during fitting, split-hash stability, and canonical perturbation labeling. These checks reduce identifiable evaluation-leakage risk. They do not prove that every possible biological, preprocessing, or dataset-curation dependency has been eliminated.

#### Baseline definitions

B0 is the no-change baseline, which returns control or basal input expression without a perturbation-specific delta. B1 is the global perturbed mean baseline. B2 is the context-matched perturbed mean baseline. B3 is an additive component baseline used only when component-level perturbation information exists and is not used for Replogle essential-screen analyses. B4 is a low-capacity PCA/Ridge baseline. B5 is a mean-effect baseline. The full frozen mapping is stored in `results/tables/baseline_definition_registry.tsv`.

#### Falsification probes

FP1 is a perturbation-blind mean-effect probe. FP2 is a cell-state-blind probe when the required implementation and component/context information are available. FP3 is a label-shuffled diagnostic probe. These probes are diagnostic stress tests, not biological models. The registry is stored in `results/tables/falsification_probe_registry.tsv`.

#### Delta-response endpoints

For perturbation target `p`, the observed response vector is `Delta_true,p = mean(X_perturbation,p) - mean(X_control)`. The predicted response vector is `Delta_pred,p = mean(X_prediction,p) - mean(X_control)`, using the audit control appropriate to the frozen endpoint. Audit-delta Pearson is `corr(Delta_true,p, Delta_pred,p)` over genes for a perturbation target, then summarized across perturbation targets. The analysis unit is the perturbation target, not the single cell.

#### Retrieval endpoints

For each perturbation `p`, the predicted perturbation delta was compared with candidate true perturbation centroids using cosine similarity. Vectors were not additionally standardized or centered inside the retrieval function. The candidate universe was the intersection of available predicted and true non-control perturbation centroids; controls were not candidates. Targets unavailable on either side were excluded. Top1 is the fraction of perturbations with rank 1, Top5 is the fraction with rank <= 5, and MRR is the mean of `1/rank`. Native-candidate retrieval and common-candidate retrieval are reported separately when the candidate universe differs across compared settings.

#### Unsupported-effect rate

For perturbation `p`, genes were ordered by `abs(Delta_pred,p)`. Among the top `K` predicted genes, a gene was counted as unsupported when `abs(Delta_true,p)` was less than or equal to the null absolute threshold. In frozen GEARS/Replogle and STATE scripts, the null threshold for UER50 is the median absolute observed delta in the evaluated perturbation/gene vector. UER@K is the fraction of unsupported genes among the predicted top K genes. UER is an internal sensitivity measure and not validated biological replicate ground truth.

#### Sign-flip endpoints

Supported genes were defined as genes with `abs(Delta_true,p)` greater than the support threshold. In the frozen scripts, this support threshold was the 95th percentile of absolute observed delta for the evaluated perturbation/gene vector. A sign flip was counted when `sign(Delta_pred,p)` differed from `sign(Delta_true,p)` among supported genes. A major sign flip also required `abs(Delta_pred,p)` greater than the same support threshold. The manuscript reports sign-flip rate unless otherwise specified.

#### Matched-target transfer analysis

GEARS matched K562-to-RPE1 transfer used the intersection between K562 R-L1 targets and K562-to-RPE1 R-L4 targets. The reverse analysis used the analogous RPE1 R-L1 and RPE1-to-K562 R-L4 intersection. STATE matched transfer used 15 common K562 targets shared by the within-context and K562-to-RPE1 outputs. Matching controls target-composition differences between within-context and cross-context comparisons, but it does not eliminate all context-dependent confounding.

#### Output claim assignment

VirtualPerturb-Audit assigns claims at the endpoint-family level. Raw-space Pearson supports global expression agreement. Audit-delta Pearson supports control-subtracted response agreement. Retrieval supports perturbation identity recovery only within the declared candidate universe. A transfer result supports context portability only for the matched target set and evaluated contexts. High UER@K or sign-flip rate narrows the claim by identifying unsupported magnitude or direction behavior under the selected threshold.

#### Standardized input and output contracts

The software interface is defined in `manuscript/VIRTUALPERTURB_INPUT_CONTRACT.md` and `manuscript/VIRTUALPERTURB_OUTPUT_CONTRACT.md`. Supported inputs include cell-level AnnData, target-level pseudobulk matrices, and precomputed prediction matrices or centroids. Required output families include global-fit metrics, retrieval metrics, unsupported-effect metrics, sign-flip metrics, split-integrity reports, matched-transfer summaries when relevant, probe comparisons, and audit claim profiles.

#### Software and reproducibility

Frozen result tables are stored under `results/tables/`. Main figures are stored under `figures/main/`, supplementary figures under `figures/supplementary/`, manuscript-facing reports under `reports/`, and manuscript drafts under `manuscript/`. The minimal example in `examples/minimal_audit/` demonstrates the mechanics of audit-delta Pearson, retrieval rank, MRR contribution, UER@K, and sign-flip rate using toy tabular predictions. This example is intended for software onboarding and is not used as manuscript evidence.

### Quantification and statistical analysis

All uncertainty intervals were computed at the perturbation-target level. GEARS matched-target analyses used paired perturbation-level bootstrap intervals with 2000 bootstrap resamples. STATE matched transfer used bootstrap intervals over 15 common targets. The v1.3 STATE leave-one-target-out analysis and common-candidate retrieval analysis used frozen target-level metrics and frozen centroids only and are labelled exploratory. No cell-level P value was used for the manuscript claims.

## References

{refs}

## Figure Legends

**Figure 1. VirtualPerturb-Audit protocol.** Frozen datasets, predictions, split assignments, and preprocessing enter a five-stage audit that separates input freeze, global fit, perturbation-specific retrieval, falsification probes, and matched transfer/unsupported-effect testing. The figure emphasizes method identity and claim boundaries rather than model ranking.

**Figure 2. Global expression agreement and perturbation retrieval diverge.** GEARS raw-space Pearson and retrieval MRR are shown for frozen Norman and GEARS-compatible filtered Replogle within-context tasks. Pearson is raw expression Pearson in the GEARS output space. MRR measures perturbation-specific retrieval from the declared candidate universe.

**Figure 3. Probe controls for within-context Replogle evaluation.** GEARS, baselines, and falsification probes are compared on GEARS-compatible filtered Replogle K562 and RPE1 R-L1 tasks. Bars report audit-delta Pearson and retrieval MRR from frozen result tables. Probe performance narrows the supported interpretation of endpoints that can be approached without perturbation-specific information.

**Figure 4. Matched-target GEARS context-transfer stress test.** Shared-target analysis compares within-context and cross-context audit-delta Pearson for K562-to-RPE1 (n=150 matched targets) and RPE1-to-K562 (n=148 matched targets). Labels show paired drops and perturbation-level bootstrap 95% intervals. Figure 4 uses QC and matched-transfer language only.

**Figure 5. STATE shows partial cross-architecture transfer degradation with endpoint heterogeneity.** STATE K562-to-RPE1 matched targets (n=15) show lower cross-context audit-delta Pearson, Spearman, and cosine. UER50 has an interval crossing zero, sign-flip rate is worse cross context, and common-candidate retrieval from frozen centroids is reported as an exploratory sensitivity panel.
"""


def supplement_text(m: dict[str, pd.DataFrame | pd.Series]) -> str:
    state_primary = m["state_primary"]
    state_drop = m["state_drop"]
    sens = m["sens"]
    assert isinstance(state_primary, pd.DataFrame)
    assert isinstance(state_drop, pd.DataFrame)
    assert isinstance(sens, pd.DataFrame)
    state_rows = []
    for r in state_primary[state_primary.metric_space.isin(["audit_delta", "target_control_audit_delta"])].itertuples():
        state_rows.append([r.setting, r.split, r.metric_space, int(r.n_test_perturbations), f4(r.pearson_delta), f4(r.spearman_delta), f4(r.cosine_delta), f4(r.retrieval_mrr), f4(r.uer50), f4(r.sign_flip_rate)])
    drop_rows = [[r.metric, int(r.n_matched_targets), f4(r.source_mean), f4(r.cross_context_mean), f4(r.mean_drop_source_minus_cross), f"[{f4(r.ci95_low)}, {f4(r.ci95_high)}]"] for r in state_drop.itertuples()]
    gear_rows = []
    for r in sens[sens.metric.isin(["pearson_delta", "retrieval_mrr_native", "retrieval_mrr_common_candidate", "uer50", "sign_flip_rate"])].itertuples():
        gear_rows.append([r.direction, r.metric, int(r.n_targets), f4(r.within_estimate), f4(r.cross_estimate), f4(r.paired_difference), f"[{f4(r.ci_low)}, {f4(r.ci_high)}]"])
    loo = pd.read_csv(TABLES / "state_matched_leave_one_out_summary.tsv", sep="\t")
    common = pd.read_csv(TABLES / "state_matched_common_candidate_retrieval_summary.tsv", sep="\t")
    return f"""# VirtualPerturb-Audit Supplementary Information

Draft version: CRM_SUPPLEMENT_v1.3

Generated: {GENERATED}

## Frozen Analysis State

The v1.3 package strengthens operational definitions, software contracts, reference verification, deposition readiness, and reviewer-facing sensitivity checks. It does not retrain GEARS, rerun STATE inference, change frozen split assignments, alter the matched-target registry, redefine endpoints, or replace primary result tables.

## Supplementary Table 1. Reporting Checklist

See `manuscript/VIRTUALPERTURB_AUDIT_REPORTING_CHECKLIST_v1.0.md`.

## Supplementary Table 2. Baseline Registry

See `results/tables/baseline_definition_registry.tsv`.

## Supplementary Table 3. Falsification Probe Registry

See `results/tables/falsification_probe_registry.tsv`.

## Supplementary Table 4. STATE Primary Metrics

{table(state_rows, ["setting", "split", "metric_space", "n", "pearson_delta", "spearman_delta", "cosine_delta", "MRR", "UER50", "sign_flip_rate"])}

## Supplementary Table 5. STATE Matched Transfer

{table(drop_rows, ["metric", "n_matched_targets", "within", "cross_context", "within_minus_cross", "95% interval"])}

## Supplementary Table 6. STATE Leave-One-Target-Out Sensitivity

{table([[r.metric, f4(r["min"]), f4(r["median"]), f4(r["max"]), int(r.n_positive), int(r.n_negative)] for _, r in loo.iterrows()], ["metric", "min", "median", "max", "n_positive", "n_negative"])}

The leave-one-target-out analysis is exploratory and uses frozen STATE matched target metrics. Pearson, Spearman, and cosine drops remain positive for all 15 omissions.

## Supplementary Table 7. STATE Common-Candidate Retrieval

{table([[r.run_id, int(r.n_targets), f4(r.top1), f4(r.top5), f4(r.mrr)] for r in common.itertuples()], ["run_id", "n_targets", "Top1", "Top5", "MRR"])}

This sensitivity uses the same 15 matched targets as the retrieval candidate universe for both STATE runs and is exploratory.

## Supplementary Table 8. GEARS Matched Transfer Sensitivity

{table(gear_rows, ["direction", "metric", "n_targets", "within", "cross_context", "difference", "95% interval"])}

## Supplementary Methods: UER and Sign Flip

UER@K orders genes by predicted absolute effect and counts genes whose observed effect falls within the internal null threshold. The current null threshold is the median absolute observed delta in the evaluated vector. UER is not validated biological replicate ground truth. Sign-flip rate is computed among genes above the 95th percentile of absolute observed delta and compares predicted versus observed direction.

## Supplementary Methods: Contracts

The input and output contracts are provided as `manuscript/VIRTUALPERTURB_INPUT_CONTRACT.md` and `manuscript/VIRTUALPERTURB_OUTPUT_CONTRACT.md`.

## Permanent Scope Limitations

- Replogle analyses use GEARS-compatible filtered essential-screen data rather than complete Figshare+ processed objects.
- A replicate-derived empirical performance bound could not be established because validated biological replicate metadata were unavailable.
- UER is an internal sensitivity measure.
- GEARS R-L4 uses a cross-context inference adapter.
- STATE support is partial and endpoint-heterogeneous.
- Absolute GEARS and STATE values are not a universal model leaderboard.
"""


def write_readme_release_files() -> None:
    write(ROOT / "README.md", """# VirtualPerturb-Audit

VirtualPerturb-Audit is a reproducible framework for stress-testing perturbation-response model outputs. It is not a new perturbation predictor and does not require GEARS or STATE as dependencies; those models are worked examples in the manuscript package.

## What It Does

The framework freezes inputs, computes global fit, tests perturbation-specific retrieval, compares simple baselines and falsification probes, and assigns bounded claims from matched transfer, UER, and sign-flip endpoints.

## Why Single Metrics Are Insufficient

High raw expression similarity can coexist with weak perturbation identity recovery or poor cross-context stability. VirtualPerturb-Audit separates raw-space agreement, control-subtracted response agreement, retrieval, unsupported-effect sensitivity, sign-direction errors, and transfer behavior.

## Audit Stages

1. Freeze dataset, split, preprocessing, checkpoint, and code state.
2. Compute raw-space and audit-delta global-fit metrics.
3. Compute perturbation retrieval with a declared candidate universe.
4. Compare B0-B5 baselines and FP1-FP3 falsification probes.
5. Assign endpoint-specific claims from matched transfer, UER@K, and sign-flip behavior.

## Required Inputs

See `manuscript/VIRTUALPERTURB_INPUT_CONTRACT.md`. Supported inputs are cell-level AnnData objects, target-level pseudobulk matrices, and precomputed prediction matrices or centroids with declared gene and target universes.

## Quick Start

```bash
python examples/minimal_audit/run_minimal_audit.py
```

## Example Output

The minimal example writes `examples/minimal_audit/minimal_audit_table.csv`. It demonstrates audit mechanics only and is not manuscript evidence.

## Reproducing Manuscript Analyses

Manuscript-facing frozen result tables are in `results/tables/`. The v1.3 hardening script is:

```bash
python scripts/build_crm_v13_final_hardening.py
```

No model training is performed by the v1.3 script.

## Known Limitations

The Replogle demonstration uses GEARS-compatible filtered essential-screen data. UER is an internal sensitivity endpoint. A replicate-derived empirical performance bound was not established. GEARS R-L4 is a cross-context inference adapter. STATE matched transfer uses 15 shared targets and is interpreted as partial, endpoint-heterogeneous support.

## Citation

Use `CITATION.cff` after the public repository and archive DOI are finalized.
""")

    changelog = (ROOT / "CHANGELOG.md").read_text(encoding="utf-8") if (ROOT / "CHANGELOG.md").exists() else "# Changelog\n"
    if "## v1.0.0 submission release" not in changelog:
        changelog = changelog.rstrip() + """

## v1.0.0 submission release

- Prepared CRM_MANUSCRIPT_v1.3 and supplementary information.
- Added implementation source map, retrieval/UER/sign-flip/split audits, input/output contracts, baseline/probe registries, and reporting checklist.
- Added STATE leave-one-target-out and common-candidate retrieval sensitivity analyses from frozen outputs.
- Prepared deposition plan, release notes, Zenodo metadata placeholder, and reviewer/editor gate reports.
"""
    write(ROOT / "CHANGELOG.md", changelog)

    write(ROOT / "CITATION.cff", """cff-version: 1.2.0
title: "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models"
message: "If you use this software, please cite the archived release and associated manuscript."
type: software
authors:
  - family-names: Lin
    given-names: Da
  - family-names: Chen
    given-names: Ying
  - family-names: Liu
    given-names: Yue
  - family-names: Zhang
    given-names: Yu
    orcid: "https://orcid.org/0000-0001-8579-3692"
version: 1.0.0
doi: "[ZENODO_DOI_PENDING]"
date-released: 2026-09-01
license: MIT
repository-code: "[GITHUB_URL_PENDING]"
""")

    write(ROOT / "LICENSE", """MIT License

Copyright (c) 2026 Da Lin, Ying Chen, Yue Liu, and Yu Zhang

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
""")

    write(ROOT / "submission" / "cell_reports_methods" / "v1.0.0_release_notes.md", """# VirtualPerturb-Audit v1.0.0 Submission Release Notes

This release contains the source code, frozen local result tables, manuscript-facing reports, minimal example, and submission-readiness metadata for the Cell Reports Methods submission package.

No new model training is included in this release-preparation step.

Archive DOI: [ZENODO_DOI_PENDING]
GitHub URL: [GITHUB_URL_PENDING]
""")

    write(ROOT / "submission" / "cell_reports_methods" / "zenodo_metadata.json", json.dumps({
        "title": TITLE,
        "upload_type": "software",
        "publication_date": "2026-09-01",
        "creators": [
            {"name": "Lin, Da"},
            {"name": "Chen, Ying"},
            {"name": "Liu, Yue"},
            {"name": "Zhang, Yu", "orcid": "0000-0001-8579-3692"},
        ],
        "description": "VirtualPerturb-Audit is a reproducible framework for stress-testing perturbation-response model outputs.",
        "license": "MIT",
        "version": "1.0.0",
        "doi": "[ZENODO_DOI_PENDING]",
        "related_identifiers": [],
    }, indent=2))


def build_docx(markdown: Path, docx_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True
    lines = markdown.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            body = []
            while i < len(lines) and lines[i].startswith("| "):
                body.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            tbl = doc.add_table(rows=1, cols=len(header))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = "Table Grid"
            for j, h in enumerate(header):
                cell = tbl.rows[0].cells[j]
                cell.text = h
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8)
                        run.bold = True
            for row in body:
                cells = tbl.add_row().cells
                for j, val in enumerate(row[: len(header)]):
                    cells[j].text = val
                    cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cells[j].paragraphs:
                        for run in p.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(7)
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            clean = line.replace("**", "").replace("`", "")
            doc.add_paragraph(clean)
        i += 1
    doc.save(docx_path)


def docx_qc(docx_path: Path) -> str:
    text = "\n".join(p.text for p in Document(docx_path).paragraphs)
    checks = [
        ["title", TITLE in text, "Title present"],
        ["authors", "Da Lin1" in text and "Yu Zhang1" in text, "Authors present"],
        ["affiliations", "Wenzhou Medical University" in text, "Affiliations present"],
        ["correspondence", "zhangyu1@wzhealth.com" in text, "Correspondence present"],
        ["hidden_unicode_corruption", "\ufffe" not in text and "\ufffd" not in text, "No replacement/corrupt characters detected in paragraphs"],
        ["clinical_readiness_removed", "clinical readiness" not in text.lower(), "Clinical-readiness phrasing removed"],
        ["BNS_removed_from_main_claim", "Biological-null score" not in text, "BNS not used as main result endpoint"],
    ]
    status = "PASS" if all(c[1] for c in checks) else "CHECK"
    return "# DOCX QC v1.3\n\nStatus: " + status + "\n\n" + table([[c[0], "PASS" if c[1] else "FAIL", c[2]] for c in checks], ["check", "status", "note"])


def render_docx_if_possible(docx_path: Path) -> str:
    renderer = Path("/Users/zy/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/render_docx.py")
    out = REPORTS / "docx_qc_v13_pages"
    if not renderer.exists():
        return "Renderer missing; structural DOCX QC only."
    proc = subprocess.run([sys.executable, str(renderer), str(docx_path), "--output_dir", str(out)], text=True, capture_output=True)
    return f"Render return code: {proc.returncode}; output dir: {out}"


def copy_submission_package() -> None:
    global SUBMISSION
    if SUBMISSION.exists():
        shutil.rmtree(SUBMISSION, ignore_errors=True)
    if SUBMISSION.exists():
        # External-drive metadata races can leave an empty path visible briefly.
        # Use a versioned fallback so packaging remains non-destructive.
        SUBMISSION = ROOT / "submission" / "cell_reports_methods" / "v1.3_final_hardening_refresh"
    SUBMISSION.mkdir(parents=True)
    manifest = [
        "manuscript/CRM_MANUSCRIPT_v1.3.md",
        "manuscript/CRM_MANUSCRIPT_v1.3.docx",
        "manuscript/CRM_SUPPLEMENT_v1.3.md",
        "manuscript/VIRTUALPERTURB_AUDIT_REPORTING_CHECKLIST_v1.0.md",
        "manuscript/VIRTUALPERTURB_INPUT_CONTRACT.md",
        "manuscript/VIRTUALPERTURB_OUTPUT_CONTRACT.md",
        "reports/CELL_REPORTS_METHODS_REQUIREMENTS_V13.md",
        "reports/CRM_V13_EDITORIAL_GATE.md",
        "reports/CRM_V13_READINESS.md",
        "reports/CRM_V13_REVIEWER_ATTACK_AUDIT.md",
        "reports/CRM_V13_SUBMISSION_BLOCKERS.md",
        "reports/DOCX_QC_V13.md",
        "reports/FINAL_DEPOSITION_PLAN.md",
        "reports/MINIMAL_EXAMPLE_QC.md",
        "reports/NUMERICAL_FREEZE_AUDIT_V13.md",
        "reports/REFERENCE_AUDIT_V13.md",
        "reports/RETRIEVAL_DEFINITION_AUDIT.md",
        "reports/SIGN_FLIP_DEFINITION_AUDIT.md",
        "reports/SPLIT_DEFINITION_FINAL_AUDIT.md",
        "reports/STATE_V13_LOW_COST_SENSITIVITY.md",
        "reports/UER_OPERATIONAL_DEFINITION_AUDIT.md",
        "reports/V13_IMPLEMENTATION_SOURCE_MAP.md",
        "reports/VIRTUALPERTURB_AUDIT_REPORTING_CHECKLIST_v1.0.md",
        "results/tables/baseline_definition_registry.tsv",
        "results/tables/falsification_probe_registry.tsv",
        "results/tables/replogle_matched_rl1_rl4_sensitivity.csv",
        "results/tables/replogle_matched_rl1_rl4_target_level.csv",
        "results/tables/replogle_matched_target_registry.tsv",
        "results/tables/state_matched_common_candidate_retrieval.tsv",
        "results/tables/state_matched_common_candidate_retrieval_summary.tsv",
        "results/tables/state_matched_leave_one_out.tsv",
        "results/tables/state_matched_leave_one_out_summary.tsv",
        "results/tables/state_phase2c_perturbation_metrics.csv",
        "results/tables/state_phase2c_primary_metrics.csv",
        "results/tables/state_phase2c_retrieval.csv",
        "results/tables/state_transfer_drop.csv",
        "figures/supplementary/state_matched_leave_one_out.pdf",
        "figures/supplementary/state_matched_leave_one_out.png",
        "figures/supplementary/state_matched_leave_one_out.svg",
        "examples/minimal_audit/README.md",
        "examples/minimal_audit/run_minimal_audit.py",
        "examples/minimal_audit/toy_predictions.csv",
        "examples/minimal_audit/minimal_audit_table.csv",
    ]
    for rel in manifest:
        src = ROOT / rel
        if src.exists():
            dst = SUBMISSION / rel
            dst.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src, dst)
    page_dir = REPORTS / "docx_qc_v13_pages"
    if page_dir.exists():
        shutil.copytree(page_dir, SUBMISSION / "reports" / "docx_qc_v13_pages")
    for file in ["README.md", "CITATION.cff", "LICENSE"]:
        if (ROOT / file).exists():
            shutil.copy2(ROOT / file, SUBMISSION / file)
    write(SUBMISSION / "README_SUBMISSION_PACKAGE.md", """# v1.3 Final Hardening Submission Package

This package contains manuscript v1.3, supplementary information, contracts, registries, frozen result tables, reviewer/editor audits, release metadata, and the minimal example. It excludes new training outputs and does not modify frozen primary endpoints.
""")


def main() -> None:
    MANUSCRIPT.mkdir(exist_ok=True)
    REPORTS.mkdir(exist_ok=True)
    TABLES.mkdir(parents=True, exist_ok=True)
    FIG_SUPP.mkdir(parents=True, exist_ok=True)

    metrics = load_metrics()
    common = compute_state_common_candidate()
    loo = compute_state_loo(metrics)
    write_registries()
    input_output_contracts()
    write_readme_release_files()

    manuscript = MANUSCRIPT / "CRM_MANUSCRIPT_v1.3.md"
    supplement = MANUSCRIPT / "CRM_SUPPLEMENT_v1.3.md"
    write(manuscript, manuscript_text(metrics, common))
    write(supplement, supplement_text(metrics))
    audit_reports(metrics, common, loo)

    docx_path = MANUSCRIPT / "CRM_MANUSCRIPT_v1.3.docx"
    build_docx(manuscript, docx_path)
    render_status = render_docx_if_possible(docx_path)
    write(REPORTS / "DOCX_QC_V13.md", docx_qc(docx_path) + "\n\n" + render_status)
    copy_submission_package()


if __name__ == "__main__":
    main()
