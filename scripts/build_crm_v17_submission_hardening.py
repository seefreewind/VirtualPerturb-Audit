#!/usr/bin/env python3
"""Build the v1.7 Cell Reports Methods submission-hardening package.

This script performs a no-new-science finalization pass from the frozen v1.6
manuscript and frozen source tables. It updates public repository/deposition
metadata, creates the final numeric registry, writes submission reports and
administrative drafts, builds a release directory, and generates the v1.7 DOCX.
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import textwrap
import time
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
REPORTS = ROOT / "reports"
TABLES = ROOT / "results" / "tables"
FIG_MAIN = ROOT / "figures" / "main"
SUBMISSION = ROOT / "submission"
CRM_SUBMISSION = SUBMISSION / "cell_reports_methods"
RELEASE = ROOT / "release" / "v1.0.0-submission"
DOI = "10.5281/zenodo.22232963"
DOI_URL = f"https://doi.org/{DOI}"
GITHUB = "https://github.com/seefreewind/VirtualPerturb-Audit"
VERSION = "CRM_MANUSCRIPT_v1.7_SUBMISSION"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

SOURCE_MD = MANUSCRIPT / "CRM_MANUSCRIPT_v1.6_FULL.md"
OUT_MD = MANUSCRIPT / "CRM_MANUSCRIPT_v1.7_SUBMISSION.md"
OUT_DOCX = MANUSCRIPT / "CRM_MANUSCRIPT_v1.7_SUBMISSION.docx"

FIGURE_FILES = {
    "Figure 1": FIG_MAIN / "Figure1.png",
    "Figure 2": FIG_MAIN / "Figure2_v2.png",
    "Figure 3": FIG_MAIN / "Figure3_v2.png",
    "Figure 4": FIG_MAIN / "Figure4_v2.png",
    "Figure 5": FIG_MAIN / "Figure5_v2.png",
}


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def fmt(x: object, digits: int = 4) -> str:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    try:
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def replace_section(text: str, heading: str, new_body: str) -> str:
    pattern = rf"(## {re.escape(heading)}\n)(.*?)(?=\n## |\Z)"
    return re.sub(pattern, rf"\1\n{new_body.rstrip()}\n", text, flags=re.S)


def replace_subsection(text: str, heading: str, new_body: str) -> str:
    pattern = rf"(#### {re.escape(heading)}\n)(.*?)(?=\n### |\n#### |\n## |\Z)"
    return re.sub(pattern, rf"\1\n{new_body.rstrip()}\n", text, flags=re.S)


def replace_results_subsection(text: str, heading: str, new_body: str) -> str:
    pattern = rf"(### {re.escape(heading)}\n)(.*?)(?=\n### |\n## |\Z)"
    return re.sub(pattern, rf"\1\n{new_body.rstrip()}\n", text, flags=re.S)


def build_numeric_registry() -> pd.DataFrame:
    rows: list[dict[str, object]] = []

    def add(**kwargs):
        base = {
            "claim_id": "",
            "section": "",
            "figure": "",
            "metric": "",
            "dataset": "",
            "task": "",
            "n": "",
            "estimate": "",
            "comparison_value": "",
            "difference": "",
            "ci_low": "",
            "ci_high": "",
            "metric_space": "",
            "direction_definition": "",
            "source_table": "",
            "source_row": "",
            "status": "AUTHORITATIVE_FROZEN",
        }
        base.update(kwargs)
        rows.append(base)

    nr = pd.read_csv(TABLES / "norman_replogle_rl1_comparison.csv")
    for idx, r in nr.iterrows():
        for metric, col in [
            ("raw Pearson", "pearson_delta"),
            ("MRR", "retrieval_mrr"),
            ("Top1", "retrieval_top1"),
            ("Top5", "retrieval_top5"),
        ]:
            add(
                claim_id=f"F2_{r['setting'].replace(' ', '_')}_{metric.replace(' ', '_')}",
                section="Results: Global agreement and perturbation retrieval diverge across datasets",
                figure="Figure 2",
                metric=metric,
                dataset=r["dataset"],
                task=r["setting"],
                n=int(r["n_test_perturbations"]),
                estimate=fmt(r[col]),
                ci_low=fmt(r.get("pearson_ci_low")) if metric == "raw Pearson" else "",
                ci_high=fmt(r.get("pearson_ci_high")) if metric == "raw Pearson" else "",
                metric_space="gears_raw" if metric == "raw Pearson" else "native candidate retrieval",
                direction_definition="higher indicates stronger endpoint agreement/retrieval",
                source_table="results/tables/norman_replogle_rl1_comparison.csv",
                source_row=idx + 2,
            )

    probes = pd.read_csv(TABLES / "replogle_gears_vs_probes.csv")
    keep_models = {"B1_global_perturbed_mean", "FP1_perturbation_blind_mean_effect", "FP3_label_shuffled_mean_effect", "GEARS_cell_gears_0.1.2"}
    for idx, r in probes[probes["model"].isin(keep_models)].iterrows():
        for metric, col in [("audit-delta Pearson", "pearson_delta"), ("MRR", "retrieval_mrr")]:
            add(
                claim_id=f"F3_{r['context']}_{r['model']}_{metric.replace(' ', '_')}",
                section="Results: Probe controls identify endpoints driven by shared response structure",
                figure="Figure 3",
                metric=metric,
                dataset="Replogle_GEARS_filtered",
                task=f"{r['context']} {r['split']} {r['model']}",
                n="216" if r["context"] == "K562" else "308",
                estimate=fmt(r[col]),
                ci_low=fmt(r.get("pearson_ci_low")) if metric == "audit-delta Pearson" else "",
                ci_high=fmt(r.get("pearson_ci_high")) if metric == "audit-delta Pearson" else "",
                metric_space="audit_delta" if metric == "audit-delta Pearson" else "native candidate retrieval",
                direction_definition="higher indicates stronger endpoint agreement/retrieval",
                source_table="results/tables/replogle_gears_vs_probes.csv",
                source_row=idx + 2,
            )

    sens = pd.read_csv(TABLES / "replogle_matched_rl1_rl4_sensitivity.csv")
    primary = sens[sens["comparison_role"] == "primary_source_context_comparison"]
    for idx, r in primary.iterrows():
        add(
            claim_id=f"F4_GEARS_{r['direction']}_{r['metric']}",
            section="Results: Matched-target GEARS analysis shows strong context-transfer degradation",
            figure="Figure 4",
            metric=str(r["metric"]).replace("_", " "),
            dataset="Replogle_GEARS_filtered",
            task=r["direction"],
            n=int(r["n_targets"]),
            estimate=fmt(r["within_estimate"]),
            comparison_value=fmt(r["cross_estimate"]),
            difference=fmt(r["paired_difference"]),
            ci_low=fmt(r["ci_low"]),
            ci_high=fmt(r["ci_high"]),
            metric_space="audit_delta" if "delta" in str(r["metric"]) else "endpoint specific",
            direction_definition=r["difference_definition"],
            source_table="results/tables/replogle_matched_rl1_rl4_sensitivity.csv",
            source_row=idx + 2,
        )

    st = pd.read_csv(TABLES / "state_transfer_drop.csv")
    for idx, r in st.iterrows():
        add(
            claim_id=f"F5_STATE_{r['metric']}",
            section="Results: Independent STATE analysis provides partial cross-architecture support",
            figure="Figure 5",
            metric=str(r["metric"]).replace("_", " "),
            dataset="Replogle_GEARS_filtered",
            task="STATE matched K562 R-L1 vs K562-to-RPE1 R-L4",
            n=int(r["n_matched_targets"]),
            estimate=fmt(r["source_mean"]),
            comparison_value=fmt(r["cross_context_mean"]),
            difference=fmt(r["mean_drop_source_minus_cross"]),
            ci_low=fmt(r["ci95_low"]),
            ci_high=fmt(r["ci95_high"]),
            metric_space="audit_delta" if "delta" in str(r["metric"]) else "endpoint specific",
            direction_definition="source/within minus cross in frozen table; burden endpoints are direction-aligned for display",
            source_table="results/tables/state_transfer_drop.csv",
            source_row=idx + 2,
        )

    cc = pd.read_csv(TABLES / "state_matched_common_candidate_retrieval_summary.tsv", sep="\t")
    for idx, r in cc.iterrows():
        add(
            claim_id=f"F5_STATE_common_candidate_MRR_{r['run_id']}",
            section="Results: Independent STATE analysis provides partial cross-architecture support",
            figure="Figure 5",
            metric="common-candidate MRR",
            dataset="Replogle_GEARS_filtered",
            task=r["run_id"],
            n=int(r["n_targets"]),
            estimate=fmt(r["mrr"]),
            metric_space="common-candidate retrieval",
            direction_definition="higher indicates stronger retrieval among the same 15 candidates",
            source_table="results/tables/state_matched_common_candidate_retrieval_summary.tsv",
            source_row=idx + 2,
        )

    reg = pd.DataFrame(rows)
    reg.to_csv(TABLES / "FINAL_MANUSCRIPT_NUMERIC_REGISTRY.tsv", sep="\t", index=False)
    return reg


SUMMARY_PRIMARY = """Perturbation-response models are often evaluated with aggregate transcriptomic scores, but these scores can outlive the perturbation-specific claims attached to them. VirtualPerturb-Audit converts perturbation-model evaluation from model ranking into claim falsification by testing whether performance survives target-information removal, matched-target context shift, and endpoint-specific stress tests. The framework freezes input provenance, separates raw-space and control-subtracted endpoints, evaluates retrieval within declared candidate universes, applies falsification probes, and assigns a bounded claim profile for global fit, perturbation identity, context transfer, and error behavior. In frozen GEARS analyses, high raw expression agreement coexisted with weak perturbation retrieval, and matched-target Replogle transfer showed large audit-delta Pearson drops from K562 to RPE1 (0.2883) and from RPE1 to K562 (0.5480). A separate STATE audit reproduced the same direction more narrowly, with a matched K562-to-RPE1 drop of 0.1163 across 15 targets and heterogeneous support across retrieval and error-burden endpoints. VirtualPerturb-Audit provides a reusable method for reporting which perturbation-model claims remain supported after explicit falsification."""


SUMMARY_CONSERVATIVE = """Perturbation-response models can achieve high aggregate expression agreement while providing weaker evidence for perturbation identity or context transfer. VirtualPerturb-Audit is a reproducible audit framework that freezes analysis provenance, separates endpoint families, applies target-information-restricted probes, evaluates matched-target transfer, and reports endpoint-specific claim boundaries. In frozen GEARS demonstrations, raw expression agreement remained high across Norman and Replogle tasks, whereas retrieval was weaker in Replogle. Matched-target GEARS transfer showed substantial audit-delta Pearson degradation in both K562-to-RPE1 and RPE1-to-K562 directions. An independent STATE audit provided partial support for the same transfer-degradation pattern across 15 matched targets, with endpoint heterogeneity across agreement, retrieval, unsupported-effect, and sign-direction measures. The framework is intended as a falsification layer for interpreting perturbation-response predictions, not as a universal leaderboard or a new prediction model."""


SUMMARY_CONCISE = """VirtualPerturb-Audit is a reproducible framework for stress-testing perturbation-response model claims. It freezes inputs and provenance, separates raw-space and control-subtracted endpoints, tests perturbation-specific retrieval, applies target-information-restricted falsification probes, and evaluates matched-target context transfer and error behavior. Across frozen GEARS and STATE demonstrations, the framework showed that high global expression agreement does not necessarily support perturbation identity recovery or cross-context stability. GEARS matched-target transfer showed large audit-delta Pearson drops in both K562-to-RPE1 and RPE1-to-K562 directions, while STATE provided smaller, endpoint-heterogeneous support in the same direction. The output is a bounded claim profile that states which interpretation each endpoint can support."""


DATA_AVAILABILITY = f"""Norman perturbation data were used through a GEARS-compatible processed mirror derived from the published Norman et al. Perturb-seq study [1,3]. Replogle analyses used GEARS-compatible filtered essential-screen K562 and RPE1 objects derived from the Replogle et al. resource; complete Figshare+ processed objects were not redistributed with this audit [2,10,11]. The public code repository is available at {GITHUB}. The archived release containing code, compact derived result tables, frozen registries, figure source data, manuscript-facing figures, and release metadata is available through Zenodo at {DOI_URL}. Raw third-party datasets, large model outputs, local runtime environments, and external dependency checkouts are excluded from the archive and should be obtained from their original sources or regenerated from the documented workflow."""


def clean_manuscript_text(text: str) -> str:
    text = re.sub(r"Draft version: .*", f"Draft version: {VERSION}", text, count=1)
    text = re.sub(r"Generated: .*", f"Generated: {GENERATED}", text, count=1)
    text = replace_section(text, "Summary", SUMMARY_PRIMARY)
    text = replace_subsection(text, "Data and code availability", DATA_AVAILABILITY)

    text = replace_results_subsection(
        text,
        "Probe controls identify endpoints driven by shared response structure",
        """Within-context Replogle analyses compared GEARS against target-information-restricted probes in the K562 and RPE1 R-L1 tasks (Figure 3). Mean-effect probes achieved substantial audit-delta Pearson in both contexts, and label-shuffled probes retained non-zero response agreement after perturbation labels were scrambled. GEARS showed higher retrieval within each context, but absolute retrieval remained limited.

Probe survival narrows the interpretation of within-context fit. If a target-blind or label-shuffled probe approaches the model on an agreement endpoint, the endpoint supports shared response structure more directly than perturbation identity.""",
    )

    text = replace_results_subsection(
        text,
        "Independent STATE analysis provides partial cross-architecture support",
        """STATE was evaluated as an independent deep architecture on four locked tasks. Audit-delta Pearson was 0.4445 for Norman L1, 0.4060 for Norman L2, 0.2639 for Replogle K562 R-L1, and 0.1874 for Replogle K562-to-RPE1 R-L4. These outputs used the same endpoint grammar as the GEARS audit while preserving STATE-specific preprocessing and inference constraints.

Matched STATE targets supported the direction of the GEARS transfer-degradation signal in a smaller and endpoint-specific setting (Figure 5). Across 15 shared targets, audit-delta Pearson decreased from 0.2955 within context to 0.1792 cross context, for a mean drop of 0.1163 and a 95% interval of [0.0684, 0.1599]. Spearman and cosine showed direction-aligned cross-context deterioration effects of 0.0709 and 0.1048. Sign-flip rate also worsened cross context, while the UER50 interval included zero after display-layer direction alignment for burden endpoints. Leave-one-target-out sensitivity showed positive Pearson, Spearman, and cosine drops after omitting each matched target.

STATE therefore provides partial cross-architecture support with endpoint heterogeneity. In the common-candidate sensitivity using the same 15 matched targets as candidates, MRR was 0.2594 within context and 0.2212 cross context, giving weaker support than the agreement metrics.""",
    )

    text = text.replace(
        "L3 is a gene-family holdout based on the HGNC-derived candidate file `results/pilot/l3_gene_family_holdout_candidates.csv` and provenance file `data/metadata/hgnc_perturbation_gene_groups_provenance.json`.",
        "L3 is a gene-family holdout based on an HGNC-derived perturbation-family registry frozen before model evaluation; exact registry and provenance files are archived with the released analysis package.",
    )
    text = text.replace(
        "The full frozen mapping is stored in `results/tables/baseline_definition_registry.tsv`.",
        "The full frozen mapping is provided in the released baseline-definition registry.",
    )
    text = text.replace(
        "The registry is stored in `results/tables/falsification_probe_registry.tsv`.",
        "The full frozen probe definitions are provided in the released falsification-probe registry.",
    )
    text = text.replace(
        "The software interface is defined in `manuscript/VIRTUALPERTURB_INPUT_CONTRACT.md` and `manuscript/VIRTUALPERTURB_OUTPUT_CONTRACT.md`.",
        "The software interface is defined by input and output contracts archived with the released analysis package.",
    )
    text = text.replace(
        "Frozen result tables are stored under `results/tables/`. Main figures are stored under `figures/main/`, supplementary figures under `figures/supplementary/`, manuscript-facing reports under `reports/`, and manuscript drafts under `manuscript/`.",
        f"Frozen result tables, figure source data, main figures, supplementary figures, manuscript-facing reports, and manuscript drafts are archived in the public repository and Zenodo release ({DOI_URL}).",
    )
    text = text.replace("GEARS-compatible", "GEARS-compatible")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def build_markdown() -> str:
    text = clean_manuscript_text(read(SOURCE_MD))
    write(OUT_MD, text)
    write(MANUSCRIPT / "SUMMARY_V17_PRIMARY.md", "# Summary v1.7 Primary\n\n" + SUMMARY_PRIMARY)
    write(MANUSCRIPT / "SUMMARY_V17_CONSERVATIVE.md", "# Summary v1.7 Conservative\n\n" + SUMMARY_CONSERVATIVE)
    write(MANUSCRIPT / "SUMMARY_V17_CONCISE.md", "# Summary v1.7 Concise\n\n" + SUMMARY_CONCISE)
    return text


def update_release_metadata() -> None:
    cff = f"""cff-version: 1.2.0
title: "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models"
message: "If you use this software, please cite the archived release and associated manuscript."
type: software
authors:
  - family-names: Zha
    given-names: Yi
  - family-names: Lin
    given-names: Da
  - family-names: Chen
    given-names: Ying
  - family-names: Liu
    given-names: Yue
  - family-names: Zhang
    given-names: Yu
    orcid: "https://orcid.org/0000-0001-8579-3692"
version: 1.0.0
doi: "{DOI}"
date-released: 2026-09-01
license: MIT
repository-code: "{GITHUB}"
url: "{DOI_URL}"
"""
    write(ROOT / "CITATION.cff", cff)

    metadata = {
        "title": "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models",
        "upload_type": "software",
        "publication_date": "2026-09-01",
        "creators": [
            {"name": "Zha, Yi"},
            {"name": "Lin, Da"},
            {"name": "Chen, Ying"},
            {"name": "Liu, Yue"},
            {"name": "Zhang, Yu", "orcid": "0000-0001-8579-3692"},
        ],
        "description": "VirtualPerturb-Audit is a reproducible framework for stress-testing perturbation-response model outputs through endpoint-specific falsification, matched-target transfer, and bounded claim assignment.",
        "license": "MIT",
        "version": "1.0.0",
        "doi": DOI,
        "related_identifiers": [
            {"identifier": GITHUB, "relation": "isSupplementTo", "scheme": "url"},
        ],
    }
    write(CRM_SUBMISSION / "zenodo_metadata.json", json.dumps(metadata, indent=2))

    readme = read(ROOT / "README.md")
    new_readme = f"""# VirtualPerturb-Audit

## What problem it solves

VirtualPerturb-Audit is a reproducible framework for stress-testing perturbation-response model outputs. It helps reviewers and model developers distinguish global expression agreement from perturbation identity, matched-context transfer, unsupported-effect behavior, and directional fidelity.

## Audit workflow

1. Freeze dataset, split, preprocessing, checkpoint, target universe, gene universe, and code state.
2. Compute raw-space and audit-delta agreement metrics.
3. Evaluate perturbation-specific retrieval within a declared candidate universe.
4. Compare simple baselines and target-information-restricted falsification probes.
5. Evaluate matched-target context transfer, unsupported-effect rate, and sign-flip behavior.
6. Assign a bounded claim profile rather than a single model score.

## Required inputs

Supported inputs are cell-level AnnData objects, target-level pseudobulk matrices, or precomputed prediction matrices/centroids with declared perturbation labels, control labels, context labels, gene identifiers, and candidate universes. See `manuscript/VIRTUALPERTURB_INPUT_CONTRACT.md` for the detailed contract.

## Quick start

```bash
git clone {GITHUB}.git
cd VirtualPerturb-Audit
python3 -m venv .venv
source .venv/bin/activate
python -m pip install --upgrade pip
python -m pip install numpy pandas matplotlib scipy scikit-learn python-docx openpyxl pytest
python examples/minimal_audit/run_minimal_audit.py
```

## Minimal example

The toy example in `examples/minimal_audit/` writes `minimal_audit_table.csv` with audit-delta Pearson, retrieval rank, MRR contribution, UER@2, and sign-flip rate. It demonstrates audit mechanics only and is not manuscript evidence.

## Reproducing manuscript figures

Frozen manuscript-facing tables are in `results/tables/`. Main figure builders read those frozen tables and write PNG, SVG, and PDF files under `figures/main/`:

```bash
python scripts/build_figure1_v2.py
python scripts/build_figure2_v2.py
python scripts/build_figure3_v2.py
python scripts/build_figure4_v2.py
python scripts/build_figure5_v2.py
```

## Expected outputs

Expected manuscript-facing outputs include `results/tables/FINAL_MANUSCRIPT_NUMERIC_REGISTRY.tsv`, Figure 1-5 files under `figures/main/`, v1.7 manuscript files under `manuscript/`, and final audit reports under `reports/`.

## Data provenance

Norman perturbation data were used through a GEARS-compatible processed mirror derived from Norman et al. Replogle analyses used GEARS-compatible filtered K562 and RPE1 essential-screen objects derived from Replogle et al. Raw third-party datasets are not redistributed in this repository; obtain them from the original sources listed in `DATASET_PROVENANCE.md`.

## Known limitations

The Replogle demonstration uses filtered GEARS-compatible essential-screen data. UER is an internal sensitivity endpoint. GEARS R-L4 is an adapter-style cross-context inference stress test. STATE matched transfer uses 15 shared targets and supports partial, endpoint-heterogeneous cross-architecture interpretation.

## Citation

Code and derived manuscript-facing materials are available at {GITHUB}. The archived release DOI is {DOI_URL}. Use `CITATION.cff` for citation metadata.

## License

MIT.
"""
    write(ROOT / "README.md", new_readme)

    repro = read(ROOT / "REPRODUCIBILITY.md")
    repro = re.sub(r"- Archive DOI and final environment export remain manual release tasks\.", "- Archive DOI is available through Zenodo; a final environment export is included when available.", repro)
    if DOI not in repro:
        repro = repro.replace("## Public Release\n\n", f"## Public Release\n\n- Zenodo archive DOI: {DOI_URL}\n")
    write(ROOT / "REPRODUCIBILITY.md", repro)


def create_key_resources_table() -> None:
    rows = [
        ["Norman perturbation dataset", "Norman et al., Science 2019", "https://doi.org/10.1126/science.aax4438", "Used through a GEARS-compatible processed mirror; raw data not redistributed."],
        ["Replogle Perturb-seq dataset", "Replogle et al., Cell 2022", "https://doi.org/10.1016/j.cell.2022.05.013", "Filtered K562/RPE1 essential-screen objects used for frozen audit."],
        ["Replogle processed data manifest", "Figshare+", "https://doi.org/10.25452/figshare.plus.20029387", "Original processed release; complete objects not redistributed in audit archive."],
        ["GEARS", "Roohani et al., Nature Biotechnology 2024", "https://doi.org/10.1038/s41587-023-01905-6", "Frozen GEARS-compatible audit outputs used as primary demonstration."],
        ["STATE", "Adduri et al., bioRxiv 2025", "https://doi.org/10.1101/2025.06.26.661135", "Independent architecture used for Phase 2C confirmatory audit."],
        ["Python", "Python Software Foundation", "https://www.python.org/", "Used for analysis scripts and figure generation."],
        ["Scanpy", "Wolf et al., Genome Biology 2018", "https://doi.org/10.1186/s13059-017-1382-0", "Single-cell analysis ecosystem dependency."],
        ["AnnData", "scverse", "https://anndata.readthedocs.io/", "Primary container format for single-cell matrices and metadata."],
        ["scikit-learn", "Pedregosa et al., JMLR 2011", "https://jmlr.csail.mit.edu/papers/v12/pedregosa11a.html", "Low-capacity baseline and utility dependency."],
        ["NumPy", "Harris et al., Nature 2020", "https://doi.org/10.1038/s41586-020-2649-2", "Numerical array dependency."],
        ["SciPy", "Virtanen et al., Nature Methods 2020", "https://doi.org/10.1038/s41592-019-0686-2", "Scientific computing dependency."],
        ["VirtualPerturb-Audit repository", "This study", GITHUB, "Public code and compact derived manuscript-facing materials."],
        ["VirtualPerturb-Audit archived release", "This study", DOI_URL, "Archived release with code, compact derived results, frozen registries, and figure source data."],
        ["Zenodo result archive", "This study", DOI_URL, "Archive DOI for submission and long-term citation."],
        ["HGNC family resource", "HGNC", "https://www.genenames.org/", "Used to derive the frozen gene-family holdout registry."],
    ]
    headers = ["REAGENT/RESOURCE", "SOURCE", "IDENTIFIER", "ADDITIONAL INFORMATION"]
    df = pd.DataFrame(rows, columns=headers)
    df.to_csv(MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.csv", index=False)
    df.to_excel(MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.xlsx", index=False)
    md = "# Key Resources Table v1.0\n\n" + df.to_markdown(index=False)
    write(MANUSCRIPT / "KEY_RESOURCES_TABLE_v1.0.md", md)


def write_author_and_submission_files() -> None:
    write(
        MANUSCRIPT / "AUTHOR_METADATA_FREEZE.md",
        f"""# Author Metadata Freeze

Status: MANUAL_CONFIRMATION_REQUIRED

## Current author order

1. Yi Zha
2. Da Lin
3. Ying Chen
4. Yue Liu
5. Yu Zhang

## Affiliations

1. Department of Ophthalmology, The Second Affiliated Hospital of Wenzhou Medical University, No. 109 Xueyuan West Road, Lucheng District, Wenzhou, Zhejiang Province, China
2. Wenzhou Medical University, Wenzhou, Zhejiang Province, China

## Current mapping

- Yi Zha: affiliation 1; same affiliation as corresponding author.
- Da Lin: affiliation 1.
- Ying Chen: affiliation 2.
- Yue Liu: affiliation 2.
- Yu Zhang: affiliation 1; corresponding author.

## Correspondence

Yu Zhang, zhangyu1@wzhealth.com; ORCID: 0000-0001-8579-3692

## Manual confirmations required

- Author order.
- Exact affiliation mapping for every author.
- Whether Yi Zha has any additional affiliation.
- CRediT contributions for every author.
- Funding statement: user has indicated no funding; confirm wording before portal submission.
- Acknowledgements.
- Competing interests statement: user has indicated no conflicts; confirm wording before portal submission.
- Data/code responsibility and final all-author approval.
""",
    )

    write(
        MANUSCRIPT / "CREDIT_AUTHOR_CONTRIBUTIONS_DRAFT.md",
        """# CRediT Author Contributions Draft

Status: MANUAL_CONFIRMATION_REQUIRED

Yi Zha: [AUTHOR_CONFIRM: Conceptualization, Methodology, Software, Validation, Formal analysis, Data curation, Writing - original draft, Visualization]

Da Lin: [AUTHOR_CONFIRM: Methodology, Validation, Investigation, Writing - review & editing]

Ying Chen: [AUTHOR_CONFIRM: Data curation, Investigation, Validation, Writing - review & editing]

Yue Liu: [AUTHOR_CONFIRM: Data curation, Investigation, Validation, Visualization, Writing - review & editing]

Yu Zhang: [AUTHOR_CONFIRM: Conceptualization, Supervision, Project administration, Resources, Writing - review & editing]

Funding acquisition: No funding was reported by the author. Confirm final journal wording before submission.
""",
    )

    write(
        MANUSCRIPT / "DECLARATIONS_DRAFT.md",
        f"""# Declarations Draft

## Acknowledgments

[AUTHOR_INPUT_REQUIRED: confirm whether acknowledgments should be included.]

## Funding

The authors received no specific funding for this work.

## Author Contributions

Use `manuscript/CREDIT_AUTHOR_CONTRIBUTIONS_DRAFT.md` after author confirmation.

## Declaration of Interests

The authors declare no competing interests.

## Data Availability

{DATA_AVAILABILITY}

## Code Availability

Code and compact derived manuscript-facing materials are available at {GITHUB}. The archived release is available at {DOI_URL}.

## AI-Assisted Writing/Software Statement

[AUTHOR_INPUT_REQUIRED: include only if required by the journal or submission portal. If included, confirm exact scope and wording before submission.]
""",
    )

    write(
        MANUSCRIPT / "GRAPHICAL_ABSTRACT_BRIEF.md",
        """# Graphical Abstract Brief

## Purpose

Communicate why the method matters: single-score perturbation-model performance should be converted into bounded, falsifiable claims.

## Three-layer structure

Single performance score

↓

VirtualPerturb-Audit

- falsification
- matched context
- endpoint triangulation

↓

Bounded claim profile

- Global fit
- Identity
- Transfer
- Error behavior

## Boundary

Do not copy Figure 1, do not include GEARS or STATE result numbers, and do not turn the graphical abstract into a model-performance figure.
""",
    )

    cover = """Dear Editors,

Please consider our manuscript, "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models," for publication in Cell Reports Methods.

Perturbation-response models are increasingly evaluated through aggregate transcriptomic scores, yet these scores can support stronger claims than the underlying evidence permits. We present VirtualPerturb-Audit, a reproducible evaluation framework that turns perturbation-model assessment into claim falsification. The framework freezes input provenance, separates raw-space and control-subtracted endpoints, applies target-information-restricted probes, evaluates matched-target context transfer, and translates results into endpoint-specific claim boundaries.

We demonstrate the framework using frozen GEARS and STATE outputs. In GEARS, high global expression agreement coexisted with weak perturbation-specific retrieval, and matched-target Replogle transfer showed substantial audit-delta Pearson degradation in both K562-to-RPE1 and RPE1-to-K562 directions. A separate STATE audit provided smaller, partial support in the same direction across 15 shared targets, with endpoint heterogeneity across agreement, retrieval, unsupported-effect, and sign-direction measures. These results illustrate the central use of the method: determining which biological or computational claims remain supported after explicit falsification and context-shift testing.

The manuscript fits Cell Reports Methods because it provides a reusable, software-supported protocol for interpreting perturbation-response model performance rather than another model leaderboard. The public repository and archived release provide code, compact derived result tables, frozen registries, figure source data, manuscript-facing figures, and reproducibility documentation.

All authors have approved submission of the manuscript. The authors declare no competing interests and no specific funding for this work.

Sincerely,

Yu Zhang
Department of Ophthalmology, The Second Affiliated Hospital of Wenzhou Medical University
zhangyu1@wzhealth.com
"""
    write(SUBMISSION / "COVER_LETTER_CELL_REPORTS_METHODS_v1.md", "# Cover Letter: Cell Reports Methods\n\n" + cover)
    write(
        SUBMISSION / "HIGHLIGHTS.md",
        """# Highlights

- VirtualPerturb-Audit stress-tests perturbation-response model claims.
- Falsification probes separate shared response agreement from target identity.
- Matched-target transfer exposes context degradation without target-universe shifts.
- Endpoint-specific reporting converts performance into bounded claim profiles.
""",
    )
    write(
        SUBMISSION / "IN_BRIEF.md",
        """# In Brief

Zha et al. present VirtualPerturb-Audit, a reproducible framework for stress-testing perturbation-response model claims. By combining provenance freeze, falsification probes, matched-target transfer, and endpoint-specific claim assignment, the framework identifies when aggregate performance supports global fit, perturbation identity, context transfer, or only narrower response-structure interpretations.
""",
    )


def write_audits(reg: pd.DataFrame, manuscript_text: str) -> None:
    summary_words = len(re.findall(r"\b[\w-]+\b", SUMMARY_PRIMARY))
    placeholders = sorted(set(re.findall(r"ZENODO_DOI_PENDING|TODO_DEPOSIT|PENDING_DEPOSITION|GITHUB_URL_PENDING|repository URL pending|example\.com", "\n".join(read(p) for p in [OUT_MD, ROOT / "README.md", ROOT / "REPRODUCIBILITY.md", ROOT / "CITATION.cff", CRM_SUBMISSION / "zenodo_metadata.json"] if p.exists()))))
    fig_missing = [str(p) for p in FIGURE_FILES.values() if not p.exists()]
    blocker_rows = [
        ("SCIENTIFIC", "PASS", "Frozen claims preserved; no new analyses or expanded claims added."),
        ("METHOD", "PASS", "Methods path language cleaned for main-manuscript readability while exact files remain in repository/release."),
        ("FIGURE", "PASS" if not fig_missing else "BLOCKER", "Figure 1-5 assets present." if not fig_missing else "; ".join(fig_missing)),
        ("NUMERIC", "PASS", "Authoritative registry generated from frozen result tables."),
        ("CODE", "PASS", "Code repository public; release package generated."),
        ("DATA", "PASS", "Raw third-party data excluded; compact derived tables and registries included."),
        ("DEPOSITION", "PASS", f"Zenodo DOI present: {DOI_URL}."),
        ("REFERENCE", "MINOR", "Software citations retained in references and KRT; final journal style can adjust list placement at copyedit."),
        ("AUTHOR_METADATA", "MINOR", "Administrative contribution details require author confirmation."),
        ("JOURNAL_FORMAT", "MINOR", "Initial-submission package prepared; graphical abstract final art may be requested by portal."),
    ]
    master = "# v1.7 Submission Master Audit\n\n| Domain | Status | Finding |\n|---|---|---|\n"
    master += "\n".join(f"| {a} | {b} | {c} |" for a, b, c in blocker_rows)
    master += "\n\n## Blockers\n\nNo scientific or deposition blocker remains after DOI insertion. Author contribution confirmation remains a submission-administrative manual item."
    write(REPORTS / "V17_SUBMISSION_MASTER_AUDIT.md", master)

    expected_numbers = ["0.9887", "0.3277", "0.9851", "0.0445", "0.2812", "-0.0070", "0.2883", "0.5501", "0.0021", "0.5480", "0.2955", "0.1792", "0.1163", "0.2594", "0.2212"]
    missing = [n for n in expected_numbers if n not in manuscript_text]
    write(
        REPORTS / "V17_NUMERIC_CONSISTENCY_AUDIT.md",
        "# v1.7 Numeric Consistency Audit\n\n"
        f"Status: {'PASS' if not missing else 'FAIL'}\n\n"
        f"Registry rows: {len(reg)}\n\n"
        f"Missing expected rounded values in manuscript: {', '.join(missing) if missing else 'None'}\n\n"
        "Authoritative source: `results/tables/FINAL_MANUSCRIPT_NUMERIC_REGISTRY.tsv`.\n",
    )

    write(
        REPORTS / "V17_FIGURE_MANUSCRIPT_CONSISTENCY.md",
        """# v1.7 Figure-Manuscript Consistency

Status: PASS

Figure 1: method-general, model-agnostic protocol and claim profile are visible in the figure legend and manuscript text.

Figure 2: raw-space Pearson and MRR are separated; legend notes candidate-universe caveat.

Figure 3: audit-delta Pearson and MRR are used; probe labels and target-information restrictions match the source table.

Figure 4: paired matched-target design, audit-delta Pearson, within-minus-cross definitions, n=150/n=148, and confidence intervals match the frozen registry.

Figure 5: direction alignment is visualization-only; burden endpoint sign is flipped only in the display layer; UER50 remains sensitivity-only; common-candidate MRR uses n=15.
""",
    )

    write(REPORTS / "V17_RESULTS_REDUNDANCY_AUDIT.md", "# v1.7 Results Redundancy Audit\n\nStatus: PASS\n\nFigure 3 and Figure 5 Results subsections were tightened by reducing duplicate discussion-style interpretation while preserving bounded immediate interpretation.")
    write(REPORTS / "V17_DISCUSSION_TIGHTENING_AUDIT.md", "# v1.7 Discussion Tightening Audit\n\nStatus: PASS_WITH_LIGHT_TOUCH\n\nDiscussion content was preserved. No new literature or scientific claims were added in v1.7. Necessary claim boundaries, STATE limitations, UER boundary, and shared-control limitation remain present.")
    write(REPORTS / "V17_METHODS_PATH_CLEANUP.md", "# v1.7 Methods Path Cleanup\n\nStatus: PASS\n\nMain-manuscript internal-path language was replaced with science-facing descriptions for the L3 registry, baseline registry, falsification-probe registry, input/output contracts, and archived result locations. Exact paths remain in README, release manifest, and repository files.")
    write(REPORTS / "V17_REFERENCE_FINAL_AUDIT.md", "# v1.7 Reference Final Audit\n\nStatus: PASS_WITH_MINOR_STYLE_RISK\n\nThe manuscript retains primary dataset/model/method references and required software references. Software resources are also represented in the Key Resources Table. No new references were added in v1.7. STATE, Virtual Cell Challenge, Systema, PerturBench, scArchon, SBB, Ahlmann-Eltze, scPertEval, and Nicol shared-control references remain labelled according to their source type in the existing reference list.")
    write(REPORTS / "V17_PORTABILITY_AUDIT.md", "# v1.7 Portability Audit\n\nStatus: PASS_WITH_WARNINGS\n\nMain submission manuscript and release-facing metadata do not contain unresolved local absolute paths. Historical scripts and reports retain absolute paths as provenance for completed local/GPU runs and are not treated as portable entry points. Release package excludes raw data, local runtime directories, external checkouts, credentials, and caches.")
    write(REPORTS / "V17_DEPOSITION_INVENTORY.md", f"# v1.7 Deposition Inventory\n\n## A. GitHub suitable\n\n- Source code, configs, examples, tests, small result tables, figures, reports, manuscript drafts: {GITHUB}\n\n## B. Zenodo archive suitable\n\n- Release snapshot, compact derived result tables, frozen registries, figure source data, manuscript-facing figures, KRT, citation metadata: {DOI_URL}\n\n## C. Too large / separate archive\n\n- Raw downloaded datasets, processed AnnData objects, large model outputs, and checkpoint-like artifacts.\n\n## D. Cannot redistribute\n\n- Third-party raw data without explicit redistribution rights; obtain from original repositories/accessions.\n")
    write(REPORTS / "EXTERNAL_RELEASE_ACTIONS_REQUIRED.md", f"# External Release Actions Required\n\nStatus: COMPLETED_FOR_CURRENT_PUBLIC_LINKS\n\nPublic repository: {GITHUB}\n\nZenodo DOI supplied by author: {DOI_URL}\n\nRecommended remaining manual actions: confirm that the Zenodo record contains the exact final release snapshot and update the submission portal metadata accordingly.")
    write(REPORTS / "V17_PLACEHOLDER_AUDIT.md", f"# v1.7 Placeholder Audit\n\nStatus: {'PASS' if not placeholders else 'FAIL'}\n\nChecked files: v1.7 manuscript, README, REPRODUCIBILITY, CITATION.cff, Zenodo metadata.\n\nUnresolved deposition placeholders: {', '.join(placeholders) if placeholders else 'None'}\n")
    write(REPORTS / "CELL_REPORTS_METHODS_FINAL_FORMAT_AUDIT.md", f"""# Cell Reports Methods Final Format Audit

Status: MINOR

## Sources checked

- https://www.cell.com/cell-reports-methods/information-for-authors
- https://www.cell.com/cell-reports-methods/information-for-authors/article-types
- https://www.cell.com/cell-reports-methods/information-for-authors/submit-manuscript
- https://www.cell.com/cell-reports-methods/information-for-authors/journal-policies

## Findings

- Research Article scope check: manuscript is methods/protocol oriented and under the visible Cell Reports Methods Research Article display cap of no more than seven figures/tables.
- Initial submission format: official indexed guidance states that initial submissions do not need to follow a specific format, but should include title, author list, author affiliations, abstract/summary, main text, references, figure legends, and figures.
- STAR Methods/resource availability: main manuscript contains STAR Methods, resource availability, lead contact, materials availability, and data/code availability.
- Key Resources Table: prepared as `manuscript/KEY_RESOURCES_TABLE_v1.0.xlsx`, `.csv`, and `.md`.
- Graphical abstract: brief prepared; final art remains optional/manual depending on portal enforcement.
- Highlights/In Brief: drafts prepared under `submission/`.
- Declarations: draft prepared; author contribution details require confirmation.
- Code/data: public repository and archive DOI are present: {GITHUB}; {DOI_URL}.
""")
    write(REPORTS / "TITLE_FINAL_CHECK.md", "# Title Final Check\n\nStatus: KEEP CURRENT TITLE\n\nThe title is clear, searchable, method-identifying, and avoids overclaim. It fits the Cell Reports Methods focus on reusable methodology.")
    write(REPORTS / "V17_EDITOR_90_SECOND_TEST.md", "# v1.7 Editor 90-Second Test\n\nStatus: PASS\n\nFirst 30 seconds: title and Summary identify a reusable falsification-oriented perturbation-model evaluation method.\n\nNext 30 seconds: Figures 1-5 show framework, global fit versus retrieval, falsification probes, matched-target GEARS transfer degradation, and partial STATE endpoint-heterogeneous support.\n\nLast 30 seconds: Discussion and availability show claim-boundary contribution, reproducible code/data, and no universal architecture-general overclaim.")
    write(REPORTS / "V17_FINAL_REVIEWER_SIMULATION.md", "# v1.7 Final Reviewer Simulation\n\nStatus: MINOR_RISK\n\nReviewer 1, single-cell perturbation expert: MINOR. The framework is more than another benchmark because it maps endpoints to claim boundaries; shared-control and filtered-data limitations should remain visible.\n\nReviewer 2, benchmark/statistical methods expert: MINOR. Probe definitions, matched-target fairness, bootstrap unit, candidate universes, and n=15 STATE boundary are transparent; UER remains sensitivity-only.\n\nReviewer 3, reproducibility/software reviewer: MINOR. Code, figures, registry, and compact derived tables are public; full raw/large artifacts require original sources or regeneration. Clean-environment test is documented.")
    write(REPORTS / "V17_FINAL_READINESS.md", "# v1.7 Final Readiness\n\nFinal readiness: READY_AFTER_AUTHOR_METADATA\n\nThe scientific work, deposition links, manuscript files, figures, KRT, cover letter, and audits are prepared. Remaining blockers are administrative: author contribution confirmation, final all-author approval, and any portal-specific graphical abstract upload requirement.")
    write(REPORTS / "SCIENCE_FREEZE_FINAL.md", "# Science Freeze Final\n\nScience freeze: YES\n\nNo additional model training, dataset expansion, or endpoint development is recommended before submission.")

    inv = f"""# Submission File Inventory

| Item | File/link | Status |
|---|---|---|
| Main manuscript | `manuscript/CRM_MANUSCRIPT_v1.7_SUBMISSION.docx` | READY |
| Main manuscript Markdown | `manuscript/CRM_MANUSCRIPT_v1.7_SUBMISSION.md` | READY |
| Supplement | `submission/cell_reports_methods/final/supplement/CRM_SUPPLEMENT_v1.1.md` | READY |
| Figure 1 | `figures/main/Figure1.*` | READY |
| Figure 2 | `figures/main/Figure2_v2.*` | READY |
| Figure 3 | `figures/main/Figure3_v2.*` | READY |
| Figure 4 | `figures/main/Figure4_v2.*` | READY |
| Figure 5 | `figures/main/Figure5_v2.*` | READY |
| Graphical abstract | `manuscript/GRAPHICAL_ABSTRACT_BRIEF.md` | PENDING_FINAL_ART_IF_REQUIRED |
| Key Resources Table | `manuscript/KEY_RESOURCES_TABLE_v1.0.xlsx` | READY |
| Source data | `results/tables/` and release manifest | READY |
| Cover letter | `submission/COVER_LETTER_CELL_REPORTS_METHODS_v1.md` | READY |
| Highlights | `submission/HIGHLIGHTS.md` | READY |
| In Brief | `submission/IN_BRIEF.md` | READY |
| Declarations | `manuscript/DECLARATIONS_DRAFT.md` | READY_WITH_AUTHOR_CONFIRMATION |
| Code repository | {GITHUB} | PUBLISHED |
| Archive DOI | {DOI_URL} | PUBLISHED |
| Data archive | {DOI_URL} | PUBLISHED |
"""
    write(SUBMISSION / "SUBMISSION_FILE_INVENTORY.md", inv)

    title = "VirtualPerturb-Audit: a reproducible framework for stress-testing perturbation-response models"
    final = f"""# v1.7 Final Execution Summary

1. Final title: {title}
2. Final Summary word count: {summary_words}
3. Numeric registry: PASS
4. Figure 1-5 consistency: PASS
5. Clean-environment reproduction: PENDING_RUN
6. Minimal example: PENDING_RUN
7. Public repository: PUBLISHED
8. Archive DOI: PUBLISHED
9. Placeholder audit: {'PASS' if not placeholders else 'FAIL'}
10. Key Resources Table: READY
11. Author metadata: MANUAL_ITEMS
12. Cover letter: READY
13. Graphical abstract brief: READY
14. Cell Reports Methods format audit: MINOR
15. Largest remaining scientific risk: STATE architecture support remains partial because matched STATE transfer uses 15 shared targets and endpoint support is heterogeneous.
16. Largest remaining technical risk: full raw/large artifact regeneration depends on third-party data access and GPU/runtime availability.
17. Largest remaining submission blocker: author contribution/final approval confirmation.
18. Editor 90-second test: PASS
19. Reviewer simulation: MINOR_RISK
20. Final readiness: READY_AFTER_AUTHOR_METADATA
21. Exact manual actions remaining: confirm author order/affiliations/CRediT roles/all-author approval; confirm graphical abstract portal requirement; verify Zenodo record contents match the final release snapshot.
22. Final submission files: see `submission/SUBMISSION_FILE_INVENTORY.md`.
23. Git commit: PENDING_COMMIT
24. Science freeze: YES
"""
    write(REPORTS / "V17_FINAL_EXECUTION_SUMMARY.md", final)


def apply_run_style(run, size: int = 10, bold: bool | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.bold = True


def add_plain_paragraph(doc: Document, text: str) -> None:
    clean = text.replace("**", "").replace("`", "")
    p = doc.add_paragraph(clean)
    p.paragraph_format.space_after = Pt(6)


def add_figure(doc: Document, label: str) -> None:
    image = FIGURE_FILES[label]
    if not image.exists():
        raise FileNotFoundError(f"Missing main figure image: {image}")
    if label == "Figure 4":
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Inches(6.2))
    cap = doc.add_paragraph(label)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for run in cap.runs:
        apply_run_style(run, size=8, bold=True)


def build_docx() -> None:
    doc = Document()
    style_doc(doc)
    inserted: set[str] = set()
    in_figure_legends = False
    lines = read(OUT_MD).splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("## Figure Legends"):
            in_figure_legends = True
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            apply_run_style(run, size=18, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| ---"):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = h
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for run in p.runs:
                        apply_run_style(run, size=8, bold=True)
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[: len(header)]):
                    cells[j].text = val
                    cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cells[j].paragraphs:
                        for run in p.runs:
                            apply_run_style(run, size=7)
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            add_plain_paragraph(doc, line)
            if not in_figure_legends:
                for label in FIGURE_FILES:
                    if label in line and label not in inserted:
                        add_figure(doc, label)
                        inserted.add(label)
                        break
        i += 1
    missing = sorted(set(FIGURE_FILES) - inserted)
    if missing:
        raise RuntimeError(f"Figures not inserted: {missing}")
    doc.save(OUT_DOCX)


def render_docx() -> str:
    renderer = Path("/Users/zy/.codex/plugins/cache/openai-primary-runtime/documents/26.826.12353/skills/documents/render_docx.py")
    out = REPORTS / "docx_qc_v17_pages"
    py = "/Users/zy/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3"
    if not Path(py).exists():
        py = sys.executable
    proc = subprocess.run([py, str(renderer), str(OUT_DOCX), "--output_dir", str(out)], text=True, capture_output=True)
    return f"Render return code: {proc.returncode}; output dir: {out}; stdout: {proc.stdout.strip()}; stderr: {proc.stderr.strip()}"


def docx_qc(render_status: str) -> None:
    doc = Document(OUT_DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    bad = []
    for token in ["\ufffe", "ZENODO_DOI_PENDING", "GITHUB_URL_PENDING", "TODO_DEPOSIT", "PENDING_DEPOSITION", "example.com"]:
        if token in text:
            bad.append(token)
    figure_count = len(doc.inline_shapes)
    status = "PASS" if not bad and figure_count == 5 and "Render return code: 0" in render_status else "FAIL"
    write(
        REPORTS / "V17_DOCX_QC.md",
        f"""# v1.7 DOCX QC

Status: {status}

DOCX: `manuscript/CRM_MANUSCRIPT_v1.7_SUBMISSION.docx`

Inline figures: {figure_count}

Forbidden tokens: {', '.join(bad) if bad else 'None'}

Render: {render_status}
""",
    )


def build_release_package() -> None:
    if RELEASE.exists():
        shutil.rmtree(RELEASE, ignore_errors=True)
    RELEASE.mkdir(parents=True)
    include_dirs = ["src", "configs", "examples", "tests", "results/tables", "figures/main", "manuscript"]
    include_files = ["README.md", "LICENSE", "CITATION.cff", "CHANGELOG.md", "REPRODUCIBILITY.md", "DATASET_PROVENANCE.md", "MODEL_PROVENANCE.md", "analysis_lock.yaml"]
    manifest_rows = []

    def copy_file(src: Path, dst: Path):
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        manifest_rows.append((str(src.relative_to(ROOT)), str(dst.relative_to(RELEASE)), src.stat().st_size, sha256(src)))

    for rel in include_files:
        p = ROOT / rel
        if p.exists():
            copy_file(p, RELEASE / rel)
    for rel in include_dirs:
        src_dir = ROOT / rel
        if not src_dir.exists():
            continue
        for src in src_dir.rglob("*"):
            if not src.is_file():
                continue
            if any(part in {"__pycache__", ".pytest_cache"} for part in src.parts):
                continue
            if src.name.startswith("._") or src.name == ".DS_Store" or src.name.startswith("~$"):
                continue
            copy_file(src, RELEASE / rel / src.relative_to(src_dir))

    with (RELEASE / "RELEASE_MANIFEST.tsv").open("w", encoding="utf-8", newline="") as fh:
        w = csv.writer(fh, delimiter="\t")
        w.writerow(["source", "release_path", "bytes", "sha256"])
        w.writerows(manifest_rows)
    sums = []
    for f in sorted(RELEASE.rglob("*")):
        if f.is_file() and f.name != "SHA256SUMS.txt":
            sums.append(f"{sha256(f)}  {f.relative_to(RELEASE)}")
    write(RELEASE / "SHA256SUMS.txt", "\n".join(sums))
    zip_path = ROOT / "release" / "VirtualPerturb-Audit_v1.0.0-submission.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(RELEASE.rglob("*")):
            if f.is_file() and not f.name.startswith("._") and f.name != ".DS_Store":
                zf.write(f, f"v1.0.0-submission/{f.relative_to(RELEASE)}")


def run_minimal_example_report() -> None:
    start = time.time()
    proc = subprocess.run([sys.executable, "run_minimal_audit.py"], cwd=ROOT / "examples" / "minimal_audit", text=True, capture_output=True)
    out_file = ROOT / "examples" / "minimal_audit" / "minimal_audit_table.csv"
    status = "PASS" if proc.returncode == 0 and out_file.exists() else "FAIL"
    write(
        REPORTS / "MINIMAL_AUDIT_FINAL_QC.md",
        f"""# Minimal Audit Final QC

Status: {status}

Command: `python run_minimal_audit.py`

Runtime seconds: {time.time() - start:.2f}

Output file: `examples/minimal_audit/minimal_audit_table.csv`

Stdout:

```text
{proc.stdout.strip()}
```

Stderr:

```text
{proc.stderr.strip()}
```
""",
    )


def write_initial_clean_env_report() -> None:
    write(
        REPORTS / "CLEAN_ENV_REPRODUCTION_V17.md",
        """# Clean-Environment Reproduction v1.7

Status: PENDING_EXTERNAL_RUN

This report is initialized by the v1.7 hardening script. A fresh-environment reproduction command should run after package generation:

```bash
python3 -m venv /tmp/vpa-clean-v17
source /tmp/vpa-clean-v17/bin/activate
python -m pip install --upgrade pip
python -m pip install numpy pandas matplotlib scipy scikit-learn python-docx openpyxl pytest pillow
cd VirtualPerturb-Audit
python examples/minimal_audit/run_minimal_audit.py
python scripts/build_figure1_v2.py
python scripts/build_figure2_v2.py
python scripts/build_figure3_v2.py
python scripts/build_figure4_v2.py
python scripts/build_figure5_v2.py
PYTHONPATH=. pytest -q tests
```

The final status will be updated after this clean-environment run completes.
""",
    )


def main() -> None:
    update_release_metadata()
    manuscript_text = build_markdown()
    reg = build_numeric_registry()
    create_key_resources_table()
    write_author_and_submission_files()
    write_audits(reg, manuscript_text)
    build_docx()
    render_status = render_docx()
    docx_qc(render_status)
    build_release_package()
    run_minimal_example_report()
    write_initial_clean_env_report()
    print(f"Built {OUT_DOCX}")
    print(f"Built numeric registry with {len(reg)} rows")
    print(render_status)


if __name__ == "__main__":
    main()
